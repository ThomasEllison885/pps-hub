"""
PPS Site Visit Report — python-docx builder
Called from hub app.py with a data dict, returns BytesIO
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ROW_HEIGHT_RULE

from pps_brand import (
    DARK_BLUE, BLUE, LIGHT_BG, GRAY, MID_GRAY, WHITE,
    DARK_BLUE_HEX, BLUE_HEX, BORDER_HEX,
    LOGO_PATH, TAGLINE_PATH,
)

BORDER_C = BORDER_HEX


def _border(para, side='bottom', color=DARK_BLUE_HEX, size=8, space=2):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(size))
    b.set(qn('w:space'), str(space))
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)


def _run(para, text, bold=False, italic=False, color=None, size=10):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = 'Arial'; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r


def _spacing(para, before=0, after=4):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color=BORDER_C):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_table_borders(table, color=BORDER_C):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _cell_margins(cell, top=55, bottom=55, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _set_col_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _sec_head(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=10, after=4)
    p.paragraph_format.keep_with_next = True
    _border(p, 'bottom', '004C8C', 8, 2)
    _run(p, text.upper(), bold=True, color=DARK_BLUE, size=10)


def _field_row(table, label, value, lw=2500, vw=6860):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    _set_col_width(lc, lw); _set_col_width(vc, vw)
    _shade_cell(lc, 'EBF6FC')
    _set_cell_borders(lc); _set_cell_borders(vc)
    _cell_margins(lc); _cell_margins(vc)
    lp = lc.paragraphs[0]; vp = vc.paragraphs[0]
    _spacing(lp, 0, 0); _spacing(vp, 0, 0)
    _run(lp, label, bold=True, color=DARK_BLUE, size=8.5)
    if value:
        _run(vp, value, color=GRAY, size=9)
    else:
        _run(vp, '', color=GRAY, size=9)


def _yesno_row(table, label, value, lw=2500, vw=6860):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    _set_col_width(lc, lw); _set_col_width(vc, vw)
    _shade_cell(lc, 'EBF6FC')
    _set_cell_borders(lc); _set_cell_borders(vc)
    _cell_margins(lc); _cell_margins(vc)
    lp = lc.paragraphs[0]; vp = vc.paragraphs[0]
    _spacing(lp, 0, 0); _spacing(vp, 0, 0)
    _run(lp, label, bold=True, color=DARK_BLUE, size=8.5)
    # Show checked value
    if value and value.lower() in ['yes','true','1']:
        _run(vp, '☑  Yes          ☐  No', color=GRAY, size=9)
    elif value and value.lower() in ['no','false','0']:
        _run(vp, '☐  Yes          ☑  No', color=GRAY, size=9)
    else:
        _run(vp, '☐  Yes          ☐  No', color=MID_GRAY, size=9)


def _status_row(table, label, value, options, lw=3800, vw=5560):
    row = table.add_row()
    lc, vc = row.cells[0], row.cells[1]
    _set_col_width(lc, lw); _set_col_width(vc, vw)
    _set_cell_borders(lc); _set_cell_borders(vc)
    _cell_margins(lc); _cell_margins(vc)
    lp = lc.paragraphs[0]; vp = vc.paragraphs[0]
    _spacing(lp, 0, 0); _spacing(vp, 0, 0)
    _run(lp, label, color=GRAY, size=9)
    parts = []
    for opt in options:
        checked = value and value.lower() == opt.lower().replace(' ','_')
        parts.append(('☑  ' if checked else '☐  ') + opt)
    _run(vp, '     '.join(parts), color=MID_GRAY, size=9)


def _header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    hp = header.add_paragraph()
    _spacing(hp, 0, 6)
    _border(hp, 'bottom', '004C8C', 8, 3)
    if os.path.exists(LOGO_PATH):
        r = hp.add_run()
        r.add_picture(LOGO_PATH, width=Inches(0.95))
    _run(hp, '    ', size=12)
    _run(hp, 'Site Visit & Inspection Report', bold=True, color=DARK_BLUE, size=12)
    _run(hp, '    \u2014    Pure Property Solutions', color=MID_GRAY, size=10)


def _footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    fp = footer.add_paragraph()
    _spacing(fp, 4, 0)
    _border(fp, 'top', '0096D6', 4, 3)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(TAGLINE_PATH):
        r = fp.add_run()
        r.add_picture(TAGLINE_PATH, width=Inches(2.8))
    else:
        _run(fp, 'The Pure Way: Trust. Quality. Results.\u2122', italic=True, color=MID_GRAY, size=9)


def build_site_visit(data):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

    _header(doc)
    _footer(doc)

    # Title
    tp = doc.add_paragraph()
    _spacing(tp, 0, 2)
    _run(tp, 'Site Visit & Inspection Report', bold=True, color=DARK_BLUE, size=18)

    sp2 = doc.add_paragraph()
    _spacing(sp2, 0, 10)
    _run(sp2, 'Pure Property Solutions  \u2014  Complete all sections during or immediately after site visit.',
         italic=True, color=MID_GRAY, size=8)

    # ── 1. PROJECT INFORMATION ─────────────────────────────────────────
    _sec_head(doc, '1. Project Information')
    t1 = doc.add_table(rows=0, cols=2)
    _set_table_borders(t1)
    _field_row(t1, 'Property Name',       data.get('property_name', ''))
    _field_row(t1, 'Property Address',     data.get('property_address', ''))
    _field_row(t1, 'Date of Visit',        data.get('visit_date', ''))
    _field_row(t1, 'Time on Site',         data.get('visit_time', ''))
    _field_row(t1, 'Visited By',           data.get('visited_by', ''))
    _field_row(t1, 'PO / Proposal Number', data.get('po_number', ''))

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # ── 2. TRADE PARTNER ───────────────────────────────────────────────
    _sec_head(doc, '2. Trade Partner & Crew')
    t2 = doc.add_table(rows=0, cols=2)
    _set_table_borders(t2)
    tp_present = data.get('trade_partner_present', '')
    _yesno_row(t2, 'Trade Partner on site?', tp_present)
    _field_row(t2, 'Company / Trade',      data.get('trade_partner_company', ''))
    _field_row(t2, 'Crew Lead & Contact',  data.get('crew_lead', ''))
    _field_row(t2, 'Number of Workers',    data.get('crew_count', ''))

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # ── 3. SCOPE & QUALITY INSPECTION ──────────────────────────────────
    _sec_head(doc, '3. Scope & Quality Inspection')
    t3 = doc.add_table(rows=0, cols=3)
    _set_table_borders(t3)

    # Header row
    hr = t3.add_row()
    for cell, text, w in [(hr.cells[0],'INSPECTION ITEM',5200),
                           (hr.cells[1],'STATUS',2160),
                           (hr.cells[2],'NOTES',2000)]:
        _shade_cell(cell, 'EBF6FC')
        _set_cell_borders(cell)
        _cell_margins(cell)
        _set_col_width(cell, w)
        p = cell.paragraphs[0]; _spacing(p, 0, 0)
        _run(p, text, bold=True, color=DARK_BLUE, size=8)

    checklist = data.get('checklist', [])
    checklist_labels = [
        'Scope being followed as approved',
        'Materials match proposal specifications',
        'Work quality meets PPS standards',
        'Safety protocols in place (PPE, fall protection)',
        'Site cleanliness maintained',
        'Property protected (landscaping, vehicles, windows)',
        'Damage to property unrelated to scope observed',
    ]
    cl_map = {item.get('label', ''): item for item in checklist}

    for label in checklist_labels:
        item = cl_map.get(label, {})
        status = item.get('status', '')
        notes  = item.get('notes', '')
        row = t3.add_row()
        ic, sc, nc = row.cells[0], row.cells[1], row.cells[2]
        for cell, w in [(ic,5200),(sc,2160),(nc,2000)]:
            _set_cell_borders(cell); _cell_margins(cell); _set_col_width(cell, w)
        ip = ic.paragraphs[0]; _spacing(ip, 0, 0)
        _run(ip, label, color=GRAY, size=9)
        sp_ = sc.paragraphs[0]; _spacing(sp_, 0, 0)
        sp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if status == 'yes':    sym = '☑ Yes  ☐ No  ☐ N/A'
        elif status == 'no':   sym = '☐ Yes  ☑ No  ☐ N/A'
        elif status == 'na':   sym = '☐ Yes  ☐ No  ☑ N/A'
        else:                   sym = '☐ Yes  ☐ No  ☐ N/A'
        _run(sp_, sym, color=MID_GRAY, size=8.5)
        np_ = nc.paragraphs[0]; _spacing(np_, 0, 0)
        _run(np_, notes, color=GRAY, size=8.5)

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # ── 4. ONSITE CONTACTS ─────────────────────────────────────────────
    _sec_head(doc, '4. Onsite Contacts')
    t4 = doc.add_table(rows=0, cols=2)
    _set_table_borders(t4)
    met = data.get('met_with_staff', '')
    _yesno_row(t4, 'Met with onsite staff, supervisor, or manager?', met)
    _field_row(t4, 'Contact Name & Title',  data.get('staff_contact', ''))
    _field_row(t4, 'Topics Discussed',      data.get('topics_discussed', ''))
    complaints = data.get('complaints_received', '')
    _yesno_row(t4, 'Resident complaints received or reported?', complaints)
    _field_row(t4, 'Complaint Details',     data.get('complaint_details', ''))

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # ── 5. PROJECT PROGRESS ────────────────────────────────────────────
    _sec_head(doc, '5. Project Progress')
    t5 = doc.add_table(rows=0, cols=2)
    _set_table_borders(t5)
    _status_row(t5, 'Overall project status',
                data.get('overall_status',''),
                ['On Track','Minor Concern','Issue'])
    _status_row(t5, 'Work quality',
                data.get('quality_status',''),
                ['Excellent','Acceptable','Needs Attention'])
    _status_row(t5, 'Schedule / timeline',
                data.get('schedule_status',''),
                ['On Schedule','Behind','At Risk'])

    doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # Observations
    t6 = doc.add_table(rows=0, cols=1)
    _set_table_borders(t6)
    hr2 = t6.add_row()
    hc = hr2.cells[0]
    _shade_cell(hc, 'EBF6FC'); _set_cell_borders(hc); _cell_margins(hc)
    _set_col_width(hc, 9360)
    hp2 = hc.paragraphs[0]; _spacing(hp2, 0, 0)
    _run(hp2, 'Observations & Recommendations:', bold=True, color=DARK_BLUE, size=8.5)

    obs_row = t6.add_row()
    obs_cell = obs_row.cells[0]
    _set_cell_borders(obs_cell); _cell_margins(obs_cell, 60, 60, 100, 100)
    _set_col_width(obs_cell, 9360)
    obs_p = obs_cell.paragraphs[0]; _spacing(obs_p, 0, 0)
    _run(obs_p, data.get('observations', ''), color=GRAY, size=9)
    # Set min row height for notes area
    tr = obs_row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '1000')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    # Photo + next visit
    ck_row = t6.add_row()
    ck_cell = ck_row.cells[0]
    _set_cell_borders(ck_cell); _cell_margins(ck_cell)
    _set_col_width(ck_cell, 9360)
    ck_p = ck_cell.paragraphs[0]; _spacing(ck_p, 0, 0)
    photos = data.get('photos_taken', '')
    nv     = data.get('next_visit_date', '')
    photo_sym = '☑' if photos else '☐'
    nv_sym    = '☑' if nv else '☐'
    nv_txt    = f'  Date: {nv}' if nv else '  Date: ___________________'
    _run(ck_p, f'{photo_sym}  Photo documentation taken          {nv_sym}  Next visit scheduled{nv_txt}',
         color=MID_GRAY, size=9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
