import os
import re
import json
import time
import secrets
import threading
import base64
from io import BytesIO
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file, send_from_directory, make_response, g, has_request_context
import pipeline_board
import office_ops
import insurance_compliance
import crm_contact_sync
import estimate_assignments
import weekly_recap
import password_campaign
import system_state
import training_overlay
import dashboard_summary
import db_pool
import db_ddl
from admin_feed import merge_activity
from werkzeug.exceptions import HTTPException
from psc_training_data import (
    PSC_TRAINING_META, PSC_TRAINING_MANAGER, get_training_curriculum,
    read_curriculum as psc_read_curriculum,
    get_week_checkin_questions as psc_get_week_checkin_questions,
    get_week_labels as psc_get_week_labels,
    get_all_item_ids, count_trackable_items,
    PSC_ROLEPLAY_SCENARIOS, PSC_ROLEPLAY_GRADER_RULES,
    get_roleplay_scenario, get_roleplay_week_links, get_roleplay_sales_links,
    get_suggested_roleplay_ids, segment_color,
    ROLEPLAY_DAILY_GRADE_LIMIT, ROLEPLAY_DAILY_TURN_LIMIT,
)
from pm_training_data import (
    PM_TRAINING_META, PM_TRAINING_MANAGER, get_pm_training_curriculum,
    get_pm_training_item_ids, count_pm_trackable_items, get_pm_week_item_ids,
    get_pm_week_checkin_questions,
)
from werkzeug.security import generate_password_hash, check_password_hash
import psycopg2
from psycopg2.extras import RealDictCursor
from auth_helpers import (
    HUB_PUBLIC_URL, PROPOSAL_URL, LOGIN_LOCKOUT_MINUTES, MAX_LOGIN_FAILURES,
    safe_next_url, client_ip, record_login_attempt, is_login_locked, login_lock_map,
    clear_login_failures,
    generate_sso_code, exchange_sso_code,
    create_password_reset_token, peek_password_reset_token,
    consume_password_reset_token, reset_url_for_token,
)
import ask_pps
def _load_dotenv():
    """Load .env into os.environ (keys already set in the environment win)."""
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    if not os.path.isfile(path):
        return
    with open(path, encoding='utf-8') as env_file:
        for line in env_file:
            line = line.strip()
            if not line or line.startswith('#') or '=' not in line:
                continue
            key, _, val = line.partition('=')
            key, val = key.strip(), val.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = val


_load_dotenv()

app = Flask(__name__)
_secret = os.environ.get('SECRET_KEY', '').strip()
if not _secret:
    print(
        'WARNING: SECRET_KEY is not set — hub will run but you should add SECRET_KEY '
        'in Render Environment (generate a random 32+ char string).'
    )
    _secret = 'pps-hub-unset-secret-key'
app.secret_key = _secret

INTERNAL_API_KEY = os.environ.get('INTERNAL_API_KEY', '').strip()
CLAUDE_API_KEY = os.environ.get('CLAUDE_API_KEY', '').strip()
CLAUDE_MODEL = 'claude-sonnet-4-6'
if not INTERNAL_API_KEY:
    print(
        'WARNING: INTERNAL_API_KEY is not set — proposal SSO and internal APIs '
        'will not work until you add the same key on hub and proposal services.'
    )
DATABASE_URL = os.environ.get('DATABASE_URL', '')
MAX_DOCUMENT_BYTES = 10 * 1024 * 1024  # 10 MB per file
VAULT_STORAGE_LIMIT_BYTES = int(os.environ.get('VAULT_STORAGE_LIMIT_MB', '512')) * 1024 * 1024

_IS_DEBUG = os.environ.get('FLASK_DEBUG', '').lower() in ('1', 'true', 'yes')
app.config.update(
    SESSION_COOKIE_SECURE=not _IS_DEBUG,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    # 30 days, idle — raised from 8 hours 2026-08-21. Flask re-signs the cookie
    # on every request (SESSION_REFRESH_EACH_REQUEST below), so this is an idle
    # window, not an absolute one. At 8 hours that sounded generous, but every
    # night is longer than eight hours: a PM who last used the Hub at 4pm was
    # signed out before he opened it next morning, and the sign-in he hit was a
    # 13-name dropdown plus a password, on a phone, on a job site.
    PERMANENT_SESSION_LIFETIME=timedelta(days=30),
    # Already Flask's default — pinned because the 30-day window above is only
    # tolerable while it stays an IDLE timeout. Flipping this to False silently
    # turns it into "log everyone out 30 days after they first signed in."
    SESSION_REFRESH_EACH_REQUEST=True,
)

GENERIC_API_ERROR = 'Something went wrong. Please try again or contact Thomas.'
GENERIC_DOWNLOAD_ERROR = 'Could not generate the file. Please try again or contact Thomas.'


def _log_exception(exc, context=''):
    import traceback
    label = f' ({context})' if context else ''
    print(f'Error{label}: {exc}')
    traceback.print_exc()


def _api_error(exc, status=500, **extra):
    _log_exception(exc)
    payload = {'error': GENERIC_API_ERROR}
    payload.update(extra)
    return jsonify(payload), status


_JSON_API_PATHS = frozenset({
    '/analyze-diff',
    '/submit-diff',
})


def _wants_json_response():
    if request.path.startswith('/api/') or request.path in _JSON_API_PATHS:
        return True
    if request.is_json:
        return True
    accept = request.accept_mimetypes.best_match(['application/json', 'text/html'])
    return accept == 'application/json'


# ── ACCESS TIERS ────────────────────────────────────────────────────────────────
#
# Rewritten 2026-08-21 (Thomas). There used to be NINE independent access
# mechanisms — role, proposal_access, ppm_access, team_view, team_view_scope,
# CONTACTS_USER_KEYS/ROLES, BOARD_ACCESS/BOARD_ACCESS_ALL, OFFICE_OPS_USER_KEYS,
# CURATORS, and the two training-manager constants — spread across four files.
# Answering "can Ben see this?" meant checking up to four of them. Two symptoms
# of how far that drifted: `ppm_access` was set True on all 13 people and read
# nowhere, and `proposal_access` stopped meaning what its name says (Jordan,
# Phil and Trey carry 'all' from old grants nobody intends), which is why
# Pipeline Board had to build a second roster to route around it.
#
# Now there is one axis — `tier` — and three values:
#
#   owner       Thomas. /admin, pricing defaults, password resets, feedback
#               inbox, vault delete, proposal diffs.
#   leadership  Stephanie, Tony, Trey. Office Ops (Numbers + Compliance),
#               PSC + PM training oversight, Ask PPS curation.
#   team        Everyone else. Every tool, every pipeline board, all history,
#               Team View. Deliberately unrestricted — Thomas 2026-08-21:
#               "I don't really want/need a lot of restrictions."
#
# The important distinction: a tier is what you may SEE. It is separate from
# ASSIGNMENT — which pipeline board opens by default, whose name prefills on a
# proposal. Assignment lives in pipeline_board.PRIMARY_PM_FOR_CONSULTANT and
# grants nothing. Most of the nine old mechanisms were assignment wearing a
# permission costume; that confusion is what this split removes.
#
# `role` survives, but only to describe WHAT SOMEONE DOES (consultant / pm /
# office_manager / admin) — proposal number initials, PSC training eligibility,
# dashboard framing, recap grouping. It no longer grants anything on its own
# except the legacy `role == 'admin'` owner checks, which are equivalent to
# tier 'owner' because Thomas is the only admin. See has_tier() below.

# Definitions live in tiers.py so app.py, office_ops.py, ask_pps.py and
# pipeline_board.py share one source instead of four copies that drift.
from tiers import (
    TIER_OWNER, TIER_LEADERSHIP, TIER_TEAM, DEFAULT_TIER,
    tier_label,
)
import tiers as _tiers


# ── USER DEFINITIONS ────────────────────────────────────────────────────────────

USERS = {
    'thomas_ellison': {
        'display': 'Thomas Ellison',
        'role': 'admin',
        'tier': TIER_OWNER,
        'title': 'President',
        'email': 'thomas@purepropsolutions.com',
    },
    'tony_cumella': {
        'display': 'Tony Cumella',
        'role': 'consultant',
        'tier': TIER_LEADERSHIP,
        'title': 'VP of Sales',
        'email': 'Tony@purepropsolutions.com',
    },
    'adam_cupito': {
        'display': 'Adam Cupito',
        'role': 'consultant',
        'tier': TIER_TEAM,
        'title': 'Property Solutions Consultant',
        'email': 'Adam@purepropsolutions.com',
    },
    'rachel_farler': {
        'display': 'Rachel Farler',
        'role': 'consultant',
        'tier': TIER_TEAM,
        'title': 'Property Solutions Consultant',
        'email': 'Rachel@purepropsolutions.com',
    },
    'andy_potts': {
        'display': 'Andy Potts',
        'role': 'consultant',
        'tier': TIER_TEAM,
        'title': 'Property Solutions Consultant',
        'email': 'Andy@purepropsolutions.com',
    },
    'phil_miller': {
        'display': 'Phil Miller',
        'role': 'pm',
        'tier': TIER_TEAM,
        # Title only (2026-08-23, Thomas) — see the note on jordan_allen.
        'title': 'Project Engineer',
        'email': 'phil@purepropsolutions.com',
    },
    'derek_kidney': {
        'display': 'Derek Kidney',
        'role': 'pm',
        'tier': TIER_TEAM,
        'title': 'Project Manager',
        'email': 'Derek@purepropsolutions.com',
    },
    'nick_triplett': {
        'display': 'Nick Triplett',
        'role': 'pm',
        'tier': TIER_TEAM,
        'title': 'Project Manager',
        'email': 'nick@purepropsolutions.com',
    },
    'trey_hollmeyer': {
        'display': 'Trey Hollmeyer',
        'role': 'pm',
        'tier': TIER_LEADERSHIP,
        'title': 'Production Manager',
        'email': 'trey@purepropsolutions.com',
    },
    'james_boling': {
        'display': 'James Boling',
        'role': 'pm',
        'tier': TIER_TEAM,
        'title': 'Project Manager',
        'email': 'James@purepropsolutions.com',
    },
    'jordan_allen': {
        'display': 'Jordan Allen',
        'role': 'pm',
        'tier': TIER_TEAM,
        # Title only (2026-08-23, Thomas). `title` is descriptive — it labels
        # people on Team View, the Admin roster and the recap. `role` and `tier`
        # are unchanged, so nothing about what he can see or do moves.
        'title': 'Senior Project Manager',
        'email': 'jordan@purepropsolutions.com',
    },
    'ben_ramsey': {
        'display': 'Ben Ramsey',
        'role': 'pm',
        'tier': TIER_TEAM,
        'title': 'Project Manager',
        'email': 'ben@purepropsolutions.com',
    },
    'stephanie_whetstone': {
        'display': 'Stephanie Whetstone',
        'role': 'office_manager',
        'tier': TIER_LEADERSHIP,
        'title': 'Office Manager',
        'email': 'Stephanie@purepropsolutions.com',
    },
}

# REMOVED 2026-08-21 — the shared "Admin" picker login (user_key 'admin',
# hardcoded password, added 2026-08-18 as the RJ login in 2f6dd58/d570f81/e07043b).
# Removed at Thomas's request: it was issued to a trusted person outside PPS, the
# credential is now known, and it lived in plaintext in this repo's history.
# Two reasons not to reintroduce a shared login of any kind:
#   1. Every action through it recorded as user_key 'admin', so per-person usage
#      for whoever used it was unrecoverable — see docs/HUB_REVIEW_2026-08-21.md F-01.
#   2. No email on the row, so Forgot Password could never work; the only recovery
#      path was editing this file.
# Give a real named USERS entry instead, and put them on the specific rosters
# (BOARD_ACCESS_ALL, OFFICE_OPS_USER_KEYS, CURATORS, CONTACTS_USER_KEYS) they need.
# Historical rows created by 'admin' still render — display names fall back to the
# key, so old proposals/board entries read "Admin" rather than going blank.

# Proposal numbers: {INITIALS}{YY}{XXX} — e.g. TE26001 (Thomas Ellison, 2026, #1)
PROPOSAL_NUMBER_SEQ_DIGITS = 3
PROPOSAL_NUMBER_INITIALS = {
    'thomas_ellison': 'TE',
    'tony_cumella': 'TC',
    'adam_cupito': 'AC',
    'rachel_farler': 'RF',
    'andy_potts': 'AP',
}

_CONSULTANT_KEY_ALIASES = {
    'thomas': 'thomas_ellison',
    'tony': 'tony_cumella',
    'adam': 'adam_cupito',
    'rachel': 'rachel_farler',
    'andy': 'andy_potts',
}


def _normalize_consultant_key(consultant_key):
    key = (consultant_key or '').strip()
    return _CONSULTANT_KEY_ALIASES.get(key, key)


def _proposal_initials(consultant_key):
    key = _normalize_consultant_key(consultant_key)
    initials = PROPOSAL_NUMBER_INITIALS.get(key)
    if initials:
        return initials
    user = USERS.get(key, {})
    display = user.get('display', '')
    parts = display.split()
    if len(parts) >= 2:
        return (parts[0][0] + parts[-1][0]).upper()
    if display:
        return display[:2].upper()
    return 'PP'


def _proposal_seq_suffix_len_ok(suffix_len):
    """Current format uses 3 digits; legacy hub entries may use 4."""
    return suffix_len in (PROPOSAL_NUMBER_SEQ_DIGITS, PROPOSAL_NUMBER_SEQ_DIGITS + 1)


def _parse_proposal_seq(raw, prefix):
    if not raw.startswith(prefix):
        return None
    suffix = raw[len(prefix):]
    if not suffix.isdigit() or not _proposal_seq_suffix_len_ok(len(suffix)):
        return None
    return int(suffix)


def _format_proposal_number(prefix, seq):
    return f"{prefix}{seq:0{PROPOSAL_NUMBER_SEQ_DIGITS}d}"


def _max_proposal_seq_from_log(cur, consultant_key, prefix):
    """Highest sequence suffix already used in proposal_log for this consultant/year."""
    cur.execute(
        '''
        SELECT proposal_number FROM proposal_log
        WHERE consultant_key = %s AND proposal_number IS NOT NULL
          AND UPPER(proposal_number) LIKE %s
        ''',
        (consultant_key, prefix + '%'),
    )
    max_seq = 0
    for row in cur.fetchall():
        num = re.sub(r'[^A-Z0-9]', '', (row[0] or '').strip().upper())
        seq = _parse_proposal_seq(num, prefix)
        if seq is not None:
            max_seq = max(max_seq, seq)
    return max_seq


def peek_next_proposal_number(consultant_key, year=None):
    """
    Suggest the next proposal number (INITIALS + 2-digit year + 3-digit seq) without
    reserving it. The sequence counter advances only on save via
    sync_proposal_number_sequence(), so edits before generate are respected.
    """
    key = _normalize_consultant_key(consultant_key)
    if key not in PROPOSAL_NUMBER_INITIALS and key not in USERS:
        return None, 'Unknown consultant'
    yr = year or datetime.now().year
    yy = yr % 100
    prefix = f"{_proposal_initials(key)}{yy:02d}"
    try:
        conn = get_db()
        if not conn:
            return None, 'Database unavailable'
        cur = conn.cursor()
        cur.execute(
            '''
            INSERT INTO proposal_number_sequence (consultant_key, seq_year, last_seq)
            VALUES (%s, %s, 0)
            ON CONFLICT (consultant_key, seq_year) DO NOTHING
            ''',
            (key, yr),
        )
        max_log = _max_proposal_seq_from_log(cur, key, prefix)
        if max_log:
            cur.execute(
                '''
                UPDATE proposal_number_sequence
                SET last_seq = GREATEST(last_seq, %s)
                WHERE consultant_key = %s AND seq_year = %s AND last_seq < %s
                ''',
                (max_log, key, yr, max_log),
            )
        cur.execute(
            '''
            SELECT last_seq FROM proposal_number_sequence
            WHERE consultant_key = %s AND seq_year = %s
            ''',
            (key, yr),
        )
        row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        if not row:
            return None, 'Could not read proposal number sequence'
        number = _format_proposal_number(prefix, row[0] + 1)
        return number, None
    except Exception as e:
        print(f"Proposal number peek error: {e}")
        return None, 'Could not read next proposal number'


def sync_proposal_number_sequence(consultant_key, proposal_number, year=None):
    """Advance the per-consultant sequence to the saved proposal number (never backward)."""
    key = _normalize_consultant_key(consultant_key)
    raw = re.sub(r'[^A-Z0-9]', '', (proposal_number or '').upper())
    if not raw:
        return
    yr = year or datetime.now().year
    yy = yr % 100
    expected_initials = _proposal_initials(key)
    prefix = f"{expected_initials}{yy:02d}"
    seq = _parse_proposal_seq(raw, prefix)
    if seq is None:
        return
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        cur.execute(
            '''
            INSERT INTO proposal_number_sequence (consultant_key, seq_year, last_seq)
            VALUES (%s, %s, %s)
            ON CONFLICT (consultant_key, seq_year) DO UPDATE SET
                last_seq = GREATEST(proposal_number_sequence.last_seq, EXCLUDED.last_seq)
            ''',
            (key, yr, seq),
        )
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Proposal number sync error: {e}")


# ── BIRTHDAYS & HIRE DATES ──────────────────────────────────────────────────
TEAM_DATES = {
    'thomas_ellison': {'birthday': (10, 8),  'hire': (5, 21, 2017)},
    'tony_cumella':   {'birthday': (8, 21),  'hire': (6, 17, 2019)},
    'phil_miller':    {'birthday': (3, 2),   'hire': (2, 24, 2020)},
    'trey_hollmeyer': {'birthday': (10, 3),  'hire': (5, 14, 2018)},
    'stephanie_whetstone': {'birthday': (5, 5), 'hire': (5, 7, 2017)},
    'derek_kidney':   {'birthday': (4, 15),  'hire': (4, 26, 2021)},
    'jordan_allen':   {'birthday': (6, 26),  'hire': (9, 13, 2021)},
    'adam_cupito':    {'birthday': (6, 18),  'hire': (2, 28, 2022)},
    'ben_ramsey':     {'birthday': (8, 6),   'hire': (7, 10, 2023)},
    'andy_potts':     {'birthday': (12, 31), 'hire': (2, 24, 2025)},
    'james_boling':   {'birthday': (3, 31),  'hire': (5, 5,  2025)},
    'rachel_farler':  {'birthday': (5, 15),  'hire': (1, 5,  2026)},
    'nick_triplett':  {'birthday': (12, 23), 'hire': (5, 4,  2026)},
}

MONTH_NAMES = ['January','February','March','April','May','June',
               'July','August','September','October','November','December']

def get_date_events(user_key, is_admin=False):
    """Returns upcoming birthday/anniversary events within 3 days for this user,
    plus team events if admin."""
    from datetime import date, timedelta
    today = date.today()
    events = []

    def days_until(month, day):
        this_year = date(today.year, month, day)
        if this_year < today:
            this_year = date(today.year + 1, month, day)
        return (this_year - today).days

    def event_label(days):
        if days == 0: return 'today'
        if days == 1: return 'tomorrow'
        return f'in {days} days'

    if is_admin:
        # Admin sees everyone's upcoming events
        for key, dates in TEAM_DATES.items():
            user = USERS.get(key, {})
            first_name = user.get('display', '').split()[0]
            bday_days = days_until(*dates['birthday'])
            hire_month, hire_day, hire_year = dates['hire']
            hire_days = days_until(hire_month, hire_day)
            years = today.year - hire_year
            if bday_days <= 3:
                events.append({
                    'type': 'birthday',
                    'name': first_name,
                    'full_name': user.get('display', ''),
                    'days': bday_days,
                    'label': event_label(bday_days),
                    'date_str': f"{MONTH_NAMES[dates['birthday'][0]-1]} {dates['birthday'][1]}",
                    'message': f"🎂 Heads up — {first_name}'s birthday is {event_label(bday_days)}, {MONTH_NAMES[dates['birthday'][0]-1]} {dates['birthday'][1]}. A good excuse to say something.",
                })
            if hire_days <= 3 and years >= 1:
                events.append({
                    'type': 'anniversary',
                    'name': first_name,
                    'full_name': user.get('display', ''),
                    'days': hire_days,
                    'label': event_label(hire_days),
                    'date_str': f"{MONTH_NAMES[hire_month-1]} {hire_day}",
                    'message': f"🏆 {first_name} is hitting {years} year{'s' if years > 1 else ''} with PPS {event_label(hire_days)}. Worth acknowledging.",
                })
    else:
        # User sees their own events (must include 'days' for sort below)
        dates = TEAM_DATES.get(user_key)
        if dates:
            user = USERS.get(user_key, {})
            first_name = user.get('display', '').split()[0]
            bday_days = days_until(*dates['birthday'])
            hire_month, hire_day, hire_year = dates['hire']
            hire_days = days_until(hire_month, hire_day)
            years = today.year - hire_year
            if bday_days <= 3:
                events.append({
                    'type': 'birthday',
                    'name': first_name,
                    'full_name': user.get('display', ''),
                    'days': bday_days,
                    'label': event_label(bday_days),
                    'date_str': f"{MONTH_NAMES[dates['birthday'][0]-1]} {dates['birthday'][1]}",
                    'message': f"🎂 Your birthday is {event_label(bday_days)} — {MONTH_NAMES[dates['birthday'][0]-1]} {dates['birthday'][1]}. Hope it's a good one.",
                })
            if hire_days <= 3 and years >= 1:
                events.append({
                    'type': 'anniversary',
                    'name': first_name,
                    'full_name': user.get('display', ''),
                    'days': hire_days,
                    'label': event_label(hire_days),
                    'date_str': f"{MONTH_NAMES[hire_month-1]} {hire_day}",
                    'message': f"🏆 {years} year{'s' if years > 1 else ''} at PPS {event_label(hire_days)}. That's worth something.",
                })

    events.sort(key=lambda x: x.get('days', 999))
    return events

CONSULTANTS = {
    'tony_cumella': 'Tony Cumella',
    'adam_cupito': 'Adam Cupito',
    'rachel_farler': 'Rachel Farler',
    'andy_potts': 'Andy Potts',
}

# ── DATABASE ────────────────────────────────────────────────────────────────────

def _database_url():
    """Normalize Render/Heroku-style postgres:// URLs for psycopg2."""
    url = (DATABASE_URL or '').strip()
    if not url:
        return ''
    if url.startswith('postgres://'):
        url = 'postgresql://' + url[len('postgres://'):]
    # Strip query sslmode so we can control it explicitly (avoids double-setting).
    if '?' in url:
        base, qs = url.split('?', 1)
        parts = [p for p in qs.split('&') if p and not p.lower().startswith('sslmode=')]
        url = base + (('?' + '&'.join(parts)) if parts else '')
    return url


# Last connect failure (sanitized) — for /health diagnostics, not shown to end users.
_DB_LAST_ERROR = ''


def _track_checkout(conn):
    """Remember a connection so teardown can return it if nobody else does.

    Seven of the twenty-one connections a dashboard load used to take were
    never closed by their caller — they were released whenever Python got
    round to collecting them, which was harmless when each one was
    disposable and is not harmless now that they come from a fixed pool.
    Rather than audit 107 call sites, the request that took them gives them
    back at the end.

    Only inside a request. A background thread (the proposal-diff emailer)
    has no request context and must close its own; tracking it here would
    hand its connection to whichever request happened to finish first.
    """
    try:
        if not has_request_context():
            return
        bucket = getattr(g, '_db_checkouts', None)
        if bucket is None:
            bucket = []
            g._db_checkouts = bucket
        bucket.append(conn)
    except Exception:
        pass


def _untrack_checkout(conn):
    try:
        if not has_request_context():
            return
        bucket = getattr(g, '_db_checkouts', None)
        if bucket:
            try:
                bucket.remove(conn)
            except ValueError:
                pass
    except Exception:
        pass


def get_db():
    """A Postgres connection. Returns None if unavailable.

    Pooled since 2026-08-25 — see `db_pool.py` for why, and for the four
    invariants that make it invisible to callers. The contract here has not
    changed: you get something you can `.cursor()` on, and you call
    `.close()` when you are done. `.close()` now means "give it back".

    Three ways this can hand back an unpooled connection, all of them
    deliberate and all of them exactly the old behaviour:
      * `DB_POOL_DISABLED=true` — the kill switch;
      * the pool is at its cap (`acquire()` returns None) — a busy moment
        must not become a queue behind gunicorn's 120s timeout;
      * building the pool failed at all.
    """
    global _DB_LAST_ERROR
    url = _database_url()
    if not url:
        _DB_LAST_ERROR = 'DATABASE_URL not set'
        return None

    if db_pool.pooling_enabled():
        try:
            pool = db_pool.get_pool(url, lambda: db_pool.connect_direct(url))
            raw = pool.acquire()
            if raw is not None:
                _DB_LAST_ERROR = ''
                conn = db_pool.PooledConnection(
                    raw, pool, on_release=_untrack_checkout)
                _track_checkout(conn)
                return conn
            # Pool at its cap. Fall through to a direct connection.
        except Exception as e:
            # Never let a pool problem be the reason a page has no database.
            _DB_LAST_ERROR = str(e)[:240]
            print(f'get_db pool error: {_DB_LAST_ERROR}')

    try:
        raw = db_pool.connect_direct(url)
    except Exception as e:
        # Never log credentials; psycopg2 errors usually omit the password.
        _DB_LAST_ERROR = str(e)[:240]
        print(f'get_db connect error: {_DB_LAST_ERROR}')
        return None
    _DB_LAST_ERROR = ''
    conn = db_pool.PooledConnection(raw, None, on_release=_untrack_checkout)
    _track_checkout(conn)
    return conn


@app.teardown_request
def _return_db_connections(exc=None):
    """Give back anything this request took and did not return."""
    bucket = getattr(g, '_db_checkouts', None)
    if not bucket:
        return
    # close() removes from the bucket via _untrack_checkout, so iterate a copy.
    for conn in list(bucket):
        try:
            # An exploded request may have left a transaction open. release()
            # rolls it back, but discarding is the honest call when we have no
            # idea what state the caller left it in.
            if exc is not None:
                conn.discard()
            else:
                conn.close()
        except Exception:
            pass
    g._db_checkouts = []


def init_db():
    conn = get_db()
    if not conn:
        return
    # Every ALTER / CREATE / DROP below runs in its own savepoint, so one that
    # cannot apply is skipped and logged instead of aborting the transaction
    # and silently killing every statement after it. That was a real bug: four
    # ALTERs against the estimate log tables run before those tables are
    # created further down this function, so on a database that has never run
    # the estimators, init_db used to half-complete and report one opaque
    # "DB init error". See db_ddl.py. Non-schema statements — the seeding
    # INSERTs, the last_login backfill — are untouched and still raise.
    cur = db_ddl.checkpointed(conn.cursor())
    try:
        _init_db_body(conn, cur)
    finally:
        # Whatever happened above, commit what worked and give the connection
        # back. See _finish_init_db.
        _finish_init_db(conn, cur)


def _init_db_body(conn, cur):
    # Users table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS hub_users (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) UNIQUE NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            password_hash VARCHAR(500) NOT NULL,
            role VARCHAR(50) NOT NULL,
            created_at TIMESTAMP DEFAULT NOW(),
            last_login TIMESTAMP
        )
    ''')

    cur.execute('''
        CREATE TABLE IF NOT EXISTS proposal_number_sequence (
            consultant_key VARCHAR(100) NOT NULL,
            seq_year INTEGER NOT NULL,
            last_seq INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (consultant_key, seq_year)
        )
    ''')

    # Proposal activity log
    cur.execute('''
        CREATE TABLE IF NOT EXISTS proposal_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            consultant_key VARCHAR(100) NOT NULL,
            consultant_name VARCHAR(255) NOT NULL,
            client_name VARCHAR(255),
            property_name VARCHAR(255),
            property_address VARCHAR(255),
            property_type VARCHAR(100),
            template_type VARCHAR(100),
            proposal_number VARCHAR(100),
            existing_issue TEXT,
            intended_outcome TEXT,
            scopes_selected TEXT,
            scope_notes TEXT,
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    # Add new columns if they don't exist (for existing databases)
    for col in [
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS property_address VARCHAR(255)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS proposal_number VARCHAR(100)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS existing_issue TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS intended_outcome TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS scopes_selected TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS scope_notes TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS document_id INTEGER",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS contact_name VARCHAR(255)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS contact_email VARCHAR(255)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS company VARCHAR(255)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS scope_details TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS other_scope TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS pricing_json TEXT",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS warranty_pps VARCHAR(100)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS warranty_mfg VARCHAR(100)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS proposal_date VARCHAR(50)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS expiry_date VARCHAR(50)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS contract_total VARCHAR(100)",
        "ALTER TABLE proposal_log ADD COLUMN IF NOT EXISTS scope_style VARCHAR(50)",
        "ALTER TABLE siding_estimate_log ADD COLUMN IF NOT EXISTS summary_meta VARCHAR(255)",
        "ALTER TABLE roofing_estimate_log ADD COLUMN IF NOT EXISTS summary_meta VARCHAR(255)",
        "ALTER TABLE gutter_estimate_log ADD COLUMN IF NOT EXISTS summary_meta VARCHAR(255)",
        "ALTER TABLE painting_estimate_log ADD COLUMN IF NOT EXISTS summary_meta VARCHAR(255)",
    ]:
        cur.execute(col)

    # Document vault — persistent file storage
    cur.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id SERIAL PRIMARY KEY,
            doc_type VARCHAR(50) NOT NULL,
            log_id INTEGER,
            user_key VARCHAR(100) NOT NULL,
            filename VARCHAR(255) NOT NULL,
            mime_type VARCHAR(100) NOT NULL,
            size_bytes INTEGER NOT NULL,
            file_data BYTEA NOT NULL,
            created_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    cur.execute("CREATE INDEX IF NOT EXISTS idx_documents_log ON documents(doc_type, log_id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_documents_user ON documents(user_key)")

    # PPM activity log
    cur.execute('''
        CREATE TABLE IF NOT EXISTS ppm_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            property_name VARCHAR(255),
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')

    # Subscope log table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS subscope_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            property_name VARCHAR(255),
            pm_name VARCHAR(255),
            consultant_name VARCHAR(255),
            language VARCHAR(50),
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')

    # SSO auth tokens table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS auth_tokens (
            token VARCHAR(64) PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            role VARCHAR(50),
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP NOT NULL,
            used BOOLEAN DEFAULT FALSE
        )
    ''')


    # Feedback table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            message TEXT NOT NULL,
            feedback_type VARCHAR(50) DEFAULT 'general',
            submitted_at TIMESTAMP DEFAULT NOW(),
            read_by_admin BOOLEAN DEFAULT FALSE
        )
    ''')
    for col in (
        "ALTER TABLE feedback ADD COLUMN IF NOT EXISTS archived BOOLEAN DEFAULT FALSE",
        "ALTER TABLE feedback ADD COLUMN IF NOT EXISTS archived_at TIMESTAMP",
    ):
        try:
            cur.execute(col)
        except Exception:
            pass

    # Proposal diffs table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS proposal_diffs (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            property_name VARCHAR(255),
            diff_analysis TEXT,
            user_notes TEXT,
            voice_recommendations TEXT,
            submitted_at TIMESTAMP DEFAULT NOW(),
            reviewed_by_admin BOOLEAN DEFAULT FALSE
        )
    ''')

    cur.execute(
        'ALTER TABLE proposal_diffs ADD COLUMN IF NOT EXISTS comparison_prompt TEXT'
    )

    # Migrate ppm_log / subscope_log metadata columns
    for col in [
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS pm_key VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS pm_name VARCHAR(255)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS property_address VARCHAR(255)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS client_name VARCHAR(255)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS proposal_number VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS total_value VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS proposal_date VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS proj_type VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS scale VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS client_type VARCHAR(100)",
        "ALTER TABLE ppm_log ADD COLUMN IF NOT EXISTS occupied VARCHAR(100)",
        "ALTER TABLE subscope_log ADD COLUMN IF NOT EXISTS property_address VARCHAR(255)",
        "ALTER TABLE subscope_log ADD COLUMN IF NOT EXISTS po_number VARCHAR(100)",
        "ALTER TABLE subscope_log ADD COLUMN IF NOT EXISTS consultant_key VARCHAR(100)",
        "ALTER TABLE subscope_log ADD COLUMN IF NOT EXISTS pm_key VARCHAR(100)",
        "ALTER TABLE subscope_log ADD COLUMN IF NOT EXISTS material_provider VARCHAR(50)",
        "ALTER TABLE subscope_log ADD COLUMN IF NOT EXISTS proposal_filename VARCHAR(255)",
        "ALTER TABLE feedback ADD COLUMN IF NOT EXISTS feedback_type VARCHAR(50) DEFAULT 'general'",
    ]:
        cur.execute(col)

    # Client / contact database
    cur.execute('''
        CREATE TABLE IF NOT EXISTS clients (
            id SERIAL PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            email VARCHAR(255),
            company VARCHAR(255),
            property_name VARCHAR(255),
            address TEXT,
            notes TEXT,
            added_by VARCHAR(100),
            updated_at TIMESTAMP DEFAULT NOW(),
            created_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    cur.execute("CREATE INDEX IF NOT EXISTS idx_clients_name ON clients(LOWER(name))")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_clients_company ON clients(LOWER(company))")

    # Site visit log
    cur.execute('''
        CREATE TABLE IF NOT EXISTS site_visit_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            property_name VARCHAR(255),
            property_address VARCHAR(255),
            visit_date VARCHAR(100),
            visit_time VARCHAR(100),
            po_number VARCHAR(100),
            trade_partner_present VARCHAR(10),
            trade_partner_company VARCHAR(255),
            crew_lead VARCHAR(255),
            crew_count VARCHAR(50),
            met_with_staff VARCHAR(10),
            staff_contact VARCHAR(255),
            topics_discussed TEXT,
            complaints_received VARCHAR(10),
            complaint_details TEXT,
            checklist JSONB,
            overall_status VARCHAR(50),
            quality_status VARCHAR(50),
            schedule_status VARCHAR(50),
            observations TEXT,
            photos_taken BOOLEAN DEFAULT FALSE,
            next_visit_date VARCHAR(100),
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')

    # Siding estimate log
    cur.execute('''
        CREATE TABLE IF NOT EXISTS siding_estimate_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            display_name VARCHAR(255),
            property_name VARCHAR(255),
            property_address VARCHAR(255),
            building_count INTEGER DEFAULT 1,
            siding_type VARCHAR(100),
            job_data JSONB NOT NULL,
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    try:
        cur.execute(
            "CREATE INDEX IF NOT EXISTS idx_siding_estimate_user ON siding_estimate_log(generated_by, generated_at DESC)"
        )
    except Exception:
        pass

    cur.execute('''
        CREATE TABLE IF NOT EXISTS roofing_estimate_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            display_name VARCHAR(255),
            property_name VARCHAR(255),
            property_address VARCHAR(255),
            report_type VARCHAR(50),
            job_data JSONB NOT NULL,
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    try:
        cur.execute(
            "CREATE INDEX IF NOT EXISTS idx_roofing_estimate_user ON roofing_estimate_log(generated_by, generated_at DESC)"
        )
    except Exception:
        pass

    cur.execute('''
        CREATE TABLE IF NOT EXISTS gutter_estimate_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            display_name VARCHAR(255),
            property_name VARCHAR(255),
            property_address VARCHAR(255),
            gutter_lf NUMERIC,
            job_data JSONB NOT NULL,
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    try:
        cur.execute(
            "CREATE INDEX IF NOT EXISTS idx_gutter_estimate_user ON gutter_estimate_log(generated_by, generated_at DESC)"
        )
    except Exception:
        pass

    cur.execute('''
        CREATE TABLE IF NOT EXISTS painting_estimate_log (
            id SERIAL PRIMARY KEY,
            generated_by VARCHAR(100) NOT NULL,
            display_name VARCHAR(255),
            property_name VARCHAR(255),
            property_address VARCHAR(255),
            line_count INTEGER,
            one_coat_bid NUMERIC,
            two_coat_bid NUMERIC,
            job_data JSONB NOT NULL,
            generated_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    try:
        cur.execute(
            "CREATE INDEX IF NOT EXISTS idx_painting_estimate_user ON painting_estimate_log(generated_by, generated_at DESC)"
        )
    except Exception:
        pass

    cur.execute('''
        CREATE TABLE IF NOT EXISTS hub_settings (
            key VARCHAR(100) PRIMARY KEY,
            value JSONB NOT NULL,
            updated_at TIMESTAMP DEFAULT NOW(),
            updated_by VARCHAR(100)
        )
    ''')

    # Auth support tables
    cur.execute('''
        CREATE TABLE IF NOT EXISTS auth_codes (
            code VARCHAR(64) PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            role VARCHAR(50),
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP NOT NULL,
            used BOOLEAN DEFAULT FALSE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS login_attempts (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            success BOOLEAN NOT NULL,
            ip_address VARCHAR(45),
            attempted_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS password_reset_tokens (
            token VARCHAR(64) PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            created_at TIMESTAMP DEFAULT NOW(),
            expires_at TIMESTAMP NOT NULL,
            used BOOLEAN DEFAULT FALSE
        )
    ''')
    try:
        # Session-invalidation stamp — see _session_password_stale(). Added here
        # as well as in _ensure_hub_users_password_schema so the column exists
        # from startup, not only after someone's first password write.
        cur.execute(
            'ALTER TABLE hub_users ADD COLUMN IF NOT EXISTS password_epoch INTEGER DEFAULT 0'
        )
    except Exception as e:
        print(f'hub_users password_epoch migrate: {e}')
    try:
        cur.execute('ALTER TABLE hub_users ADD COLUMN IF NOT EXISTS must_change_password BOOLEAN DEFAULT FALSE')
    except Exception:
        pass
    try:
        cur.execute('ALTER TABLE hub_users ALTER COLUMN password_hash TYPE TEXT')
    except Exception:
        pass

    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_training_progress (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            item_id VARCHAR(100) NOT NULL,
            completed BOOLEAN DEFAULT FALSE,
            completed_at TIMESTAMP,
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(user_key, item_id)
        )
    ''')
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_psc_training_user ON psc_training_progress(user_key)")
    except Exception:
        pass

    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_training_notes (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            week_num INTEGER NOT NULL,
            notes TEXT,
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(user_key, week_num)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_training_feedback (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            week_num INTEGER,
            message TEXT NOT NULL,
            feedback_type VARCHAR(50) DEFAULT 'improvement',
            submitted_at TIMESTAMP DEFAULT NOW(),
            read_by_admin BOOLEAN DEFAULT FALSE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_training_enrollment (
            user_key VARCHAR(100) PRIMARY KEY,
            enrolled_at TIMESTAMP DEFAULT NOW(),
            enrolled_by VARCHAR(100),
            manager_key VARCHAR(100) NOT NULL DEFAULT 'tony_cumella',
            target_weeks INTEGER DEFAULT 12,
            last_activity_at TIMESTAMP,
            graduated_at TIMESTAMP,
            active BOOLEAN DEFAULT TRUE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_training_notifications (
            user_key VARCHAR(100) NOT NULL,
            notification_type VARCHAR(50) NOT NULL,
            week_num INTEGER NOT NULL DEFAULT -1,
            created_at TIMESTAMP DEFAULT NOW(),
            PRIMARY KEY (user_key, notification_type, week_num)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_training_manager_signoffs (
            user_key VARCHAR(100) NOT NULL,
            week_num INTEGER NOT NULL,
            signed_by VARCHAR(100) NOT NULL,
            signed_at TIMESTAMP DEFAULT NOW(),
            PRIMARY KEY (user_key, week_num)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_roleplay_sessions (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            scenario_id VARCHAR(100) NOT NULL,
            transcript TEXT NOT NULL,
            feedback TEXT NOT NULL,
            overall REAL NOT NULL,
            result VARCHAR(32) NOT NULL,
            turn_count INTEGER NOT NULL,
            created_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS psc_roleplay_daily_usage (
            user_key VARCHAR(100) NOT NULL,
            usage_date DATE NOT NULL DEFAULT CURRENT_DATE,
            turn_count INTEGER NOT NULL DEFAULT 0,
            grade_count INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (user_key, usage_date)
        )
    ''')
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_roleplay_sessions_user ON psc_roleplay_sessions(user_key)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_roleplay_sessions_created ON psc_roleplay_sessions(created_at DESC)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_psc_training_notes_user ON psc_training_notes(user_key)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_psc_training_feedback_time ON psc_training_feedback(submitted_at DESC)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_psc_training_enrollment_active ON psc_training_enrollment(active)")
    except Exception:
        pass

    # PM training (mirror of PSC — separate item-id namespace)
    cur.execute('''
        CREATE TABLE IF NOT EXISTS pm_training_progress (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            item_id VARCHAR(100) NOT NULL,
            completed BOOLEAN DEFAULT FALSE,
            completed_at TIMESTAMP,
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(user_key, item_id)
        )
    ''')
    try:
        cur.execute("CREATE INDEX IF NOT EXISTS idx_pm_training_user ON pm_training_progress(user_key)")
    except Exception:
        pass
    cur.execute('''
        CREATE TABLE IF NOT EXISTS pm_training_notes (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            week_num INTEGER NOT NULL,
            notes TEXT,
            updated_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(user_key, week_num)
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS pm_training_feedback (
            id SERIAL PRIMARY KEY,
            user_key VARCHAR(100) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            week_num INTEGER,
            message TEXT NOT NULL,
            feedback_type VARCHAR(50) DEFAULT 'improvement',
            submitted_at TIMESTAMP DEFAULT NOW(),
            read_by_admin BOOLEAN DEFAULT FALSE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS pm_training_enrollment (
            user_key VARCHAR(100) PRIMARY KEY,
            enrolled_at TIMESTAMP DEFAULT NOW(),
            enrolled_by VARCHAR(100),
            manager_key VARCHAR(100) NOT NULL DEFAULT 'trey_hollmeyer',
            target_weeks INTEGER DEFAULT 4,
            last_activity_at TIMESTAMP,
            graduated_at TIMESTAMP,
            active BOOLEAN DEFAULT TRUE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS pm_training_manager_signoffs (
            user_key VARCHAR(100) NOT NULL,
            week_num INTEGER NOT NULL,
            signed_by VARCHAR(100) NOT NULL,
            signed_at TIMESTAMP DEFAULT NOW(),
            PRIMARY KEY (user_key, week_num)
        )
    ''')

    ask_pps.init_tables(cur)
    pipeline_board.init_tables(cur)
    pipeline_board.cleanup_legacy_import_notes(cur)
    office_ops.init_tables(cur)
    insurance_compliance.init_tables(cur)
    import hub_usage
    hub_usage.init_tables(cur)
    training_overlay.init_tables(cur)

    # Seed missing accounts with an UNUSABLE password, never a shared one.
    #
    # This used to plant os.environ['DEFAULT_PASSWORD'] on every account that
    # had no row yet, which is where the shared password came from in the first
    # place. Retiring the current one (password_campaign.py) would have been
    # pointless while this stayed: the next person added to USERS would be
    # seeded with the same secret everyone already knew, including anyone who
    # had since left. DEFAULT_PASSWORD is now ignored even if it is still set
    # on the Render service — deleting the env var is good hygiene, but the
    # code no longer depends on anyone remembering to.
    #
    # A new hire gets a row they cannot log into, and reaches the Hub the same
    # way everyone else does after a reset: Forgot Password, or Thomas sending
    # them a link from /admin. One less shared secret, permanently.
    for key, user in USERS.items():
        cur.execute('SELECT id FROM hub_users WHERE user_key = %s', (key,))
        if not cur.fetchone():
            unusable = generate_password_hash(secrets.token_urlsafe(48), method='pbkdf2:sha256')
            cur.execute(
                '''INSERT INTO hub_users
                   (user_key, display_name, password_hash, role, must_change_password)
                   VALUES (%s, %s, %s, %s, TRUE)''',
                (key, user['display'], unusable, user['role'])
            )
            print(f'seeded {key} with no usable password — send them a reset link')

    # Revoke the retired shared "Admin" login (removed 2026-08-21, see USERS above).
    # Dropping it from USERS already blocks sign-in — the login route rejects any
    # user_key not in USERS — but the hub_users row would otherwise keep a working
    # password hash for a known credential. Idempotent; a no-op once the row is gone.
    db_ddl.optional_step(
        cur, "DELETE FROM hub_users WHERE user_key = 'admin'",
        label='retired admin login cleanup')

    # Backfill last_login from tool activity where logs are more recent.
    # optional_step, not try/except: this reads four log tables, and on a
    # database where any of them does not exist yet the failure would abort
    # the transaction and make the commit below silently roll back every
    # table init_db had just created.
    db_ddl.optional_step(cur, '''
            UPDATE hub_users u
            SET last_login = activity.last_seen
            FROM (
                SELECT user_key, MAX(last_seen) AS last_seen
                FROM (
                    SELECT generated_by AS user_key, MAX(generated_at) AS last_seen
                    FROM proposal_log GROUP BY generated_by
                    UNION ALL
                    SELECT generated_by, MAX(generated_at) FROM ppm_log GROUP BY generated_by
                    UNION ALL
                    SELECT generated_by, MAX(generated_at) FROM subscope_log GROUP BY generated_by
                    UNION ALL
                    SELECT generated_by, MAX(generated_at) FROM site_visit_log GROUP BY generated_by
                ) combined
                GROUP BY user_key
            ) activity
            WHERE u.user_key = activity.user_key
              AND (u.last_login IS NULL OR u.last_login < activity.last_seen)
        ''', label='last_login backfill')


def _finish_init_db(conn, cur):
    """Commit and hand the connection back, whatever happened above.

    This used to be three bare lines at the end of init_db, which meant that
    if anything in the several hundred statements before them raised, the
    connection was never closed. That was survivable when every connection
    was disposable; with a pool it permanently costs one of ten slots per
    worker, and `init_db` runs outside a request so the teardown hook cannot
    clean up after it. Measured: `in_use` sat at 1 forever from process
    start.
    """
    try:
        conn.commit()
    except Exception as e:
        print(f'init_db commit failed: {e}')
        try:
            conn.rollback()
        except Exception:
            pass
    try:
        cur.close()
    except Exception:
        pass
    try:
        conn.close()
    except Exception:
        pass


_db_startup_lock = threading.Lock()
_db_startup_done = False


def _run_db_startup():
    """Run migrations once — gunicorn imports app in every worker."""
    global _db_startup_done
    with _db_startup_lock:
        if _db_startup_done:
            return
        try:
            init_db()
        except Exception as e:
            print(f"DB init error: {e}")
        # try/finally, not try/except: startup runs outside a request, so the
        # teardown hook cannot give this connection back if something here
        # raises before the close. One permanently checked-out connection per
        # worker is one of ten pool slots gone for the life of the process.
        _conn = None
        try:
            _conn = get_db()
            if _conn:
                _cur = _conn.cursor()
                _cur.execute('''
                    CREATE TABLE IF NOT EXISTS auth_tokens (
                        token VARCHAR(64) PRIMARY KEY,
                        user_key VARCHAR(100) NOT NULL,
                        display_name VARCHAR(255) NOT NULL,
                        role VARCHAR(50),
                        created_at TIMESTAMP DEFAULT NOW(),
                        expires_at TIMESTAMP NOT NULL,
                        used BOOLEAN DEFAULT FALSE
                    )
                ''')
                _conn.commit()
                _cur.close()
                print("auth_tokens table ready")
        except Exception as _e:
            print(f"auth_tokens migration error: {_e}")
        finally:
            if _conn is not None:
                try:
                    _conn.close()
                except Exception:
                    pass
        _db_startup_done = True


_run_db_startup()


# ── HELPERS ─────────────────────────────────────────────────────────────────────

def get_current_user():
    return session.get('user_key')


# Thin wrappers that bind tiers.py's functions to this app's USERS roster, so
# call sites read `is_owner(user_key)` rather than threading USERS through.
def user_tier(user_key):
    return _tiers.user_tier(USERS, user_key)


def has_tier(user_key, minimum):
    return _tiers.has_tier(USERS, user_key, minimum)


def is_owner(user_key):
    return _tiers.is_owner(USERS, user_key)


def is_leadership(user_key):
    return _tiers.is_leadership(USERS, user_key)


# How often a live session re-checks that its password is still current.
# Not per-request on purpose: there is no connection pool yet (see
# docs/HUB_REVIEW_2026-08-21.md F-05), so a SELECT on every request would mean
# a fresh Postgres connect on every request. The cost of the throttle is that a
# password reset can take up to this long to evict other devices — acceptable
# for "someone lost a phone", and worth revisiting once pooling lands, at which
# point this can safely drop to every request.
PASSWORD_EPOCH_RECHECK_SECONDS = 15 * 60


def _session_password_stale(user_key):
    """True when this session was opened with a password that has since changed.

    Returns False on any DB trouble — an unreachable database must not sign the
    whole company out. Fails open by design; the roster check above fails closed.
    """
    try:
        now = time.time()
        checked_at = float(session.get('pw_checked_at') or 0)
        if now - checked_at < PASSWORD_EPOCH_RECHECK_SECONDS:
            return False
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            'SELECT COALESCE(password_epoch, 0) FROM hub_users WHERE user_key = %s',
            (user_key,),
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return False
        session['pw_checked_at'] = now
        return int(row[0] or 0) != int(session.get('pw_epoch') or 0)
    except Exception as e:
        print(f'password epoch check error for {user_key}: {e}')
        return False


def require_login(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        current_key = session.get('user_key')
        # Offboarding actually revokes access. A signed session cookie is valid
        # until it expires, so before this check, deleting someone from USERS
        # left their open session working for the rest of its lifetime — the
        # cookie carries the identity and nothing re-checked the roster. That
        # matters more the longer sessions live (see PERMANENT_SESSION_LIFETIME).
        if current_key and current_key not in USERS:
            session.clear()
            current_key = None
        if current_key and _session_password_stale(current_key):
            session.clear()
            current_key = None
        if not current_key:
            # Prefer a clean relative next so mobile Safari doesn't loop on full
            # auth URLs (login?next=login?next=…).
            nxt = safe_next_url(request.full_path if request.query_string else request.path)
            if nxt:
                return redirect(url_for('login', next=nxt))
            return redirect(url_for('login'))
        if session.get('must_change_password') and request.endpoint not in ('change_password', 'logout'):
            return redirect(url_for('change_password'))
        return f(*args, **kwargs)
    return decorated


def require_admin(f):
    """Owner tier only (Thomas). Name kept — 10 routes and the 25 inline
    `session.get('role') != 'admin'` checks elsewhere all mean the same thing,
    since Thomas is the only admin-role user. Renaming those is churn with real
    regression risk and no behavior change; left as a follow-up."""
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not is_owner(session.get('user_key')):
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated


def can_manage_contacts(user_key):
    """Who may open /clients and create/update contact records.

    Everyone on the roster (2026-08-21, tier rework). Contacts are shared sales
    infrastructure; the old CONTACTS_USER_KEYS/CONTACTS_ROLES pair only ever
    produced 403s for people who needed to save a client mid-proposal — Trey
    reported exactly that, and the fix at the time was widening the role list,
    which is how these accumulate.
    """
    return bool(user_key) and user_key in USERS


def get_user_proposal_access(user_key):
    """Which consultants' books this person may work in — everyone, all of them.

    Was a per-user `proposal_access` field (2026-08-21 tier rework). It had five
    different shapes across thirteen people and had drifted far enough from its
    name that Pipeline Board built a second roster rather than read it. Kept as
    a function, not inlined, because a dozen call sites want the consultant list
    and this is the one place to narrow it again if that day comes.
    """
    if not user_key or user_key not in USERS:
        return []
    return list(CONSULTANTS.keys())


def user_can_access_consultant_proposals(user_key, consultant_key):
    """True if this user may view/prefill proposals for that consultant's book."""
    if not user_key or not consultant_key:
        return False
    return consultant_key in get_user_proposal_access(user_key)


def get_recent_proposals(user_key, limit=5):
    """Proposals this user generated, plus their consultants' book (for PMs)."""
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor(cursor_factory=RealDictCursor)
        access = get_user_proposal_access(user_key)
        if access:
            cur.execute('''
                SELECT * FROM proposal_log
                WHERE generated_by = %s OR consultant_key = ANY(%s)
                ORDER BY generated_at DESC LIMIT %s
            ''', (user_key, access, limit))
        else:
            cur.execute('''
                SELECT * FROM proposal_log
                WHERE generated_by = %s
                ORDER BY generated_at DESC LIMIT %s
            ''', (user_key, limit))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return rows
    except Exception as e:
        print(f'get_recent_proposals error: {e}')
        return []


def get_recent_ppms(user_key, limit=5):
    try:
        conn = get_db()
        if not conn:
            return []
        user_def = USERS.get(user_key, {})
        cur = conn.cursor(cursor_factory=RealDictCursor)
        if user_def.get('role') in ('pm', 'admin'):
            cur.execute('''
                SELECT * FROM ppm_log
                WHERE generated_by = %s OR pm_key = %s
                ORDER BY generated_at DESC LIMIT %s
            ''', (user_key, user_key, limit))
        else:
            cur.execute('''
                SELECT * FROM ppm_log
                WHERE generated_by = %s
                ORDER BY generated_at DESC LIMIT %s
            ''', (user_key, limit))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return rows
    except:
        return []


def get_recent_tpscopes(user_key, limit=5):
    """Trade Partner Scopes this user generated or is listed as PM on."""
    try:
        conn = get_db()
        if not conn:
            return []
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            '''SELECT * FROM subscope_log
               WHERE generated_by = %s OR pm_key = %s
               ORDER BY generated_at DESC LIMIT %s''',
            (user_key, user_key, limit),
        )
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return rows
    except Exception as e:
        print(f'get_recent_tpscopes error: {e}')
        return []


def get_profile_result(user_key):
    user = USERS.get(user_key, {})
    display_name = user.get('display', '')
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('''
            SELECT * FROM profile_results
            WHERE LOWER(name) = LOWER(%s)
            ORDER BY taken_date DESC LIMIT 1
        ''', (display_name,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return row
    except:
        return None


def get_psc_training_progress(user_key):
    """Return {item_id: True} for completed training items."""
    progress = {}
    try:
        conn = get_db()
        if not conn:
            return progress
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT item_id FROM psc_training_progress WHERE user_key = %s AND completed = TRUE',
            (user_key,),
        )
        for row in cur.fetchall():
            progress[row['item_id']] = True
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PSC training progress read error: {e}")
    return progress


def _psc_accountability_recipients():
    """President and PSC training manager — accountability email recipients."""
    recipients = []
    for key in ('thomas_ellison', PSC_TRAINING_MANAGER):
        email = (USERS.get(key, {}).get('email') or '').strip()
        if email and email.lower() not in {r.lower() for r in recipients}:
            recipients.append(email)
    return recipients


def _send_psc_accountability_email(subject, text_body, html_body=None):
    """Email President and training manager about PSC onboarding events."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    recipients = _psc_accountability_recipients()
    if not recipients:
        print(f"PSC accountability (no recipients): {subject}\n{text_body}")
        return False

    smtp_host = os.environ.get('SMTP_HOST', '')
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')
    if not smtp_host:
        print(f"PSC accountability email:\nSubject: {subject}\nTo: {', '.join(recipients)}\n{text_body}")
        return False

    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = smtp_user
        msg['To'] = ', '.join(recipients)
        msg.attach(MIMEText(text_body, 'plain'))
        if html_body:
            msg.attach(MIMEText(html_body, 'html'))
        with smtplib.SMTP_SSL(smtp_host, 465, timeout=30) as s:
            s.login(smtp_user, smtp_pass)
            s.send_message(msg)
        return True
    except Exception as e:
        print(f"PSC accountability email failed: {e}")
        return False


def _psc_week_labels():
    """Read-only walk — see psc_training_data.read_curriculum on why this does
    not go through get_training_curriculum()."""
    return psc_get_week_labels()


def _psc_week_trainee_complete(progress, week_map):
    """Return {week_num: bool} for whether the trainee finished every item in that week."""
    status = {}
    for week_num, ids in week_map.items():
        if not ids:
            status[week_num] = False
            continue
        done = sum(1 for item_id in ids if progress.get(item_id))
        status[week_num] = done == len(ids)
    return status


def get_psc_manager_signoffs(user_key):
    """Return {week_num: {signed_by, signed_at, signed_by_display}}."""
    signoffs = {}
    try:
        conn = get_db()
        if not conn:
            return signoffs
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT week_num, signed_by, signed_at FROM psc_training_manager_signoffs WHERE user_key = %s',
            (user_key,),
        )
        for row in cur.fetchall():
            signer = row['signed_by']
            signoffs[row['week_num']] = {
                'signed_by': signer,
                'signed_at': row['signed_at'],
                'signed_by_display': USERS.get(signer, {}).get('display', signer),
            }
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PSC manager signoffs read error: {e}")
    return signoffs


def _psc_week_checkin_questions():
    """Read-only. Reached on every dashboard load via compute_psc_training_stats."""
    return psc_get_week_checkin_questions()


def manager_signoff_psc_week(trainee_key, week_num, signed_by):
    """Manager verifies a week after the trainee completes all items."""
    week_map = _psc_week_item_ids()
    week_num = int(week_num)
    if week_num not in week_map:
        return False, 'Invalid week'
    progress = get_psc_training_progress(trainee_key)
    if not _psc_week_trainee_complete(progress, week_map).get(week_num):
        return False, 'Trainee has not completed all items for this week yet'
    if not is_psc_training_enrolled(trainee_key):
        return False, 'Trainee is not actively enrolled'
    try:
        conn = get_db()
        if not conn:
            return False, 'Database unavailable'
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO psc_training_manager_signoffs (user_key, week_num, signed_by, signed_at)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (user_key, week_num)
            DO UPDATE SET signed_by = EXCLUDED.signed_by, signed_at = NOW()
        ''', (trainee_key, week_num, signed_by))
        conn.commit()
        cur.close()
        conn.close()
        touch_psc_training_activity(trainee_key)
        _notify_psc_week_signed_off(trainee_key, week_num, signed_by)
        return True, None
    except Exception as e:
        print(f"PSC manager signoff error: {e}")
        return False, str(e)


def revoke_psc_manager_signoff(trainee_key, week_num):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            'DELETE FROM psc_training_manager_signoffs WHERE user_key = %s AND week_num = %s',
            (trainee_key, int(week_num)),
        )
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PSC signoff revoke error: {e}")
        return False


def _psc_week_notification_sent(user_key, week_num):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            SELECT 1 FROM psc_training_notifications
            WHERE user_key = %s AND notification_type = %s AND week_num = %s
        ''', (user_key, 'week_complete', week_num))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return bool(row)
    except Exception:
        return False


def _record_psc_week_notification(user_key, week_num):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO psc_training_notifications (user_key, notification_type, week_num)
            VALUES (%s, %s, %s)
            ON CONFLICT (user_key, notification_type, week_num) DO NOTHING
        ''', (user_key, 'week_complete', week_num))
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PSC week notification record error: {e}")
        return False


def _notify_psc_week_signed_off(trainee_key, week_num, signed_by):
    """Email accountability owners when a manager officially signs off a week."""
    from html import escape

    if _psc_week_notification_sent(trainee_key, week_num):
        return
    labels = _psc_week_labels()
    checkins = _psc_week_checkin_questions()
    display = USERS.get(trainee_key, {}).get('display', trainee_key)
    signer = USERS.get(signed_by, {}).get('display', signed_by)
    stats = compute_psc_training_stats(trainee_key)
    oversight_url = f"{HUB_PUBLIC_URL.rstrip('/')}/psc-training/oversight"
    label = labels.get(week_num, f'Week {week_num}')
    checkin = checkins.get(week_num, '')
    subject = f'PSC Training — {display} completed {label} (manager sign-off)'
    text_body = (
        f'{signer} signed off {label} for {display}.\n\n'
        f'Manager check-in: {checkin}\n\n'
        f'Overall progress: {stats["pct"]}% ({stats["done"]} / {stats["total"]} items)\n\n'
        f'Review progress: {oversight_url}'
    )
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:0 auto;">
      <div style="background:#004C8C;padding:18px 22px;border-radius:8px 8px 0 0;">
        <p style="color:white;font-size:17px;font-weight:600;margin:0;">PSC Training — Week Signed Off</p>
      </div>
      <div style="background:#f8fafc;padding:22px;border:1px solid #e2e8f0;border-top:none;border-radius:0 0 8px 8px;">
        <p style="color:#334155;font-size:15px;margin:0 0 12px;">
          <strong>{signer}</strong> signed off <strong>{label}</strong> for <strong>{display}</strong>.
        </p>
        <p style="color:#64748b;font-size:14px;margin:0 0 12px;"><strong>Check-in covered:</strong> {escape(checkin)}</p>
        <p style="color:#64748b;font-size:14px;margin:0 0 16px;">
          Overall progress: {stats['pct']}% ({stats['done']} / {stats['total']} items)
        </p>
        <p style="margin:0;">
          <a href="{oversight_url}" style="color:#004C8C;font-weight:600;">Open PSC Accountability dashboard →</a>
        </p>
      </div>
    </div>
    """
    if _send_psc_accountability_email(subject, text_body, html_body):
        _record_psc_week_notification(trainee_key, week_num)


def _notify_psc_training_feedback(user_key, display_name, message, week_num=None):
    """Email accountability owners when a trainee submits module feedback."""
    from html import escape

    labels = _psc_week_labels()
    if week_num is None or week_num == '':
        week_label = 'General — whole module'
    else:
        week_label = labels.get(int(week_num), f'Week {week_num}')
    oversight_url = f"{HUB_PUBLIC_URL.rstrip('/')}/psc-training/oversight"
    subject = f'PSC Training Feedback — {display_name}'
    text_body = (
        f'New PSC training feedback from {display_name}\n'
        f'Week: {week_label}\n\n'
        f'{message.strip()}\n\n'
        f'Review in hub: {oversight_url}'
    )
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:0 auto;">
      <div style="background:#004C8C;padding:18px 22px;border-radius:8px 8px 0 0;">
        <p style="color:white;font-size:17px;font-weight:600;margin:0;">PSC Training Feedback</p>
      </div>
      <div style="background:#f8fafc;padding:22px;border:1px solid #e2e8f0;border-top:none;border-radius:0 0 8px 8px;">
        <p style="color:#334155;font-size:14px;margin:0 0 8px;"><strong>From:</strong> {display_name}</p>
        <p style="color:#334155;font-size:14px;margin:0 0 16px;"><strong>Week:</strong> {week_label}</p>
        <div style="background:white;border:1px solid #e2e8f0;border-radius:8px;padding:14px;color:#334155;font-size:14px;line-height:1.55;white-space:pre-wrap;">{escape(message.strip())}</div>
        <p style="margin:16px 0 0;">
          <a href="{oversight_url}" style="color:#004C8C;font-weight:600;">Open PSC Accountability dashboard →</a>
        </p>
      </div>
    </div>
    """
    _send_psc_accountability_email(subject, text_body, html_body)


def save_psc_training_progress(user_key, progress_dict):
    """Upsert completion state for training items."""
    if not progress_dict or not isinstance(progress_dict, dict):
        return False
    valid_ids = set(get_all_item_ids())
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        for item_id, completed in progress_dict.items():
            if item_id not in valid_ids:
                continue
            if completed:
                cur.execute('''
                    INSERT INTO psc_training_progress (user_key, item_id, completed, completed_at, updated_at)
                    VALUES (%s, %s, TRUE, NOW(), NOW())
                    ON CONFLICT (user_key, item_id)
                    DO UPDATE SET completed = TRUE, completed_at = NOW(), updated_at = NOW()
                ''', (user_key, item_id))
            else:
                cur.execute('''
                    INSERT INTO psc_training_progress (user_key, item_id, completed, completed_at, updated_at)
                    VALUES (%s, %s, FALSE, NULL, NOW())
                    ON CONFLICT (user_key, item_id)
                    DO UPDATE SET completed = FALSE, completed_at = NULL, updated_at = NOW()
                ''', (user_key, item_id))
        conn.commit()
        cur.close()
        conn.close()
        touch_psc_training_activity(user_key)
        return True
    except Exception as e:
        print(f"PSC training progress save error: {e}")
        return False


def can_psc_training_oversight(user_key):
    """President (admin) and VP Sales track enrolled trainee progress."""
    return is_leadership(user_key)


def is_psc_training_enrolled(user_key):
    """Active enrollment — enrolled consultants (and admin self-preview) see the training module."""
    role = USERS.get(user_key, {}).get('role')
    if role not in ('consultant', 'admin'):
        return False
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            'SELECT 1 FROM psc_training_enrollment WHERE user_key = %s AND active = TRUE AND graduated_at IS NULL',
            (user_key,),
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
        return bool(row)
    except Exception as e:
        print(f"PSC enrollment check error: {e}")
        return False


def get_psc_enrollment(user_key):
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('SELECT * FROM psc_training_enrollment WHERE user_key = %s', (user_key,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return row
    except Exception:
        return None


def list_psc_enrolled_trainees():
    rows = []
    try:
        conn = get_db()
        if not conn:
            return rows
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('''
            SELECT * FROM psc_training_enrollment
            WHERE active = TRUE AND graduated_at IS NULL
            ORDER BY enrolled_at DESC
        ''')
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PSC enrollment list error: {e}")
    result = []
    for row in rows:
        key = row['user_key']
        u = USERS.get(key, {})
        result.append({
            **dict(row),
            'display': u.get('display', row.get('display_name') or key),
        })
    return result


def enroll_psc_trainee(user_key, enrolled_by, manager_key=None):
    user = USERS.get(user_key, {})
    role = user.get('role')
    if role == 'consultant':
        pass
    elif role == 'admin' and user_key == enrolled_by:
        pass  # President previewing the trainee experience
    else:
        return False, 'User cannot be enrolled in PSC training'
    mgr = manager_key or PSC_TRAINING_MANAGER
    try:
        conn = get_db()
        if not conn:
            return False, 'Database unavailable'
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO psc_training_enrollment
            (user_key, enrolled_by, manager_key, enrolled_at, active, graduated_at, last_activity_at)
            VALUES (%s, %s, %s, NOW(), TRUE, NULL, NOW())
            ON CONFLICT (user_key) DO UPDATE SET
                enrolled_by = EXCLUDED.enrolled_by,
                manager_key = EXCLUDED.manager_key,
                enrolled_at = NOW(),
                active = TRUE,
                graduated_at = NULL,
                last_activity_at = NOW()
        ''', (user_key, enrolled_by, mgr))
        conn.commit()
        cur.close()
        conn.close()
        return True, None
    except Exception as e:
        print(f"PSC enroll error: {e}")
        return False, str(e)


def graduate_psc_trainee(user_key):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            UPDATE psc_training_enrollment
            SET graduated_at = NOW(), active = FALSE, last_activity_at = NOW()
            WHERE user_key = %s
        ''', (user_key,))
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PSC graduate error: {e}")
        return False


def unenroll_psc_trainee(user_key):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            'UPDATE psc_training_enrollment SET active = FALSE WHERE user_key = %s',
            (user_key,),
        )
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PSC unenroll error: {e}")
        return False


def touch_psc_training_activity(user_key):
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        cur.execute(
            'UPDATE psc_training_enrollment SET last_activity_at = NOW() WHERE user_key = %s AND active = TRUE',
            (user_key,),
        )
        conn.commit()
        cur.close()
        conn.close()
    except Exception:
        pass


def get_psc_training_notes(user_key):
    """Return {week_num: notes_text} for a trainee."""
    notes = {}
    try:
        conn = get_db()
        if not conn:
            return notes
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT week_num, notes FROM psc_training_notes WHERE user_key = %s',
            (user_key,),
        )
        for row in cur.fetchall():
            notes[row['week_num']] = row['notes'] or ''
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PSC training notes read error: {e}")
    return notes


def save_psc_training_notes(user_key, week_num, notes_text):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO psc_training_notes (user_key, week_num, notes, updated_at)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (user_key, week_num)
            DO UPDATE SET notes = EXCLUDED.notes, updated_at = NOW()
        ''', (user_key, int(week_num), notes_text or ''))
        conn.commit()
        cur.close()
        conn.close()
        touch_psc_training_activity(user_key)
        return True
    except Exception as e:
        print(f"PSC training notes save error: {e}")
        return False


def submit_psc_training_feedback(user_key, display_name, message, week_num=None, feedback_type='improvement'):
    if not message or not message.strip():
        return False
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO psc_training_feedback
            (user_key, display_name, week_num, message, feedback_type)
            VALUES (%s, %s, %s, %s, %s)
        ''', (user_key, display_name, week_num, message.strip(), feedback_type))
        conn.commit()
        cur.close()
        conn.close()
        _notify_psc_training_feedback(user_key, display_name, message, week_num)
        return True
    except Exception as e:
        print(f"PSC training feedback error: {e}")
        return False


def _psc_week_item_ids():
    """Map week number -> list of trainee item IDs for that week.

    Reads the cached curriculum: this only collects ID strings, and it runs on
    every dashboard load. Nothing below writes to the structure.
    """
    onboarding, weeks, core_values, sales_training, company_operations = psc_read_curriculum()
    result = {0: []}

    def collect(week_data):
        ids = []
        for v in week_data.get('videos', []):
            ids.append(v['id'])
        for s in week_data.get('shadowing', []):
            ids.append(s['id'] if isinstance(s, dict) else s)
        for a in week_data.get('additional', []):
            ids.append(a['id'])
        for f in week_data.get('pps_focus', []):
            ids.append(f['id'])
        if week_data.get('book_id'):
            ids.append(week_data['book_id'])
        return ids

    ops_by_week = {}
    for module in company_operations['modules']:
        week_num = module.get('assigned_week', 0)
        ops_by_week.setdefault(week_num, []).extend(item['id'] for item in module['items'])

    result[0] = collect(onboarding) + ops_by_week.get(0, [])
    for section in core_values['sections']:
        for act in section.get('activities', []):
            result[0].append(act['id'])
    for module in sales_training['modules']:
        for item in module['items']:
            result[0].append(item['id'])
    for w in weeks:
        week_num = w['week']
        result[week_num] = collect(w) + ops_by_week.get(week_num, [])
    return result


def compute_psc_training_stats(user_key):
    """Overall and per-week completion stats."""
    progress = get_psc_training_progress(user_key)
    week_map = _psc_week_item_ids()
    signoffs = get_psc_manager_signoffs(user_key)
    checkins = _psc_week_checkin_questions()
    total = count_trackable_items()
    done = sum(1 for i in get_all_item_ids() if progress.get(i))
    week_pcts = []
    for week_num in sorted(week_map.keys()):
        ids = week_map[week_num]
        w_done = sum(1 for i in ids if progress.get(i))
        w_total = len(ids)
        trainee_pct = round((w_done / w_total) * 100) if w_total else 0
        manager_signed = week_num in signoffs
        entry = {
            'week': week_num,
            'done': w_done,
            'total': w_total,
            'trainee_pct': trainee_pct,
            'manager_signed': manager_signed,
            'ready_for_signoff': trainee_pct == 100 and not manager_signed,
            'pct': 100 if manager_signed and trainee_pct == 100 else trainee_pct,
            'checkin': checkins.get(week_num, ''),
        }
        if manager_signed:
            entry['signed_by_display'] = signoffs[week_num]['signed_by_display']
            entry['signed_at'] = signoffs[week_num]['signed_at']
        week_pcts.append(entry)
    pct = round((done / total) * 100) if total else 0
    signed_weeks = sum(1 for w in week_pcts if w['manager_signed'] and w['trainee_pct'] == 100)
    return {
        'done': done,
        'total': total,
        'pct': pct,
        'week_pcts': week_pcts,
        'signed_weeks': signed_weeks,
        'total_weeks': len(week_map),
    }


def get_pm_training_progress(user_key):
    """Return {item_id: True} for completed PM training items."""
    progress = {}
    try:
        conn = get_db()
        if not conn:
            return progress
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT item_id FROM pm_training_progress WHERE user_key = %s AND completed = TRUE',
            (user_key,),
        )
        for row in cur.fetchall():
            progress[row['item_id']] = True
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PM training progress read error: {e}")
    return progress


def save_pm_training_progress(user_key, progress_dict):
    """Upsert completed flags. Open to all logged-in users (under-construction rollout)."""
    if not isinstance(progress_dict, dict):
        return False
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        for item_id, completed in progress_dict.items():
            item_id = str(item_id)[:100]
            done = bool(completed)
            cur.execute('''
                INSERT INTO pm_training_progress (user_key, item_id, completed, completed_at, updated_at)
                VALUES (%s, %s, %s, CASE WHEN %s THEN NOW() ELSE NULL END, NOW())
                ON CONFLICT (user_key, item_id)
                DO UPDATE SET
                    completed = EXCLUDED.completed,
                    completed_at = CASE WHEN EXCLUDED.completed THEN COALESCE(pm_training_progress.completed_at, NOW()) ELSE NULL END,
                    updated_at = NOW()
            ''', (user_key, item_id, done, done))
        conn.commit()
        cur.close()
        conn.close()
        touch_pm_training_activity(user_key)
        return True
    except Exception as e:
        print(f"PM training progress save error: {e}")
        return False


def can_pm_training_oversight(user_key):
    """Admin and Production Manager (Trey) track PM training progress."""
    return is_leadership(user_key)


def is_pm_training_enrolled(user_key):
    """Optional formal enrollment. Module is open to everyone; enrollment is for accountability tracking."""
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            'SELECT 1 FROM pm_training_enrollment WHERE user_key = %s AND active = TRUE AND graduated_at IS NULL',
            (user_key,),
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
        return bool(row)
    except Exception as e:
        print(f"PM enrollment check error: {e}")
        return False


def get_pm_enrollment(user_key):
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('SELECT * FROM pm_training_enrollment WHERE user_key = %s', (user_key,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        return row
    except Exception:
        return None


def list_pm_enrolled_trainees():
    rows = []
    try:
        conn = get_db()
        if not conn:
            return rows
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('''
            SELECT * FROM pm_training_enrollment
            WHERE active = TRUE AND graduated_at IS NULL
            ORDER BY enrolled_at DESC
        ''')
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PM enrollment list error: {e}")
    result = []
    for row in rows:
        key = row['user_key']
        u = USERS.get(key, {})
        result.append({
            **dict(row),
            'display': u.get('display', row.get('display_name') or key),
        })
    return result


def enroll_pm_trainee(user_key, enrolled_by, manager_key=None):
    if user_key not in USERS:
        return False, 'Unknown user'
    mgr = manager_key or PM_TRAINING_MANAGER
    try:
        conn = get_db()
        if not conn:
            return False, 'Database unavailable'
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO pm_training_enrollment
            (user_key, enrolled_by, manager_key, enrolled_at, active, graduated_at, last_activity_at, target_weeks)
            VALUES (%s, %s, %s, NOW(), TRUE, NULL, NOW(), 4)
            ON CONFLICT (user_key) DO UPDATE SET
                enrolled_by = EXCLUDED.enrolled_by,
                manager_key = EXCLUDED.manager_key,
                enrolled_at = NOW(),
                active = TRUE,
                graduated_at = NULL,
                last_activity_at = NOW(),
                target_weeks = 4
        ''', (user_key, enrolled_by, mgr))
        conn.commit()
        cur.close()
        conn.close()
        return True, None
    except Exception as e:
        print(f"PM enroll error: {e}")
        return False, str(e)


def graduate_pm_trainee(user_key):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            UPDATE pm_training_enrollment
            SET graduated_at = NOW(), active = FALSE, last_activity_at = NOW()
            WHERE user_key = %s
        ''', (user_key,))
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PM graduate error: {e}")
        return False


def unenroll_pm_trainee(user_key):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            'UPDATE pm_training_enrollment SET active = FALSE WHERE user_key = %s',
            (user_key,),
        )
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PM unenroll error: {e}")
        return False


def touch_pm_training_activity(user_key):
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        cur.execute(
            'UPDATE pm_training_enrollment SET last_activity_at = NOW() WHERE user_key = %s AND active = TRUE',
            (user_key,),
        )
        conn.commit()
        cur.close()
        conn.close()
    except Exception:
        pass


def get_pm_training_notes(user_key):
    notes = {}
    try:
        conn = get_db()
        if not conn:
            return notes
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT week_num, notes FROM pm_training_notes WHERE user_key = %s',
            (user_key,),
        )
        for row in cur.fetchall():
            notes[row['week_num']] = row['notes'] or ''
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PM training notes read error: {e}")
    return notes


def save_pm_training_notes(user_key, week_num, notes_text):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO pm_training_notes (user_key, week_num, notes, updated_at)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (user_key, week_num)
            DO UPDATE SET notes = EXCLUDED.notes, updated_at = NOW()
        ''', (user_key, int(week_num), notes_text or ''))
        conn.commit()
        cur.close()
        conn.close()
        touch_pm_training_activity(user_key)
        return True
    except Exception as e:
        print(f"PM training notes save error: {e}")
        return False


def submit_pm_training_feedback(user_key, display_name, message, week_num=None, feedback_type='improvement'):
    if not message or not message.strip():
        return False
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO pm_training_feedback
            (user_key, display_name, week_num, message, feedback_type)
            VALUES (%s, %s, %s, %s, %s)
        ''', (user_key, display_name, week_num, message.strip(), feedback_type))
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"PM training feedback error: {e}")
        return False


def get_pm_manager_signoffs(user_key):
    signoffs = {}
    try:
        conn = get_db()
        if not conn:
            return signoffs
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT week_num, signed_by, signed_at FROM pm_training_manager_signoffs WHERE user_key = %s',
            (user_key,),
        )
        for row in cur.fetchall():
            signer = row['signed_by']
            signoffs[row['week_num']] = {
                'signed_by': signer,
                'signed_at': row['signed_at'],
                'signed_by_display': USERS.get(signer, {}).get('display', signer),
            }
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PM manager signoffs read error: {e}")
    return signoffs


def compute_pm_training_stats(user_key):
    progress = get_pm_training_progress(user_key)
    week_map = get_pm_week_item_ids()
    signoffs = get_pm_manager_signoffs(user_key)
    checkins = get_pm_week_checkin_questions()
    total = count_pm_trackable_items()
    all_ids = get_pm_training_item_ids()
    done = sum(1 for i in all_ids if progress.get(i))
    week_pcts = []
    for week_num in sorted(week_map.keys()):
        ids = week_map[week_num]
        w_done = sum(1 for i in ids if progress.get(i))
        w_total = len(ids)
        trainee_pct = round((w_done / w_total) * 100) if w_total else 0
        manager_signed = week_num in signoffs
        entry = {
            'week': week_num,
            'done': w_done,
            'total': w_total,
            'trainee_pct': trainee_pct,
            'manager_signed': manager_signed,
            'ready_for_signoff': trainee_pct == 100 and not manager_signed,
            'pct': 100 if manager_signed and trainee_pct == 100 else trainee_pct,
            'checkin': checkins.get(week_num, ''),
        }
        if manager_signed:
            entry['signed_by_display'] = signoffs[week_num]['signed_by_display']
            entry['signed_at'] = signoffs[week_num]['signed_at']
        week_pcts.append(entry)
    pct = round((done / total) * 100) if total else 0
    signed_weeks = sum(1 for w in week_pcts if w['manager_signed'] and w['trainee_pct'] == 100)
    return {
        'done': done,
        'total': total,
        'pct': pct,
        'week_pcts': week_pcts,
        'signed_weeks': signed_weeks,
        'total_weeks': len(week_map),
    }


def list_pm_training_feedback(limit=50):
    rows = []
    try:
        conn = get_db()
        if not conn:
            return rows
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('''
            SELECT * FROM pm_training_feedback
            ORDER BY submitted_at DESC
            LIMIT %s
        ''', (limit,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PM feedback list error: {e}")
    return rows



def can_access_psc_roleplay(user_key):
    """Enrolled trainees and oversight managers (VP Sales / admin) may use Practice Arena."""
    return is_psc_training_enrolled(user_key) or can_psc_training_oversight(user_key)


_ROLEPLAY_PERSONA_SYSTEM = """You are playing a role in a sales training simulation for Pure Property Solutions (PPS),
a multi-family and commercial property contractor. Stay in character at all times.

CHARACTER: {persona}

RULES:
- Speak only as this character. Never mention being an AI, a simulation, or training.
- Reply in 2–4 sentences, conversational and realistic. One point or question at a time.
- Be realistic, not a pushover: raise natural objections, ask follow-up questions,
  and push back on vague or salesy answers.
- Reward good behavior realistically: if the trainee asks smart discovery questions,
  explains value concretely, or proposes a clear next step, warm up gradually.
- If the trainee makes a promise a contractor shouldn't make on the spot (pricing changes,
  schedule guarantees, scope additions without approval), react the way a real client would —
  take them up on it or press for it in writing. Do not correct or coach them mid-conversation.
- If the trainee is unprofessional or gives up, respond in character and let the
  conversation land where it naturally would.
- Never discuss topics outside the scenario. If the trainee goes off-topic, steer back
  in character ("Let's stay on the project — ...").
"""


def _strip_json_fences(raw):
    text = (raw or '').strip()
    if text.startswith('```'):
        text = text.split('\n', 1)[-1]
        if text.endswith('```'):
            text = text.rsplit('```', 1)[0]
        text = text.strip()
        if text.lower().startswith('json'):
            text = text[4:].strip()
    return text


def _claude_roleplay_call(system_prompt, messages, max_tokens, timeout=60.0):
    import time
    import anthropic
    cl = anthropic.Anthropic(api_key=CLAUDE_API_KEY, timeout=timeout)
    last_err = None
    for attempt in range(2):
        try:
            msg = cl.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=max_tokens,
                system=system_prompt,
                messages=messages,
            )
            return msg.content[0].text.strip()
        except Exception as e:
            last_err = e
            err_name = type(e).__name__
            transient = err_name in (
                'APITimeoutError', 'APIConnectionError', 'RateLimitError', 'InternalServerError',
            ) or 'timeout' in str(e).lower() or 'overloaded' in str(e).lower()
            if attempt == 0 and transient:
                time.sleep(1.5)
                continue
            print(f"Roleplay Claude error ({err_name}): {e}")
            raise last_err


def _validate_roleplay_messages(messages, max_turns):
    if not isinstance(messages, list):
        return None, 'Invalid messages format'
    if len(messages) > max_turns * 2:
        return None, f'Too many messages (max {max_turns * 2})'
    total_chars = 0
    cleaned = []
    for m in messages:
        if not isinstance(m, dict):
            return None, 'Invalid message entry'
        role = m.get('role')
        content = m.get('content')
        if role not in ('user', 'assistant'):
            return None, 'Invalid message role'
        if not isinstance(content, str):
            return None, 'Invalid message content'
        if len(content) > 2000:
            return None, 'Each message must be 2,000 characters or fewer'
        total_chars += len(content)
        cleaned.append({'role': role, 'content': content})
    if total_chars > 50000:
        return None, 'Conversation payload too large'
    return cleaned, None


def _count_trainee_turns(messages):
    return sum(1 for m in messages if m.get('role') == 'user')


def _roleplay_usage_today(user_key):
    try:
        conn = get_db()
        if not conn:
            return 0, 0
        cur = conn.cursor()
        cur.execute(
            '''SELECT turn_count, grade_count FROM psc_roleplay_daily_usage
               WHERE user_key = %s AND usage_date = CURRENT_DATE''',
            (user_key,),
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
        if row:
            return row[0] or 0, row[1] or 0
        return 0, 0
    except Exception as e:
        print(f"Roleplay usage read error: {e}")
        return 0, 0


def _increment_roleplay_usage(user_key, turn_delta=0, grade_delta=0):
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute(
            '''
            INSERT INTO psc_roleplay_daily_usage (user_key, usage_date, turn_count, grade_count)
            VALUES (%s, CURRENT_DATE, %s, %s)
            ON CONFLICT (user_key, usage_date) DO UPDATE SET
                turn_count = psc_roleplay_daily_usage.turn_count + EXCLUDED.turn_count,
                grade_count = psc_roleplay_daily_usage.grade_count + EXCLUDED.grade_count
            ''',
            (user_key, turn_delta, grade_delta),
        )
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception as e:
        print(f"Roleplay usage increment error: {e}")
        return False


def get_roleplay_user_stats(user_key):
    """Best score and attempt count per scenario for picker display."""
    stats = {}
    for sc in PSC_ROLEPLAY_SCENARIOS:
        stats[sc['id']] = {'attempts': 0, 'best_overall': None, 'best_result': None}
    try:
        conn = get_db()
        if not conn:
            return stats
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            '''SELECT scenario_id, COUNT(*) AS attempts, MAX(overall) AS best_overall
               FROM psc_roleplay_sessions WHERE user_key = %s GROUP BY scenario_id''',
            (user_key,),
        )
        agg = {r['scenario_id']: r for r in cur.fetchall()}
        cur.execute(
            '''SELECT DISTINCT ON (scenario_id) scenario_id, overall, result
               FROM psc_roleplay_sessions
               WHERE user_key = %s
               ORDER BY scenario_id, overall DESC, created_at DESC''',
            (user_key,),
        )
        best_rows = {r['scenario_id']: r for r in cur.fetchall()}
        cur.close()
        conn.close()
        for sid, row in agg.items():
            if sid in stats:
                stats[sid]['attempts'] = row['attempts']
                best = best_rows.get(sid)
                if best:
                    stats[sid]['best_overall'] = round(float(best['overall']), 1)
                    stats[sid]['best_result'] = best['result']
    except Exception as e:
        print(f"Roleplay user stats error: {e}")
    return stats


def get_roleplay_summary_for_oversight(user_key):
    """Aggregate role-play data for accountability page."""
    summary = {
        'total_attempts': 0,
        'pass_count': 0,
        'scenarios': {},
        'latest_key_moment': None,
        'latest_at': None,
        'latest_scenario_title': None,
    }
    try:
        conn = get_db()
        if not conn:
            return summary
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            'SELECT COUNT(*) AS cnt FROM psc_roleplay_sessions WHERE user_key = %s',
            (user_key,),
        )
        summary['total_attempts'] = cur.fetchone()['cnt'] or 0
        cur.execute(
            '''SELECT COUNT(*) AS cnt FROM psc_roleplay_sessions
               WHERE user_key = %s AND result = %s''',
            (user_key, 'pass'),
        )
        summary['pass_count'] = cur.fetchone()['cnt'] or 0
        cur.execute(
            '''SELECT scenario_id, COUNT(*) AS attempts, MAX(overall) AS best_overall,
                      BOOL_OR(result = 'pass') AS ever_passed
               FROM psc_roleplay_sessions WHERE user_key = %s GROUP BY scenario_id''',
            (user_key,),
        )
        for row in cur.fetchall():
            sc = get_roleplay_scenario(row['scenario_id'])
            summary['scenarios'][row['scenario_id']] = {
                'title': sc['title'] if sc else row['scenario_id'],
                'attempts': row['attempts'],
                'best_overall': round(float(row['best_overall']), 1) if row['best_overall'] is not None else None,
                'ever_passed': bool(row['ever_passed']),
            }
        cur.execute(
            '''SELECT scenario_id, feedback, created_at FROM psc_roleplay_sessions
               WHERE user_key = %s ORDER BY created_at DESC LIMIT 1''',
            (user_key,),
        )
        latest = cur.fetchone()
        cur.close()
        conn.close()
        if latest:
            summary['latest_at'] = latest['created_at']
            sc = get_roleplay_scenario(latest['scenario_id'])
            summary['latest_scenario_title'] = sc['title'] if sc else latest['scenario_id']
            try:
                fb = json.loads(latest['feedback'])
                summary['latest_key_moment'] = (fb.get('key_moment') or '').strip() or None
            except (json.JSONDecodeError, TypeError):
                pass
    except Exception as e:
        print(f"Roleplay oversight summary error: {e}")
    return summary


def get_roleplay_history_rows(user_key):
    rows = []
    try:
        conn = get_db()
        if not conn:
            return rows
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            '''SELECT scenario_id, overall, result, created_at
               FROM psc_roleplay_sessions WHERE user_key = %s
               ORDER BY created_at DESC LIMIT 100''',
            (user_key,),
        )
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Roleplay history error: {e}")
    return rows


def _grade_roleplay_session(scenario, messages):
    objectives = '\n'.join(f'- {o}' for o in scenario.get('objectives', []))
    transcript = '\n'.join(
        f"{'TRAINEE' if m['role'] == 'user' else 'CLIENT'}: {m['content']}"
        for m in messages
    )
    grader_focus = scenario.get('grader_focus', '')
    system = f"""You are grading a PSC sales training role-play for Pure Property Solutions (PPS).

SCENARIO: {scenario['title']}
TRAINEE BRIEF: {scenario['trainee_brief']}
OBJECTIVES:
{objectives}

SCENARIO-SPECIFIC GRADING EMPHASIS:
{grader_focus}

PPS VOICE RULES:
{PSC_ROLEPLAY_GRADER_RULES}

Score each category 1–5 (5 = excellent). Integrity measures whether the trainee made promises a contractor
should not make on the spot (pricing, scope additions without approval, schedule guarantees).

PASS CRITERIA (you must apply): result is "pass" only if overall >= 3.5 AND integrity score >= 4.
Otherwise result is "practice_again". Integrity is a hard gate.

Respond with ONLY valid JSON (no markdown fences) in exactly this shape:
{{
  "scores": {{
    "discovery": {{"score": 4, "note": "..."}},
    "value_communication": {{"score": 3, "note": "..."}},
    "pps_voice": {{"score": 5, "note": "..."}},
    "integrity": {{"score": 5, "note": "..."}},
    "next_step_close": {{"score": 2, "note": "..."}}
  }},
  "overall": 3.8,
  "result": "pass",
  "strengths": ["...", "..."],
  "improvements": ["...", "..."],
  "key_moment": "Short quote of the trainee's best or most costly line and why it mattered."
}}

Recalculate overall as the average of the five category scores. Set result from the pass criteria above."""

    user_content = f"TRANSCRIPT:\n{transcript[:30000]}"
    raw = None
    for attempt in range(2):
        try:
            raw = _claude_roleplay_call(
                system,
                [{'role': 'user', 'content': user_content}],
                max_tokens=1200,
            )
            parsed = json.loads(_strip_json_fences(raw))
            scores = parsed.get('scores') or {}
            required = ('discovery', 'value_communication', 'pps_voice', 'integrity', 'next_step_close')
            for key in required:
                if key not in scores or 'score' not in scores[key]:
                    raise ValueError(f'Missing score: {key}')
            vals = [float(scores[k]['score']) for k in required]
            overall = round(sum(vals) / len(vals), 1)
            parsed['overall'] = overall
            integrity = float(scores['integrity']['score'])
            parsed['result'] = 'pass' if overall >= 3.5 and integrity >= 4 else 'practice_again'
            return parsed, None
        except (json.JSONDecodeError, ValueError, KeyError, TypeError) as e:
            print(f"Roleplay grade parse error (attempt {attempt + 1}): {e}")
            if attempt == 1:
                return None, 'Could not parse feedback — please try ending the session again.'
        except Exception:
            return None, 'The practice partner is unavailable — try again in a minute.'
    return None, 'Could not parse feedback — please try ending the session again.'


def _save_roleplay_session(user_key, scenario_id, messages, feedback, turn_count):
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor()
        cur.execute(
            '''
            INSERT INTO psc_roleplay_sessions
            (user_key, scenario_id, transcript, feedback, overall, result, turn_count)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING id
            ''',
            (
                user_key,
                scenario_id,
                json.dumps(messages),
                json.dumps(feedback),
                float(feedback['overall']),
                feedback['result'],
                turn_count,
            ),
        )
        session_id = cur.fetchone()[0]
        conn.commit()
        cur.close()
        conn.close()
        return session_id
    except Exception as e:
        print(f"Roleplay session save error: {e}")
        return None


def _internal_api_ok():
    if not INTERNAL_API_KEY:
        return False
    api_key = (request.headers.get('X-API-Key') or '').strip()
    return api_key == INTERNAL_API_KEY


# ── ROUTES ──────────────────────────────────────────────────────────────────────

def _http_get_json(url, headers=None, timeout=8, retries=3):
    import time
    import urllib.error
    import urllib.request

    hdrs = {'User-Agent': 'PPS-Hub-HealthCheck/1.0'}
    if headers:
        hdrs.update(headers)

    last_err = None
    for attempt in range(retries):
        try:
            req = urllib.request.Request(url, headers=hdrs, method='GET')
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                return json.loads(resp.read().decode('utf-8'))
        except urllib.error.HTTPError as e:
            last_err = e
            if e.code in (429, 502, 503, 504) and attempt < retries - 1:
                time.sleep(1.5 * (attempt + 1))
                continue
            raise
        except Exception as e:
            last_err = e
            if attempt < retries - 1:
                time.sleep(1.0)
                continue
            raise
    if last_err:
        raise last_err


def _run_system_health_checks():
    """Verify env alignment and cross-service connectivity."""
    import urllib.error
    checks = []

    def add(name, ok, detail='', error='', transient=False):
        checks.append({
            'name': name,
            'ok': bool(ok),
            'detail': detail,
            'error': error,
            'transient': bool(transient),
        })

    add('secret_key', os.environ.get('SECRET_KEY', '').strip() != '')
    add('internal_api_key', bool(INTERNAL_API_KEY))
    add('database_url', bool(DATABASE_URL))

    smtp_host = os.environ.get('SMTP_HOST', '').strip()
    smtp_user = os.environ.get('SMTP_USER', '').strip()
    smtp_pass = os.environ.get('SMTP_PASS', '').strip()
    smtp_ok = bool(smtp_host and smtp_user and smtp_pass)
    add(
        'smtp_configured',
        smtp_ok,
        detail='feedback & comparison emails' if smtp_ok else '',
        error='' if smtp_ok else 'Set SMTP_HOST, SMTP_USER, and SMTP_PASS on Render',
    )
    notify = _hub_notify_recipients()
    add(
        'hub_notify_email',
        bool(notify),
        detail=', '.join(notify) if notify else '',
        error='' if notify else 'Set HUB_NOTIFY_EMAIL or use default',
    )

    if DATABASE_URL:
        try:
            conn = get_db()
            if conn:
                cur = conn.cursor()
                cur.execute('SELECT 1')
                add('database_connect', True)
                for table in ('feedback', 'proposal_diffs', 'proposal_log', 'ppm_log', 'subscope_log'):
                    cur.execute(
                        "SELECT EXISTS (SELECT FROM information_schema.tables WHERE table_name = %s)",
                        (table,),
                    )
                    exists = cur.fetchone()[0]
                    add(f'table_{table}', bool(exists), error='' if exists else 'missing')
                cur.close()
                conn.close()
            else:
                add('database_connect', False, error='get_db returned None')
        except Exception as e:
            add('database_connect', False, error=str(e))
    else:
        add('database_connect', False, error='DATABASE_URL not set')

    # Local check only — do not HTTP-call this hub from the same gunicorn worker (deadlocks with 1 worker).
    add('hub_internal_ping', bool(INTERNAL_API_KEY), detail='configured locally' if INTERNAL_API_KEY else '')

    for label, base_url in (('proposal_tool', PROPOSAL_URL),):
        if not base_url:
            add(label, False, error=f'{label.upper()} URL not configured')
            continue
        try:
            data = _http_get_json(base_url.rstrip('/') + '/health', timeout=6)
            add(label, data.get('ok'), detail=base_url)
            hub_url = (data.get('hub_url') or data.get('hub_public_url') or '').rstrip('/')
            hub_expected = HUB_PUBLIC_URL.rstrip('/')
            if hub_url:
                add(
                    f'{label}_hub_link',
                    hub_url == hub_expected,
                    detail=hub_url,
                    error='' if hub_url == hub_expected else f'expected {hub_expected}',
                )
            sso_ok = data.get('sso_configured')
            if sso_ok is not None:
                add(f'{label}_sso', bool(sso_ok))
        except urllib.error.HTTPError as e:
            if e.code == 429:
                add(
                    label, False,
                    detail=base_url,
                    error='Rate limited — service is up but temporarily throttled; retry shortly',
                    transient=True,
                )
            else:
                add(label, False, error=str(e))
        except Exception as e:
            add(label, False, error=str(e))

    try:
        from daily_digest import (
            digest_recipients,
            _load_last_run,
            _load_sent_date,
            report_date_for_run,
        )
        recips = digest_recipients()
        add(
            'daily_digest_recipients',
            bool(recips),
            detail=', '.join(recips) if recips else '',
            error='' if recips else 'Set DAILY_DIGEST_EMAIL or HUB_NOTIFY_EMAIL',
        )
        enabled = os.environ.get('DAILY_DIGEST_ENABLED', 'true').strip().lower() in ('1', 'true', 'yes')
        add('daily_digest_enabled', enabled)
        if DATABASE_URL:
            last_run = _load_last_run(get_db)
            last_sent = _load_sent_date(get_db)
            expected = report_date_for_run().isoformat()
            if last_run:
                lr_detail = []
                if last_run.get('skipped'):
                    lr_detail.append(f"skipped: {last_run.get('reason', '?')}")
                elif last_run.get('sent'):
                    lr_detail.append(f"sent {last_run.get('report_date', '?')} ({last_run.get('item_count', 0)} items)")
                elif last_run.get('email_failed'):
                    lr_detail.append('email failed')
                if last_run.get('at'):
                    lr_detail.append(f"at {last_run['at']}")
                add(
                    'daily_digest_last_run',
                    bool(last_run.get('sent')) or not last_run.get('email_failed'),
                    detail=' · '.join(lr_detail) if lr_detail else 'no runs recorded',
                    error=last_run.get('warning') or '',
                )
            else:
                add('daily_digest_last_run', False, error='No digest run logged yet')
            add(
                'daily_digest_last_sent',
                last_sent == expected,
                detail=f'last sent for {last_sent or "never"}; expect {expected}',
                error='' if last_sent == expected else 'Yesterday digest may not have been delivered',
            )
    except Exception as e:
        add('daily_digest_status', False, error=str(e))

    ok = all(c['ok'] or c.get('transient') for c in checks)
    return {'ok': ok, 'checks': checks}


# ── PWA service worker ──────────────────────────────────────────────────────
#
# Served from the root because a service worker cannot control anything above
# its own path — one at /static/sw.js would only ever cover /static/.

SERVICE_WORKER_DISABLED = os.environ.get('SERVICE_WORKER_DISABLED', '').strip().lower() in ('1', 'true', 'yes')

# Every installed copy re-fetches /sw.js periodically and installs it if the
# bytes differ. This stub is how a bad worker gets removed from phones we do not
# have: set SERVICE_WORKER_DISABLED=true and each device unregisters itself on
# its next check. A normal deploy cannot do that — the old worker keeps running.
_SW_KILL_STUB = (
    "self.addEventListener('install', () => self.skipWaiting());\n"
    "self.addEventListener('activate', (e) => {\n"
    "  e.waitUntil((async () => {\n"
    "    const names = await caches.keys();\n"
    "    await Promise.all(names.map((n) => caches.delete(n)));\n"
    "    await self.registration.unregister();\n"
    "    const clients = await self.clients.matchAll({type: 'window'});\n"
    "    clients.forEach((c) => c.navigate(c.url));\n"
    "  })());\n"
    "});\n"
)


def _sw_version():
    """Changes the file's bytes per deploy so browsers see an update.

    Falls back to a date stamp because RENDER_GIT_COMMIT is not set on the
    service (see the System State panel). With the commit set this is exact;
    without it, a deploy on the same day will not bust the static cache — bump
    the fallback by hand if that ever matters.
    """
    return (os.environ.get('RENDER_GIT_COMMIT') or '')[:12] or '2026-08-23a'


@app.route('/sw.js')
def service_worker():
    if SERVICE_WORKER_DISABLED:
        body = _SW_KILL_STUB
    else:
        try:
            with open(os.path.join(app.static_folder, 'sw.js'), 'r', encoding='utf-8') as fh:
                body = fh.read().replace('__SW_VERSION__', _sw_version())
        except Exception as e:
            print(f'service worker read error: {e}')
            body = _SW_KILL_STUB
    resp = make_response(body)
    resp.headers['Content-Type'] = 'application/javascript'
    # Never let a CDN or browser pin the worker itself — that is how a bad one
    # becomes unremovable.
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Service-Worker-Allowed'] = '/'
    return resp


@app.route('/offline')
def offline_page():
    """Shown only when a navigation fails with no network. Precached by the
    worker, so it must not reference anything it would have to fetch."""
    resp = make_response(render_template('offline.html'))
    resp.headers['Cache-Control'] = 'no-cache'
    return resp


@app.route('/health')
def health():
    db_ok = False
    if DATABASE_URL:
        try:
            conn = get_db()
            if conn:
                cur = conn.cursor()
                cur.execute('SELECT 1')
                cur.close()
                conn.close()
                db_ok = True
        except Exception:
            pass
    # Public /health must not list recipient emails (ops detail belongs on
    # /health/deep, which is admin-or-API-key gated). Cron only needs a 200
    # body to wake a sleeping instance — counts and booleans are enough.
    digest_recipient_count = 0
    notify_recipient_count = 0
    digest_last = None
    digest_last_sent = None
    try:
        from daily_digest import digest_recipients, _load_last_run, _load_sent_date
        digest_recipient_count = len(digest_recipients() or [])
        if db_ok:
            digest_last = _load_last_run(get_db)
            digest_last_sent = _load_sent_date(get_db)
    except Exception:
        pass
    try:
        notify_recipient_count = len(_hub_notify_recipients() or [])
    except Exception:
        pass
    return jsonify({
        'ok': True,
        'service': 'hub',
        'hub_public_url': HUB_PUBLIC_URL,
        'proposal_url': PROPOSAL_URL,
        'secret_configured': os.environ.get('SECRET_KEY', '').strip() != '',
        'internal_api_configured': bool(INTERNAL_API_KEY),
        'database_configured': bool(DATABASE_URL),
        'database_connected': db_ok,
        # Sanitized connect failure for ops (no password; may include host/role).
        'database_error': ('' if db_ok else (_DB_LAST_ERROR or ''))[:240],
        'smtp_configured': bool(
            os.environ.get('SMTP_HOST', '').strip()
            and os.environ.get('SMTP_USER', '').strip()
            and os.environ.get('SMTP_PASS', '').strip()
        ),
        'daily_digest_enabled': os.environ.get('DAILY_DIGEST_ENABLED', 'true').strip().lower() in ('1', 'true', 'yes'),
        'daily_digest_recipient_count': digest_recipient_count,
        'daily_digest_last_run': digest_last,
        'daily_digest_last_sent_date': digest_last_sent,
        'hub_notify_recipient_count': notify_recipient_count,
        'resend_configured': bool(os.environ.get('RESEND_API_KEY', '').strip()),
        'claude_configured': bool(CLAUDE_API_KEY),
    })


@app.route('/health/deep')
def health_deep():
    if not (_internal_api_ok() or session.get('role') == 'admin'):
        return jsonify({'error': 'Unauthorized'}), 401
    report = _run_system_health_checks()
    return jsonify(report)


@app.route('/api/internal/ping')
def internal_ping():
    if not _internal_api_ok():
        return jsonify({'ok': False, 'error': 'Invalid or missing API key'}), 401
    return jsonify({'ok': True, 'service': 'hub'})


@app.route('/api/cron/daily-digest', methods=['POST'])
def cron_daily_digest():
    """Nightly team activity digest — triggered by Render cron at midnight US/Eastern."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401

    from datetime import date as date_type
    from daily_digest import run_daily_digest

    force = request.args.get('force', '').lower() in ('1', 'true', 'yes')
    date_param = request.args.get('date', '').strip()
    date_override = None
    if date_param:
        try:
            date_override = date_type.fromisoformat(date_param)
        except ValueError:
            return jsonify({'error': 'Invalid date (use YYYY-MM-DD)'}), 400

    try:
        result = run_daily_digest(
            get_db,
            USERS,
            _format_template_label,
            _send_digest_email,
            force=force,
            date_override=date_override,
        )
        return jsonify(result), 200
    except Exception as e:
        print(f'Daily digest cron error: {e}')
        import traceback
        traceback.print_exc()
        return _api_error(e, ok=False)


@app.route('/api/cron/weekly-tp-compliance', methods=['POST'])
def cron_weekly_tp_compliance():
    """Weekly Trade Partner insurance compliance digest — Stephanie + Thomas only."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401

    try:
        recipients = [USERS['stephanie_whetstone']['email'], USERS['thomas_ellison']['email']]
        result = insurance_compliance.run_weekly_compliance_check(
            get_db,
            _send_digest_email,
            recipients,
        )
        system_state.record_job_run(get_db, 'weekly_tp_compliance', result)
        return jsonify(result), 200
    except Exception as e:
        print(f'Weekly TP compliance cron error: {e}')
        import traceback
        traceback.print_exc()
        return _api_error(e, ok=False)


@app.route('/api/cron/weekly-recap', methods=['POST'])
def cron_weekly_recap():
    """Monday-morning team recap — the one Hub email that goes TO the team.

    Everyone with an email on the roster gets their own copy, with their own
    row highlighted. Unlike the nightly digest (Thomas only), this is not gated
    on a send window: the cron fires Mondays and this endpoint does what it is
    asked, so a manual Trigger Run in Render works without a force flag.
    """
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401
    force = (request.args.get('force') or '').strip().lower() in ('1', 'true', 'yes')
    try:
        result = weekly_recap.run_weekly_recap(
            get_db, USERS, _send_digest_email, force=force,
        )
        system_state.record_job_run(get_db, 'weekly_recap', result)
        return jsonify(result), 200
    except Exception as e:
        print(f'Weekly recap cron error: {e}')
        import traceback
        traceback.print_exc()
        return _api_error(e, ok=False)


@app.route('/api/cron/weekly-crm-sync', methods=['POST'])
def cron_weekly_crm_sync():
    """Weekly sync: new Monday CRM contacts into the Hub /clients picker."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401

    try:
        result = crm_contact_sync.run_weekly_crm_sync(
            get_db,
            _send_digest_email,
            [USERS['thomas_ellison']['email']],
        )
        system_state.record_job_run(get_db, 'weekly_crm_sync', result)
        return jsonify(result), 200
    except Exception as e:
        print(f'Weekly CRM contact sync cron error: {e}')
        import traceback
        traceback.print_exc()
        return _api_error(e, ok=False)


@app.route('/api/cron/daily-estimate-check', methods=['POST'])
def cron_daily_estimate_check():
    """Daily: Estimates board new-assignment notification + open-work reminder."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401

    try:
        result = estimate_assignments.run_daily_estimate_check(get_db, _send_digest_email)
        system_state.record_job_run(get_db, 'daily_estimate_check', result)
        return jsonify(result), 200
    except Exception as e:
        print(f'Daily estimate check cron error: {e}')
        import traceback
        traceback.print_exc()
        return _api_error(e, ok=False)


@app.route('/')
def index():
    if session.get('user_key'):
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))


def _post_login_redirect():
    nxt = safe_next_url(session.pop('login_next', None) or request.args.get('next', ''))
    # Never send a logged-in user back into the login screen (redirect loop).
    if nxt:
        return redirect(nxt)
    return redirect(url_for('dashboard'))


def _safe_check_password(stored_hash, password):
    if not stored_hash or not password:
        return False
    try:
        return check_password_hash(stored_hash, password)
    except (ValueError, TypeError) as e:
        print(f'Password hash check failed: {e}')
        return False


def _ensure_hub_users_password_schema(cur):
    """Idempotent schema fixes so password resets don't fail on older DBs."""
    try:
        cur.execute('ALTER TABLE hub_users ADD COLUMN IF NOT EXISTS must_change_password BOOLEAN DEFAULT FALSE')
    except Exception as e:
        print(f'hub_users must_change_password migrate: {e}')
    try:
        # scrypt hashes are fine under 500, but TEXT avoids edge-case truncations forever
        cur.execute('ALTER TABLE hub_users ALTER COLUMN password_hash TYPE TEXT')
    except Exception as e:
        print(f'hub_users password_hash migrate: {e}')
    try:
        # Bumped on every password write; stamped into the session at login and
        # re-checked in require_login. This is what makes a password reset log
        # out other devices. It matters now that sessions live 30 days: without
        # it, a reset would not evict the session on a lost phone until the
        # phone had been idle a month, which is not a remedy.
        cur.execute(
            'ALTER TABLE hub_users ADD COLUMN IF NOT EXISTS password_epoch INTEGER DEFAULT 0'
        )
    except Exception as e:
        print(f'hub_users password_epoch migrate: {e}')


def _upsert_hub_user_password(user_key, new_password, must_change=False):
    """Create or update hub_users row — fixes login when user is in USERS but missing from DB.

    Returns (ok: bool, action_or_error: str).
    """
    user_def = USERS.get(user_key)
    if not user_def:
        return False, 'unknown_user'
    if not new_password or len(new_password) < 6:
        return False, 'password_too_short'
    try:
        # Explicit method — stable length/format across Werkzeug versions
        hashed = generate_password_hash(new_password, method='pbkdf2:sha256')
    except Exception as e:
        print(f'generate_password_hash failed: {e}')
        return False, 'hash_failed'

    conn = None
    try:
        conn = get_db()
        if not conn:
            return False, 'no_db'
        cur = conn.cursor()
        _ensure_hub_users_password_schema(cur)

        cur.execute('SELECT id FROM hub_users WHERE user_key = %s', (user_key,))
        exists = cur.fetchone()
        if exists:
            cur.execute(
                'UPDATE hub_users SET password_hash = %s, must_change_password = %s, '
                'password_epoch = COALESCE(password_epoch, 0) + 1 WHERE user_key = %s',
                (hashed, bool(must_change), user_key),
            )
            action = 'updated'
        else:
            cur.execute(
                '''INSERT INTO hub_users
                   (user_key, display_name, password_hash, role, must_change_password)
                   VALUES (%s, %s, %s, %s, %s)''',
                (
                    user_key,
                    user_def['display'],
                    hashed,
                    user_def.get('role', 'consultant'),
                    bool(must_change),
                ),
            )
            action = 'created'
        conn.commit()
        cur.close()
        return True, action
    except Exception as e:
        print(f'_upsert_hub_user_password error for {user_key}: {e}')
        try:
            if conn:
                conn.rollback()
        except Exception:
            pass
        return False, f'db_error:{type(e).__name__}'
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass


def _establish_session(user_key, user_def, db_user=None):
    session.permanent = True
    session['user_key'] = user_key
    session['display_name'] = user_def.get('display', '')
    session['user_email'] = user_def.get('email', '')
    session['role'] = user_def.get('role', 'consultant')
    session['tier'] = user_def.get('tier', DEFAULT_TIER)
    session['proposal_access'] = get_user_proposal_access(user_key)
    # Stamp the password generation so a reset invalidates sessions on other
    # devices. Without this a 30-day session outlives the password that opened
    # it, and "reset their password" stops being a real remedy for a lost phone.
    session['pw_epoch'] = int((db_user or {}).get('password_epoch') or 0)
    session['pw_checked_at'] = time.time()
    if db_user and db_user.get('must_change_password'):
        session['must_change_password'] = True


def _pricing_defaults():
    from estimators.pricing_defaults import get_pricing_defaults
    return get_pricing_defaults(get_db)


def _admin_inbox_counts():
    """Unread feedback and proposal comparison submissions for admin badges."""
    unread_feedback = 0
    unread_diffs = 0
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute(
                '''SELECT COUNT(*) FROM feedback
                   WHERE read_by_admin = FALSE
                   AND COALESCE(archived, FALSE) = FALSE'''
            )
            unread_feedback = cur.fetchone()[0] or 0
            cur.execute('SELECT COUNT(*) FROM proposal_diffs WHERE reviewed_by_admin = FALSE')
            unread_diffs = cur.fetchone()[0] or 0
            cur.close()
            conn.close()
    except Exception as e:
        print(f'Admin inbox counts error: {e}')
    return unread_feedback, unread_diffs


def _pricing_summary_for_dashboard():
    """Compact pricing meta for the admin dashboard lane."""
    d = _pricing_defaults()
    sd, rd = d.get('siding', {}), d.get('roofing', {})
    updated = d.get('updated_at')
    updated_label = ''
    if updated and hasattr(updated, 'strftime'):
        updated_label = updated.strftime('%b %d, %Y')
    elif updated:
        updated_label = str(updated)[:10]
    return {
        'updated_label': updated_label,
        'updated_by_name': d.get('updated_by_name') or '',
        'is_custom': bool(updated),
        'siding_labor': sd.get('labor_per_sq'),
        'roofing_labor': rd.get('labor_per_sq'),
        'gutter_lf': d.get('gutter', {}).get('gutter_price_per_lf'),
        'painting_hour': d.get('painting', {}).get('labor_per_hour'),
    }


def _resolve_login_user_key(raw):
    """Map form value or email to a USERS key (case-insensitive email OK)."""
    raw = (raw or '').strip()
    if not raw:
        return ''
    if raw in USERS:
        return raw
    lower = raw.lower()
    for key, u in USERS.items():
        email = (u.get('email') or '').strip().lower()
        if email and email == lower:
            return key
        if key.lower() == lower:
            return key
        if (u.get('display') or '').strip().lower() == lower:
            return key
    return raw


def _login_redirect_with_error(error, selected_user='', next_url=''):
    """Post/Redirect/Get so a refresh never re-submits the password form."""
    session['login_flash_error'] = error or ''
    session['login_flash_user'] = selected_user or ''
    if next_url:
        session['login_next'] = next_url
    resp = redirect(url_for('login'))
    # Phones cache POST/GET aggressively — never cache auth screens.
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, private'
    resp.headers['Pragma'] = 'no-cache'
    return resp


def _no_store_html(template_name, **ctx):
    """Render HTML that mobile browsers must not keep as a stale error screen."""
    resp = make_response(render_template(template_name, **ctx))
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, private'
    resp.headers['Pragma'] = 'no-cache'
    return resp


@app.route('/login', methods=['GET', 'POST'])
def login():
    if session.get('user_key'):
        # Already signed in — never bounce through /login as a "next" target
        return _post_login_redirect()

    next_url = safe_next_url(request.args.get('next', ''))
    if next_url:
        session['login_next'] = next_url

    # GET and HEAD: show form (and any one-shot flash from a failed POST).
    # HEAD must not fall through to POST logic — that 302-looped some mobile/Safari
    # prefetches and "too many redirects" errors after a wrong password.
    if request.method in ('GET', 'HEAD'):
        error = session.pop('login_flash_error', None) or None
        success = session.pop('login_flash_success', None) or None
        selected_user = session.pop('login_flash_user', '') or ''
        return _no_store_html(
            'login.html',
            users=sorted(USERS.items(), key=lambda x: x[1]['display']),
            error=error,
            success=success,
            selected_user=selected_user,
            next_url=next_url or session.get('login_next') or '',
        )

    # POST
    next_from_form = safe_next_url(request.form.get('next', ''))
    if next_from_form:
        session['login_next'] = next_from_form
        next_url = next_from_form
    # Do NOT strip password — spaces can be intentional; strip only identity.
    user_key = _resolve_login_user_key(request.form.get('user_key', ''))
    password = request.form.get('password', '') or ''
    selected_user = user_key if user_key in USERS else ''
    ip = client_ip(request)

    if not user_key or user_key not in USERS:
        return _login_redirect_with_error(
            'Select your name from the list, then enter your password.',
            '',
            next_url,
        )

    if not password:
        return _login_redirect_with_error(
            'Enter your password.',
            selected_user,
            next_url,
        )

    # Always verify credentials first. A correct password must work even during
    # a temporary lockout (lockout only blocks guessing, not real sign-in).
    try:
        conn = get_db()
        if not conn:
            return _login_redirect_with_error(
                'Sign-in is temporarily unavailable (database). '
                'Please try again in a moment — this is not a wrong password.',
                selected_user,
                next_url,
            )

        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('SELECT * FROM hub_users WHERE user_key = %s', (user_key,))
        db_user = cur.fetchone()
        cur.close()
        conn.close()

        if db_user and _safe_check_password(db_user.get('password_hash'), password):
            user = USERS.get(user_key, {})
            _establish_session(user_key, user, db_user)
            record_login_attempt(get_db, user_key, True, ip)
            clear_login_failures(get_db, user_key)
            _update_last_login(user_key)
            if session.get('must_change_password'):
                return redirect(url_for('change_password'))
            return _post_login_redirect()

        # Wrong password / inactive account — apply lockout only to real wrong passwords
        if not db_user:
            # Do NOT count toward lockout: nothing to guess; activation is the fix.
            return _login_redirect_with_error(
                'Your Hub account is not activated yet. '
                'Use Forgot Password, or ask Thomas or Stephanie to set your password from Admin.',
                selected_user,
                next_url,
            )

        locked, fail_count, mins_left = is_login_locked(get_db, user_key)
        if locked:
            wait = mins_left if mins_left is not None else LOGIN_LOCKOUT_MINUTES
            # Password already checked above and was wrong. Correct password
            # always succeeds before this branch (lockout is for guesses only).
            return _login_redirect_with_error(
                f'That password is not correct, and this account is temporarily locked '
                f'after {fail_count} failed tries (~{wait} min left). '
                f'Use Forgot Password to set a new one (that unlocks you), '
                f'or ask Thomas/Stephanie to Unlock or Reset from Admin.',
                selected_user,
                next_url,
            )

        record_login_attempt(get_db, user_key, False, ip)
        locked_now, fails, mins = is_login_locked(get_db, user_key)
        remaining = max(0, MAX_LOGIN_FAILURES - fails)
        if locked_now:
            wait = mins if mins is not None else LOGIN_LOCKOUT_MINUTES
            return _login_redirect_with_error(
                f'Incorrect password. Account locked for about {wait} minute'
                f'{"s" if wait != 1 else ""}. Use Forgot Password (unlocks you) '
                f'or ask Admin to Unlock.',
                selected_user,
                next_url,
            )
        return _login_redirect_with_error(
            f'Incorrect password. {remaining} attempt'
            f'{"s" if remaining != 1 else ""} left before a temporary lock. '
            f'Use Forgot Password if you are unsure.',
            selected_user,
            next_url,
        )
    except Exception as e:
        print(f"Login error for {user_key or 'unknown'}: {e}")
        return _login_redirect_with_error(
            'Something went wrong on our side. Please try again or contact Thomas.',
            selected_user,
            next_url,
        )

def _touch_last_active(user_key, force=False):
    """Record user activity. force=True on explicit login; otherwise throttle to 30 min."""
    if not user_key:
        return
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            if force:
                cur.execute(
                    'UPDATE hub_users SET last_login = NOW() WHERE user_key = %s',
                    (user_key,)
                )
            else:
                cur.execute(
                    '''UPDATE hub_users SET last_login = NOW()
                       WHERE user_key = %s
                         AND (last_login IS NULL OR last_login < NOW() - INTERVAL '30 minutes')''',
                    (user_key,)
                )
            conn.commit()
            cur.close()
            conn.close()
    except:
        pass


def _update_last_login(user_key):
    _touch_last_active(user_key, force=True)


def _build_dashboard_recent_feed(
    recent_proposals,
    recent_ppms,
    recent_tpscopes,
    recent_siding,
    recent_roofing,
    recent_gutter,
    recent_painting,
    limit=8,
):
    """Merge recent hub activity into one chronological feed for the dashboard."""
    from datetime import datetime

    items = []

    def _add(dt, payload):
        items.append({**payload, '_ts': dt or datetime.min})

    for i, p in enumerate(recent_proposals or []):
        _add(p.get('generated_at'), {
            'kind': 'proposal',
            'kind_label': 'Proposal',
            'title': p.get('property_name') or p.get('client_name') or 'Unnamed',
            'meta': ' · '.join(x for x in [p.get('consultant_name'), p.get('property_type')] if x),
            'modal': {'type': 'proposal', 'id': p.get('id')},
        })
    for i, p in enumerate(recent_ppms or []):
        _add(p.get('generated_at'), {
            'kind': 'ppm',
            'kind_label': 'PPM',
            'title': p.get('client_name') or p.get('property_name') or 'Unnamed',
            'meta': p.get('proj_type') or p.get('pm_name') or '',
            'modal': {'type': 'ppm', 'id': p.get('id')},
        })
    for i, s in enumerate(recent_tpscopes or []):
        _add(s.get('generated_at'), {
            'kind': 'tps',
            'kind_label': 'TPS',
            'title': s.get('property_name') or 'Unnamed',
            'meta': ' · '.join(x for x in [
                s.get('language', '').title() if s.get('language') else '',
                s.get('pm_name'),
            ] if x),
            'modal': {'type': 'tps', 'id': s.get('id')},
        })
    for e in recent_siding or []:
        _add(e.get('generated_at'), {
            'kind': 'estimate',
            'kind_label': 'Siding',
            'title': e.get('property_name') or 'Unnamed',
            'meta': f"{e.get('building_count') or 1} building{'s' if (e.get('building_count') or 1) != 1 else ''}",
            'url': f"/siding-estimator/result/{e.get('id')}",
        })
    for e in recent_roofing or []:
        _add(e.get('generated_at'), {
            'kind': 'estimate',
            'kind_label': 'Roofing',
            'title': e.get('property_name') or 'Unnamed',
            'meta': e.get('report_type') or 'report',
            'url': f"/roofing-estimator/result/{e.get('id')}",
        })
    for e in recent_gutter or []:
        _add(e.get('generated_at'), {
            'kind': 'estimate',
            'kind_label': 'Gutters',
            'title': e.get('property_name') or 'Unnamed',
            'meta': f"{float(e.get('gutter_lf') or 0):.0f} LF",
            'url': f"/gutter-estimator/result/{e.get('id')}",
        })
    for e in recent_painting or []:
        _add(e.get('generated_at'), {
            'kind': 'estimate',
            'kind_label': 'Painting',
            'title': e.get('property_name') or 'Unnamed',
            'meta': f"{e.get('line_count') or 0} lines",
            'url': f"/painting-estimator/result/{e.get('id')}",
        })

    items.sort(key=lambda x: x['_ts'], reverse=True)
    feed = []
    for it in items[:limit]:
        dt = it['_ts']
        feed.append({
            'kind': it['kind'],
            'kind_label': it['kind_label'],
            'title': it['title'],
            'meta': it['meta'],
            'date': dt.strftime('%b %d') if hasattr(dt, 'strftime') and dt != datetime.min else '',
            'url': it.get('url'),
            'modal': it.get('modal'),
        })
    return feed


@app.route('/dashboard')
@require_login
def dashboard():
    user_key = session['user_key']
    _touch_last_active(user_key)
    user = USERS.get(user_key, {})
    proposal_access = get_user_proposal_access(user_key)
    accessible_consultants = {k: CONSULTANTS[k] for k in proposal_access if k in CONSULTANTS}
    recent_proposals = get_recent_proposals(user_key)
    recent_ppms = get_recent_ppms(user_key)
    recent_siding_estimates = []
    recent_roofing_estimates = []
    recent_gutter_estimates = []
    recent_painting_estimates = []
    # Recent Trade Partner Scopes (generated by user or listed as PM)
    recent_tpscopes = get_recent_tpscopes(user_key)
    try:
        conn_tps = get_db()
        if conn_tps:
            cur_tps = conn_tps.cursor(cursor_factory=RealDictCursor)
            cur_tps.execute(
                '''SELECT id, property_name, property_address, building_count, siding_type, generated_at
                   FROM siding_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_siding_estimates = cur_tps.fetchall()
            cur_tps.execute(
                '''SELECT id, property_name, property_address, report_type, generated_at
                   FROM roofing_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_roofing_estimates = cur_tps.fetchall()
            cur_tps.execute(
                '''SELECT id, property_name, property_address, gutter_lf, generated_at
                   FROM gutter_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_gutter_estimates = cur_tps.fetchall()
            cur_tps.execute(
                '''SELECT id, property_name, property_address, line_count, one_coat_bid, generated_at
                   FROM painting_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_painting_estimates = cur_tps.fetchall()
            cur_tps.close()
            conn_tps.close()
    except Exception as e:
        print(f"Recent activity error: {e}")
    real_is_admin = (user.get('role') == 'admin')
    real_role = user.get('role', '')

    # Role preview (?view_as=sales|pm) removed 2026-08-22 at Thomas's request.
    # It let the owner render the dashboard as a "standard PSC" or "standard PM"
    # to check what those people saw. The tier rework made that question mostly
    # moot — everyone below owner now sees the same dashboard — and the preview
    # had become actively misleading: it faked a role while the real session
    # permissions still applied underneath, so it showed a layout no actual
    # person would get. `/admin/member/<user_key>` answers "what does this
    # person have" honestly; do not rebuild a fake-identity preview.
    is_admin = real_is_admin
    user_role = real_role
    team_view = True
    psc_training_enrolled = is_psc_training_enrolled(user_key)
    psc_training_stats = (
        compute_psc_training_stats(user_key) if psc_training_enrolled else None
    )
    psc_training_oversight = can_psc_training_oversight(user_key)
    pm_training_oversight = can_pm_training_oversight(user_key)
    unread_feedback = 0
    unread_diffs = 0
    pricing_summary = None
    if is_admin:
        unread_feedback, unread_diffs = _admin_inbox_counts()
        pricing_summary = _pricing_summary_for_dashboard()

    # Sales and Production both open by default for field roles (and admin)
    sales_lane_open = user_role in ('consultant', 'pm', 'office_manager', 'admin')
    production_lane_open = user_role in ('consultant', 'pm', 'office_manager', 'admin')

    # Pipeline Board: every board is open to everyone since the tier rework;
    # this list decides which one opens by default and what the lane shows.
    pipeline_boards = pipeline_board.list_accessible_boards(USERS, user_key)
    pipeline_board_pair_key = pipeline_board.get_pair_key(USERS, user_key)
    pipeline_board_access = bool(pipeline_boards)
    # Office Ops: leadership tier and up (Thomas, Stephanie, Tony, Trey).
    office_ops_access = office_ops.can_access_office_ops(USERS, user_key)

    date_events = get_date_events(user_key, is_admin=real_is_admin)
    recent_feed = _build_dashboard_recent_feed(
        recent_proposals,
        recent_ppms,
        recent_tpscopes,
        recent_siding_estimates,
        recent_roofing_estimates,
        recent_gutter_estimates,
        recent_painting_estimates,
    )
    # PM training: open to everyone (under construction)
    pm_training_stats = compute_pm_training_stats(user_key)
    pm_training_open = True
    # Field Ask PPS only on dashboard — same queue rules for everyone (no curator admin UI).
    ask_pps_prompt = ask_pps.get_next_prompt_for_user(get_db, USERS, user_key, user_role)
    ask_pps_prompt_queue = len(
        ask_pps.get_prompts_for_user(
            get_db, USERS, user_key, user_role, include_all_for_curator=False,
        )
    )
    user_notifications = ask_pps.get_unread_notifications(get_db, user_key)

    # ── Above the lanes (2026-08-25) ────────────────────────────────────────
    #
    # The dashboard was 4.1 phone screens with the first tool link 258px down
    # and nothing above it but a greeting. These two blocks put something
    # worth reading in that space. Both are assembled by dashboard_summary
    # from data this route already holds — see that module for why the week
    # number is the recap's number and why nothing renders at zero.
    #
    # Cost is two extra queries per load (the pipeline count and the usage
    # read); the recap score is cached across the whole company for five
    # minutes, so it is amortised to roughly nothing.
    summary_pills = []
    dashboard_recent_tools = []
    try:
        week_score = dashboard_summary.week_scores(get_db, USERS).get(user_key)
        summary_pills = dashboard_summary.build_pills(
            week_score=week_score,
            pipeline_open=dashboard_summary.pipeline_in_progress(
                get_db, pipeline_board_pair_key, pipeline_board.COMPLETED_STATUSES,
            ),
            pipeline_url=(
                f'/pipeline-board?pair={pipeline_board_pair_key}'
                if pipeline_board_pair_key else None
            ),
            psc_pct=(psc_training_stats or {}).get('pct'),
            pm_pct=(pm_training_stats or {}).get('pct'),
            unread_feedback=unread_feedback,
            unread_diffs=unread_diffs,
            is_owner=is_admin,
            show_week=(
                dashboard_summary.SHOW_WEEK_SCORE_TO_EVERYONE or is_admin
            ),
        )
        # History is not permission: this set is what the person may open
        # today, not what they once used. A tier change must not leave a card
        # pointing into a tool they have lost.
        allowed_tools = {'ppm', 'estimate', 'site_visit'}
        if accessible_consultants:
            allowed_tools |= {'proposal', 'tps'}
        if pipeline_boards:
            allowed_tools.add('pipeline')
        if office_ops_access:
            allowed_tools |= {'office_ops', 'compliance'}
        if psc_training_enrolled:
            allowed_tools.add('psc_training')
        if pm_training_open:
            allowed_tools.add('pm_training')
        dashboard_recent_tools = dashboard_summary.recent_tools(
            {
                'proposal': recent_proposals,
                'ppm': recent_ppms,
                'tps': recent_tpscopes,
                'estimate': (
                    list(recent_siding_estimates)
                    + list(recent_roofing_estimates)
                    + list(recent_gutter_estimates)
                    + list(recent_painting_estimates)
                ),
            },
            usage_rows=dashboard_summary.recent_usage_features(get_db, user_key),
            proposal_url=os.environ.get(
                'PROPOSAL_URL', 'https://pps-proposal-tool.onrender.com'),
            allowed=allowed_tools,
        )
    except Exception as e:
        # The strip is a convenience on top of a page that worked without it
        # for a year. It must never be the reason a dashboard 500s.
        print(f'Dashboard summary error: {e}')

    return render_template(
        'dashboard.html',
        summary_pills=summary_pills,
        dashboard_recent_tools=dashboard_recent_tools,
        user=user,
        user_key=user_key,
        user_role=user_role,
        real_role=real_role,
        real_is_admin=real_is_admin,
        ask_pps_prompt=ask_pps_prompt,
        ask_pps_prompt_queue=ask_pps_prompt_queue,
        user_notifications=user_notifications,
        sales_lane_open=sales_lane_open,
        production_lane_open=production_lane_open,
        admin_lane_open=is_admin,
        team_view=team_view,
        consultants=accessible_consultants,
        recent_proposals=recent_proposals,
        recent_ppms=recent_ppms,
        recent_tpscopes=recent_tpscopes,
        recent_feed=recent_feed,
        date_events=date_events,
        psc_training_stats=psc_training_stats,
        psc_training_enrolled=psc_training_enrolled,
        psc_training_oversight=psc_training_oversight,
        pm_training_stats=pm_training_stats,
        pm_training_oversight=pm_training_oversight,
        pm_training_open=pm_training_open,
        pipeline_board_access=pipeline_board_access,
        pipeline_board_pair_key=pipeline_board_pair_key,
        pipeline_boards=pipeline_boards,
        office_ops_access=office_ops_access,
        unread_feedback=unread_feedback,
        unread_diffs=unread_diffs,
        pricing_summary=pricing_summary,
        proposal_url=os.environ.get('PROPOSAL_URL', 'https://pps-proposal-tool.onrender.com'),
    )


@app.route('/estimating/property-lookup')
@require_login
def estimating_property_lookup():
    """Search hub clients and past jobs by address — not external measurement APIs."""
    q = (request.args.get('q') or '').strip()
    try:
        from estimators.property_lookup import lookup_property_by_address
        return jsonify(lookup_property_by_address(get_db, q))
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e, results=[])


@app.route('/estimating/confidence', methods=['POST'])
@require_login
def estimating_confidence():
    """Return takeoff reliability metadata for gutter, roofing, or siding."""
    data = request.get_json(silent=True) or {}
    tool = (data.get('tool') or '').strip()
    try:
        from estimators.reliability import (
            build_gutter_reliability,
            build_painting_reliability,
            build_roofing_reliability,
            build_siding_job_reliability,
            build_siding_reliability,
        )
        if tool == 'gutter':
            confidence = build_gutter_reliability(
                data.get('measurements') or {},
                data.get('inputs') or {},
                data.get('user_overrides') or {},
            )
        elif tool == 'roofing':
            confidence = build_roofing_reliability(
                data.get('measurements') or {},
                data.get('inputs') or {},
            )
        elif tool == 'siding':
            buildings = data.get('buildings') or []
            if buildings:
                confidence = build_siding_job_reliability(
                    buildings,
                    int(data.get('pricing_loaded') or 0),
                )
            else:
                confidence = build_siding_reliability(
                    data.get('measurements') or {},
                    data.get('source') or 'eagleview',
                    int(data.get('pricing_loaded') or 0),
                )
        elif tool == 'painting':
            confidence = build_painting_reliability(
                data.get('measurements') or {},
                data.get('line_items') or [],
                data.get('user_overrides') or {},
            )
        else:
            return jsonify({'error': 'Unknown tool'}), 400
        return jsonify({'confidence': confidence})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/estimating')
@require_login
def estimating_hub():
    user_key = session['user_key']
    recent_siding = []
    recent_roofing = []
    recent_gutters = []
    recent_painting = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute(
                '''SELECT id, property_name, building_count, generated_at
                   FROM siding_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_siding = cur.fetchall()
            cur.execute(
                '''SELECT id, property_name, report_type, generated_at
                   FROM roofing_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_roofing = cur.fetchall()
            cur.execute(
                '''SELECT id, property_name, gutter_lf, generated_at
                   FROM gutter_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_gutters = cur.fetchall()
            cur.execute(
                '''SELECT id, property_name, line_count, one_coat_bid, two_coat_bid, generated_at
                   FROM painting_estimate_log WHERE generated_by = %s
                   ORDER BY generated_at DESC LIMIT 5''',
                (user_key,),
            )
            recent_painting = cur.fetchall()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Estimating hub error: {e}")
    return render_template(
        'estimating.html',
        recent_siding=recent_siding,
        recent_roofing=recent_roofing,
        recent_gutters=recent_gutters,
        recent_painting=recent_painting,
    )


@app.route('/log-proposal', methods=['POST'])
def log_proposal():
    """Called by proposal tool after successful generation."""
    data = request.get_json()
    api_key = request.headers.get('X-API-Key', '')
    if api_key != INTERNAL_API_KEY:
        return jsonify({'error': 'Unauthorized'}), 401
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute('''
                INSERT INTO proposal_log
                (generated_by, consultant_key, consultant_name, client_name,
                 property_name, property_address, property_type, template_type,
                 proposal_number, existing_issue, intended_outcome, scopes_selected, scope_notes)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''', (
                data.get('generated_by'),
                data.get('consultant_key'),
                data.get('consultant_name'),
                data.get('client_name'),
                data.get('property_name'),
                data.get('property_address', ''),
                data.get('property_type'),
                data.get('template_type'),
                data.get('proposal_number', ''),
                data.get('existing_issue', ''),
                data.get('intended_outcome', ''),
                data.get('scopes_selected', ''),
                data.get('scope_notes', ''),
            ))
            conn.commit()
            cur.close()
            conn.close()
            sync_proposal_number_sequence(
                data.get('consultant_key') or '',
                data.get('proposal_number') or '',
            )
            _touch_last_active(data.get('generated_by'))
        return jsonify({'success': True})
    except Exception as e:
        print(f"Log proposal error: {e}")
        return _api_error(e)


def _get_proposal_log_for_document(cur, document_id):
    cur.execute('''
        SELECT pl.*, d.filename, d.mime_type, d.size_bytes, d.user_key AS doc_user_key
        FROM documents d
        JOIN proposal_log pl ON pl.id = d.log_id
        WHERE d.id = %s AND d.doc_type = 'proposal'
    ''', (document_id,))
    return cur.fetchone()


def _user_can_download_proposal(user_key, role, row):
    if not row:
        return False
    if role == 'admin':
        return True
    if row.get('generated_by') == user_key:
        return True
    if row.get('doc_user_key') == user_key:
        return True
    # PMs / shared access: download proposals for consultants they are paired to
    if user_can_access_consultant_proposals(user_key, row.get('consultant_key')):
        return True
    return False


def _user_can_access_proposal_log(user_key, role, row):
    """View/prefill history: own generations, or any proposal for accessible consultants."""
    if not row:
        return False
    if role == 'admin':
        return True
    if row.get('generated_by') == user_key:
        return True
    return user_can_access_consultant_proposals(user_key, row.get('consultant_key'))


@app.route('/api/proposals/next-number')
def proposal_next_number():
    """Suggest the next proposal number for a consultant (INITIALS + YY + XXX); does not reserve."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401
    consultant_key = _normalize_consultant_key(request.args.get('consultant_key', ''))
    if not consultant_key:
        return jsonify({'error': 'consultant_key required'}), 400
    number, err = peek_next_proposal_number(consultant_key)
    if err:
        return jsonify({'error': err}), 400
    return jsonify({
        'success': True,
        'proposal_number': number,
        'consultant_key': consultant_key,
        'format': 'INITIALS + 2-digit year + 3-digit sequence (e.g. TE26001)',
    })


@app.route('/api/proposals/<int:log_id>/prefill')
def proposal_prefill(log_id):
    """Return saved proposal metadata for regenerating in the proposal tool."""
    user_key = session.get('user_key')
    role = session.get('role', '')
    if _internal_api_ok():
        user_key = user_key or request.args.get('user_key', '').strip()
        if user_key:
            role = USERS.get(user_key, {}).get('role', role)
    if not user_key:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('SELECT * FROM proposal_log WHERE id = %s', (log_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return jsonify({'error': 'Proposal not found'}), 404
        if not _user_can_access_proposal_log(user_key, role, row):
            return jsonify({'error': 'Permission denied'}), 403

        scopes = []
        if row.get('scopes_selected'):
            scopes = [s.strip() for s in row['scopes_selected'].split(',') if s.strip()]

        scope_details = []
        if row.get('scope_details'):
            try:
                parsed = json.loads(row['scope_details'])
                if isinstance(parsed, list):
                    scope_details = parsed
            except Exception:
                scope_details = [s.strip() for s in row['scope_details'].split(',') if s.strip()]

        pricing_lines = []
        if row.get('pricing_json'):
            try:
                parsed = json.loads(row['pricing_json'])
                if isinstance(parsed, list):
                    pricing_lines = parsed
            except Exception:
                pass

        return jsonify({
            'success': True,
            'consultant_key': row.get('consultant_key') or '',
            'client_name': row.get('client_name') or row.get('property_name') or '',
            'contact_name': row.get('contact_name') or '',
            'contact_email': row.get('contact_email') or '',
            'company': row.get('company') or '',
            'address': row.get('property_address') or '',
            'property_type': row.get('property_type') or '',
            'template_type': row.get('template_type') or 'short',
            'scope_style': row.get('scope_style') or 'bullets',
            'proposal_number': row.get('proposal_number') or '',
            'proposal_date': row.get('proposal_date') or '',
            'expiry_date': row.get('expiry_date') or '',
            'existing_issue': row.get('existing_issue') or '',
            'intended_outcome': row.get('intended_outcome') or '',
            'scopes': scopes,
            'scope_details': scope_details,
            'other_scope': row.get('other_scope') or '',
            'scope_notes': row.get('scope_notes') or '',
            'pricing_lines': pricing_lines,
            'warranty_pps': row.get('warranty_pps') or '',
            'warranty_mfg': row.get('warranty_mfg') or '',
            'contract_total': row.get('contract_total') or '',
            'has_file': bool(row.get('document_id')),
            'document_id': row.get('document_id'),
        })
    except Exception as e:
        print(f"Proposal prefill error: {e}")
        return _api_error(e)


@app.route('/api/vault/proposals', methods=['POST'])
def vault_store_proposal():
    """Store a generated proposal file and metadata atomically."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401

    data = request.get_json() or {}
    user_key = (data.get('user_key') or '').strip()
    if not user_key:
        return jsonify({'error': 'user_key required'}), 400

    filename = (data.get('filename') or 'PPS_Proposal.docx').strip()
    file_b64 = data.get('file_base64') or ''
    if not file_b64:
        return jsonify({'error': 'file_base64 required'}), 400

    try:
        file_bytes = base64.b64decode(file_b64)
    except Exception:
        return jsonify({'error': 'Invalid file_base64 payload'}), 400

    if len(file_bytes) > MAX_DOCUMENT_BYTES:
        return jsonify({'error': f'File exceeds {MAX_DOCUMENT_BYTES // (1024*1024)} MB limit'}), 413
    if len(file_bytes) < 100:
        return jsonify({'error': 'File too small or empty'}), 400

    mime_type = data.get('mime_type') or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO proposal_log
            (generated_by, consultant_key, consultant_name, client_name,
             property_name, property_address, property_type, template_type, scope_style,
             proposal_number, existing_issue, intended_outcome, scopes_selected, scope_notes,
             contact_name, contact_email, company, scope_details, other_scope,
             pricing_json, warranty_pps, warranty_mfg, proposal_date, expiry_date, contract_total)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        ''', (
            data.get('generated_by') or user_key,
            data.get('consultant_key'),
            data.get('consultant_name'),
            data.get('client_name'),
            data.get('property_name'),
            data.get('property_address', ''),
            data.get('property_type'),
            data.get('template_type'),
            data.get('scope_style', 'bullets'),
            data.get('proposal_number', ''),
            data.get('existing_issue', ''),
            data.get('intended_outcome', ''),
            data.get('scopes_selected', ''),
            data.get('scope_notes', ''),
            data.get('contact_name', ''),
            data.get('contact_email', ''),
            data.get('company', ''),
            data.get('scope_details', ''),
            data.get('other_scope', ''),
            data.get('pricing_json', ''),
            data.get('warranty_pps', ''),
            data.get('warranty_mfg', ''),
            data.get('proposal_date', ''),
            data.get('expiry_date', ''),
            data.get('contract_total', ''),
        ))
        log_id = cur.fetchone()[0]

        cur.execute('''
            INSERT INTO documents
            (doc_type, log_id, user_key, filename, mime_type, size_bytes, file_data)
            VALUES ('proposal', %s, %s, %s, %s, %s, %s)
            RETURNING id
        ''', (log_id, user_key, filename, mime_type, len(file_bytes), psycopg2.Binary(file_bytes)))
        document_id = cur.fetchone()[0]

        cur.execute('UPDATE proposal_log SET document_id = %s WHERE id = %s', (document_id, log_id))
        conn.commit()
        cur.close()
        conn.close()
        sync_proposal_number_sequence(
            data.get('consultant_key') or '',
            data.get('proposal_number') or '',
        )
        _touch_last_active(data.get('generated_by') or user_key)
        return jsonify({'success': True, 'log_id': log_id, 'document_id': document_id})
    except Exception as e:
        print(f"Vault store error: {e}")
        return _api_error(e)


@app.route('/api/documents/<int:document_id>/download')
def document_download(document_id):
    """Download a vaulted document — hub session or internal API proxy."""
    user_key = session.get('user_key')
    role = session.get('role', '')
    internal = _internal_api_ok()
    if internal:
        user_key = request.args.get('user_key', user_key)
        role = USERS.get(user_key, {}).get('role', role)

    if not user_key:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor(cursor_factory=RealDictCursor)
        row = _get_proposal_log_for_document(cur, document_id)
        if not row:
            cur.close()
            conn.close()
            return jsonify({'error': 'Document not found'}), 404
        if not _user_can_download_proposal(user_key, role, row):
            cur.close()
            conn.close()
            return jsonify({'error': 'Permission denied'}), 403

        cur.execute(
            'SELECT filename, mime_type, file_data FROM documents WHERE id = %s',
            (document_id,)
        )
        doc = cur.fetchone()
        cur.close()
        conn.close()
        if not doc or not doc.get('file_data'):
            return jsonify({'error': 'File data missing'}), 404

        return send_file(
            BytesIO(bytes(doc['file_data'])),
            as_attachment=True,
            download_name=doc['filename'],
            mimetype=doc['mime_type'],
        )
    except Exception as e:
        print(f"Document download error: {e}")
        return _api_error(e)


@app.route('/log-ppm', methods=['POST'])
def log_ppm():
    """Called by PPM tool after generation."""
    data = request.get_json()
    api_key = request.headers.get('X-API-Key', '')
    if api_key != INTERNAL_API_KEY:
        return jsonify({'error': 'Unauthorized'}), 401
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute(
                '''INSERT INTO ppm_log
                   (generated_by, property_name, pm_key, pm_name,
                    property_address, client_name, proposal_number, total_value,
                    proposal_date, proj_type, scale, client_type, occupied)
                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)''',
                (
                    data.get('generated_by'),
                    data.get('property_name'),
                    data.get('pm_key', ''),
                    data.get('pm_name', ''),
                    data.get('property_address', ''),
                    data.get('client_name', ''),
                    data.get('proposal_number', ''),
                    data.get('total_value', ''),
                    data.get('proposal_date', ''),
                    data.get('proj_type', ''),
                    data.get('scale', ''),
                    data.get('client_type', ''),
                    data.get('occupied', ''),
                )
            )
            conn.commit()
            cur.close()
            conn.close()
            _touch_last_active(data.get('generated_by'))
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


@app.route('/session-info')
def session_info():
    """Called by proposal/profile tools to get current user session."""
    user_key = session.get('user_key')
    if not user_key:
        return jsonify({'authenticated': False})
    user = USERS.get(user_key, {})
    return jsonify({
        'authenticated': True,
        'user_key': user_key,
        'display_name': session.get('display_name'),
        'role': session.get('role'),
        'proposal_access': get_user_proposal_access(user_key),
    })


def _fetch_admin_breakdown(cur):
    """Trade/template activity breakdown for admin dashboard."""
    breakdown = {}
    queries = {
        'property_types': '''
            SELECT COALESCE(NULLIF(TRIM(property_type), ''), 'Unknown') AS label, COUNT(*) AS cnt
            FROM proposal_log GROUP BY 1 ORDER BY cnt DESC LIMIT 14
        ''',
        'templates': '''
            SELECT COALESCE(NULLIF(TRIM(template_type), ''), 'Unknown') AS label, COUNT(*) AS cnt
            FROM proposal_log GROUP BY 1 ORDER BY cnt DESC LIMIT 14
        ''',
        'ppm_types': '''
            SELECT COALESCE(NULLIF(TRIM(proj_type), ''), 'Unknown') AS label, COUNT(*) AS cnt
            FROM ppm_log GROUP BY 1 ORDER BY cnt DESC LIMIT 14
        ''',
        'tps_languages': '''
            SELECT COALESCE(NULLIF(TRIM(language), ''), 'Unknown') AS label, COUNT(*) AS cnt
            FROM subscope_log GROUP BY 1 ORDER BY cnt DESC LIMIT 14
        ''',
    }
    for key, sql in queries.items():
        try:
            cur.execute(sql)
            breakdown[key] = cur.fetchall()
        except Exception:
            breakdown[key] = []
    return breakdown


def _fetch_vault_summary(cur):
    """Vault storage stats and file list (no binary payload)."""
    vault = {
        'files': [],
        'file_count': 0,
        'total_bytes': 0,
        'proposals_with_file': 0,
        'proposals_total': 0,
    }
    try:
        cur.execute('SELECT COUNT(*) AS c, COALESCE(SUM(size_bytes), 0) AS b FROM documents')
        row = cur.fetchone()
        vault['file_count'] = row['c'] or 0
        vault['total_bytes'] = row['b'] or 0
        cur.execute('SELECT COUNT(*) AS c FROM proposal_log')
        vault['proposals_total'] = cur.fetchone()['c'] or 0
        cur.execute('SELECT COUNT(*) AS c FROM proposal_log WHERE document_id IS NOT NULL')
        vault['proposals_with_file'] = cur.fetchone()['c'] or 0
        vault['proposals_missing_file'] = max(0, vault['proposals_total'] - vault['proposals_with_file'])
        cur.execute(
            """SELECT COUNT(*) AS c FROM proposal_log
               WHERE document_id IS NULL AND generated_at >= NOW() - INTERVAL '7 days'"""
        )
        vault['recent_missing_file'] = cur.fetchone()['c'] or 0
        cur.execute('''
            SELECT d.id, d.doc_type, d.filename, d.size_bytes, d.created_at, d.user_key, d.log_id,
                   pl.property_name, pl.consultant_name, pl.generated_by
            FROM documents d
            LEFT JOIN proposal_log pl ON pl.id = d.log_id AND d.doc_type = 'proposal'
            ORDER BY d.created_at DESC
            LIMIT 100
        ''')
        vault['files'] = cur.fetchall()

        limit = VAULT_STORAGE_LIMIT_BYTES
        used = vault['total_bytes'] or 0
        used_pct = min(100, round(used / limit * 100)) if limit else 0
        vault['storage_limit_bytes'] = limit
        vault['storage_used_bytes'] = used
        vault['storage_used_pct'] = used_pct
        vault['storage_remaining_pct'] = max(0, 100 - used_pct)
        vault['storage_used_mb'] = round(used / (1024 * 1024), 1)
        vault['storage_limit_mb'] = int(limit / (1024 * 1024))
        vault['storage_remaining_mb'] = round(max(0, limit - used) / (1024 * 1024), 1)
    except Exception as e:
        print(f"Vault summary error: {e}")
    return vault


def _serialize_dt(val):
    if not val:
        return ''
    return val.strftime('%Y-%m-%d') if hasattr(val, 'strftime') else str(val)


def _format_activity_date(val):
    if not val:
        return '—'
    if hasattr(val, 'strftime'):
        return val.strftime('%B %d, %Y')
    return str(val)


def _format_activity_by(user_key):
    if not user_key:
        return '—'
    return str(user_key).replace('_', ' ').title()


def _format_template_label(val):
    if not val:
        return '—'
    v = str(val).strip().lower()
    if v in ('short', 'standard'):
        return 'Standard'
    if v in ('full', 'long', 'comprehensive'):
        return 'Comprehensive'
    return str(val)


@app.template_filter('template_label')
def _template_label_filter(val):
    return _format_template_label(val)


def _can_view_activity_row(user_key, role, activity_type, row):
    """Any signed-in person may open any activity row (2026-08-21 tier rework).

    Deliberate: the same proposal/PPM/TPS activity is now on Team View for
    everyone and in the weekly team recap, so gating the detail drawer while
    publishing the summary would just be inconsistent. Signature kept so
    callers don't change; `role` and `activity_type` are unused now.
    """
    return bool(user_key) and user_key in USERS


def _activity_detail_row(label, value):
    val = value if value not in (None, '') else '—'
    return {'label': label, 'value': val}


def _format_contact(row, client=''):
    """Contact name and email on one line — 'Dana Reed · dana@acme.com'.

    Two rows for one person padded the modal without telling you anything the
    joined line doesn't. Returns '' when neither is set, which is common on
    older rows: `/log-proposal` (the original logging path) never wrote these
    columns, only `/api/vault/proposals` does. Callers drop the row entirely
    rather than printing a dash, so an old proposal looks the same as it always
    did instead of gaining an empty field.

    When the contact IS the client — the common case for a single-owner
    property — the name is dropped from this line and only the email remains,
    so the modal doesn't print "Dana Reed" twice in a row.
    """
    name = (row.get('contact_name') or '').strip()
    email = (row.get('contact_email') or '').strip()
    if name and client and name.casefold() == client.strip().casefold():
        name = ''
    return ' · '.join(p for p in (name, email) if p)


def _activity_detail_payload(activity_type, row):
    """Normalize a DB row into a JSON payload for activity detail modals."""
    if activity_type == 'proposal':
        # Client / Contact / Company sit directly under the proposal number:
        # "who is this for" before "what was in it". The data has been arriving
        # from the proposal tool all along and was simply never displayed.
        client = (row.get('client_name') or '').strip()
        company = (row.get('company') or '').strip()
        contact = _format_contact(row, client)
        detail_rows = [_activity_detail_row('Proposal #', row.get('proposal_number'))]
        if client:
            detail_rows.append(_activity_detail_row('Client', client))
        if contact:
            detail_rows.append(_activity_detail_row('Contact', contact))
        # Company is usually the client name again; only worth a row when it
        # actually says something different.
        if company and company.casefold() != client.casefold():
            detail_rows.append(_activity_detail_row('Company', company))
        return {
            'type': 'proposal',
            'title': row.get('property_name') or row.get('client_name') or 'Unnamed',
            'rows': detail_rows + [
                _activity_detail_row('Address', row.get('property_address')),
                _activity_detail_row('Consultant', row.get('consultant_name')),
                _activity_detail_row('Generated By', _format_activity_by(row.get('generated_by'))),
                _activity_detail_row('Property Type', row.get('property_type')),
                _activity_detail_row('Template', _format_template_label(row.get('template_type'))),
                _activity_detail_row('Scopes', row.get('scopes_selected')),
                _activity_detail_row('Date', _format_activity_date(row.get('generated_at'))),
            ],
            'issue': row.get('existing_issue') or '',
            'outcome': row.get('intended_outcome') or '',
            'notes': row.get('scope_notes') or '',
            'document_id': row.get('document_id'),
            'log_id': row.get('id'),
        }
    if activity_type == 'ppm':
        details = ' · '.join(x for x in [
            row.get('scale'),
            row.get('client_type'),
            row.get('occupied'),
            row.get('total_value'),
        ] if x and x != '—') or '—'
        prop_date = row.get('proposal_date')
        generated = _format_activity_date(row.get('generated_at'))
        date_val = f"{prop_date} (generated {generated})" if prop_date and prop_date != '—' else generated
        return {
            'type': 'ppm',
            'title': row.get('property_name') or row.get('client_name') or 'Unnamed',
            'rows': [
                _activity_detail_row('Proposal #', row.get('proposal_number')),
                _activity_detail_row('Address', row.get('property_address')),
                _activity_detail_row('Client', row.get('client_name')),
                _activity_detail_row('Project Manager', row.get('pm_name')),
                _activity_detail_row('Project Type', row.get('proj_type')),
                _activity_detail_row('Details', details),
                _activity_detail_row('Generated By', _format_activity_by(row.get('generated_by'))),
                _activity_detail_row('Date', date_val),
            ],
        }
    if activity_type == 'tps':
        lang = (row.get('language') or '').title() or '—'
        mat = row.get('material_provider')
        mat_label = 'TP materials' if mat == 'tp' else 'PPS materials' if mat == 'pps' else (mat or '—')
        return {
            'type': 'tps',
            'title': row.get('property_name') or 'Unnamed',
            'rows': [
                _activity_detail_row('PO #', row.get('po_number')),
                _activity_detail_row('Address', row.get('property_address')),
                _activity_detail_row('Consultant', row.get('consultant_name')),
                _activity_detail_row('Project Manager', row.get('pm_name')),
                _activity_detail_row('Details', ' · '.join(x for x in [lang, mat_label] if x and x != '—') or '—'),
                _activity_detail_row('Generated By', _format_activity_by(row.get('generated_by'))),
                _activity_detail_row('Date', _format_activity_date(row.get('generated_at'))),
            ],
            'notes': (
                f"Source proposal: {row.get('proposal_filename')}"
                if row.get('proposal_filename') and row.get('proposal_filename') != '—'
                else ''
            ),
        }
    if activity_type == 'site_visit':
        rows_out = [
            _activity_detail_row('Address', row.get('property_address')),
            _activity_detail_row('Visit Date', row.get('visit_date')),
            _activity_detail_row('Time', row.get('visit_time')),
            _activity_detail_row('Visited By', row.get('display_name')),
            _activity_detail_row('PO / Proposal #', row.get('po_number')),
        ]
        if row.get('trade_partner_company'):
            rows_out.append(_activity_detail_row('Trade Partner', row.get('trade_partner_company')))
        if row.get('crew_lead'):
            rows_out.append(_activity_detail_row('Crew Lead', row.get('crew_lead')))
        if row.get('staff_contact'):
            rows_out.append(_activity_detail_row('Onsite Contact', row.get('staff_contact')))
        if row.get('topics_discussed'):
            rows_out.append(_activity_detail_row('Topics', row.get('topics_discussed')))
        if row.get('complaints_received') == 'yes' and row.get('complaint_details'):
            rows_out.append(_activity_detail_row('Complaint', row.get('complaint_details')))
        rows_out.append(_activity_detail_row('Submitted', _format_activity_date(row.get('generated_at'))))
        checklist = row.get('checklist')
        if isinstance(checklist, str):
            try:
                checklist = json.loads(checklist)
            except Exception:
                checklist = []
        return {
            'type': 'site_visit',
            'title': row.get('property_name') or 'Unnamed',
            'rows': rows_out,
            'observations': row.get('observations') or '',
            'checklist': checklist or [],
            'download_url': f"/site-visit/download/{row.get('id')}",
        }
    return None


ACTIVITY_DETAIL_TABLES = {
    'proposal': 'proposal_log',
    'ppm': 'ppm_log',
    'tps': 'subscope_log',
    'site_visit': 'site_visit_log',
}


@app.route('/api/activity-detail/<activity_type>/<int:record_id>')
@require_login
def activity_detail(activity_type, record_id):
    table = ACTIVITY_DETAIL_TABLES.get(activity_type)
    if not table:
        return jsonify({'error': 'Unknown activity type'}), 400
    user_key = session.get('user_key')
    role = session.get('role')
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(f'SELECT * FROM {table} WHERE id = %s', (record_id,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return jsonify({'error': 'Not found'}), 404
        if not _can_view_activity_row(user_key, role, activity_type, row):
            return jsonify({'error': 'Forbidden'}), 403
        payload = _activity_detail_payload(activity_type, row)
        if not payload:
            return jsonify({'error': 'Unsupported activity type'}), 400
        return jsonify(payload)
    except Exception as e:
        print(f"Activity detail error: {e}")
        return _api_error(e)


def _admin_search(cur, query, limit=50):
    """Search proposals, PPMs, TPS, site visits, and clients."""
    q = (query or '').strip()
    if len(q) < 2:
        return []
    pattern = f'%{q}%'
    per_type = max(6, limit // 5)
    results = []

    def _add(result_type, row, title, meta, date_val, link='', document_id=None):
        results.append({
            'type': result_type,
            'id': row.get('id'),
            'title': title or 'Unnamed',
            'meta': meta,
            'date': _serialize_dt(date_val),
            'link': link,
            'document_id': document_id,
        })

    try:
        cur.execute('''
            SELECT id, property_name, client_name, property_address, consultant_name,
                   proposal_number, generated_by, generated_at, document_id
            FROM proposal_log
            WHERE property_name ILIKE %s OR client_name ILIKE %s OR property_address ILIKE %s
               OR proposal_number ILIKE %s OR consultant_name ILIKE %s OR scopes_selected ILIKE %s
               OR COALESCE(generated_by, '') ILIKE %s
            ORDER BY generated_at DESC LIMIT %s
        ''', (pattern,) * 7 + (per_type,))
        for r in cur.fetchall():
            _add(
                'proposal',
                r,
                r.get('property_name') or r.get('client_name'),
                ' · '.join(filter(None, [
                    r.get('consultant_name'),
                    r.get('proposal_number'),
                    (r.get('generated_by') or '').replace('_', ' ').title(),
                ])),
                r.get('generated_at'),
                '/admin/proposals',
                r.get('document_id'),
            )
    except Exception as e:
        print(f"Search proposals error: {e}")

    try:
        cur.execute('''
            SELECT id, property_name, property_address, client_name, proposal_number,
                   pm_name, proj_type, generated_by, generated_at
            FROM ppm_log
            WHERE property_name ILIKE %s OR property_address ILIKE %s OR client_name ILIKE %s
               OR proposal_number ILIKE %s OR pm_name ILIKE %s OR proj_type ILIKE %s
               OR COALESCE(generated_by, '') ILIKE %s
            ORDER BY generated_at DESC LIMIT %s
        ''', (pattern,) * 7 + (per_type,))
        for r in cur.fetchall():
            _add(
                'ppm',
                r,
                r.get('property_name'),
                ' · '.join(filter(None, [r.get('pm_name'), r.get('proj_type'), (r.get('generated_by') or '').replace('_', ' ').title()])),
                r.get('generated_at'),
            )
    except Exception as e:
        print(f"Search PPM error: {e}")

    try:
        cur.execute('''
            SELECT id, property_name, property_address, consultant_name, pm_name,
                   po_number, language, generated_by, generated_at
            FROM subscope_log
            WHERE property_name ILIKE %s OR property_address ILIKE %s OR consultant_name ILIKE %s
               OR pm_name ILIKE %s OR po_number ILIKE %s OR language ILIKE %s
               OR COALESCE(generated_by, '') ILIKE %s
            ORDER BY generated_at DESC LIMIT %s
        ''', (pattern,) * 7 + (per_type,))
        for r in cur.fetchall():
            _add(
                'tps',
                r,
                r.get('property_name'),
                ' · '.join(filter(None, [
                    r.get('consultant_name'),
                    r.get('language', '').title() if r.get('language') else '',
                    f"PO {r['po_number']}" if r.get('po_number') else '',
                ])),
                r.get('generated_at'),
                '/admin/tpscopes',
            )
    except Exception as e:
        print(f"Search TPS error: {e}")

    try:
        cur.execute('''
            SELECT id, property_name, property_address, display_name, po_number, generated_by, generated_at
            FROM site_visit_log
            WHERE property_name ILIKE %s OR property_address ILIKE %s OR display_name ILIKE %s
               OR po_number ILIKE %s OR COALESCE(generated_by, '') ILIKE %s
               OR COALESCE(observations, '') ILIKE %s
            ORDER BY generated_at DESC LIMIT %s
        ''', (pattern,) * 6 + (per_type,))
        for r in cur.fetchall():
            _add(
                'site_visit',
                r,
                r.get('property_name'),
                ' · '.join(filter(None, [r.get('display_name'), r.get('po_number')])),
                r.get('generated_at'),
                '/admin/site-visits',
            )
    except Exception as e:
        print(f"Search site visit error: {e}")

    try:
        cur.execute('''
            SELECT id, name, company, property_name, address, email, added_by, updated_at
            FROM clients
            WHERE name ILIKE %s OR company ILIKE %s OR property_name ILIKE %s
               OR address ILIKE %s OR email ILIKE %s OR COALESCE(added_by, '') ILIKE %s
            ORDER BY updated_at DESC LIMIT %s
        ''', (pattern,) * 6 + (per_type,))
        for r in cur.fetchall():
            _add(
                'client',
                r,
                r.get('name') or r.get('property_name'),
                ' · '.join(filter(None, [r.get('company'), r.get('email'), (r.get('added_by') or '').replace('_', ' ').title()])),
                r.get('updated_at'),
                '/clients',
            )
    except Exception as e:
        print(f"Search clients error: {e}")

    return results[:limit]


@app.route('/admin')
@require_admin
def admin():
    rows = []
    all_proposals = []
    all_ppms = []
    all_subscopes = []
    unread_feedback, _ = _admin_inbox_counts()
    client_count = 0
    proposals_30d = ppms_30d = subscopes_30d = 0
    vault = {
        'files': [], 'file_count': 0, 'total_bytes': 0,
        'proposals_with_file': 0, 'proposals_total': 0,
        'proposals_missing_file': 0, 'recent_missing_file': 0,
        'storage_limit_bytes': VAULT_STORAGE_LIMIT_BYTES,
        'storage_used_bytes': 0, 'storage_used_pct': 0,
        'storage_remaining_pct': 100, 'storage_used_mb': 0,
        'storage_limit_mb': int(VAULT_STORAGE_LIMIT_BYTES / (1024 * 1024)),
        'storage_remaining_mb': int(VAULT_STORAGE_LIMIT_BYTES / (1024 * 1024)),
    }
    proposals_all = ppms_all = subscopes_all = 0
    ask_pps_pending_cnt = 0
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM hub_users ORDER BY display_name')
            rows = cur.fetchall()
            cur.execute('SELECT * FROM proposal_log ORDER BY generated_at DESC LIMIT 50')
            all_proposals = cur.fetchall()
            cur.execute('SELECT * FROM ppm_log ORDER BY generated_at DESC LIMIT 50')
            all_ppms = cur.fetchall()
            cur.execute('SELECT * FROM subscope_log ORDER BY generated_at DESC LIMIT 50')
            all_subscopes = cur.fetchall()
            cur.execute("SELECT COUNT(*) AS c FROM proposal_log WHERE generated_at >= NOW() - INTERVAL '30 days'")
            proposals_30d = cur.fetchone()['c'] or 0
            cur.execute("SELECT COUNT(*) AS c FROM ppm_log WHERE generated_at >= NOW() - INTERVAL '30 days'")
            ppms_30d = cur.fetchone()['c'] or 0
            cur.execute("SELECT COUNT(*) AS c FROM subscope_log WHERE generated_at >= NOW() - INTERVAL '30 days'")
            subscopes_30d = cur.fetchone()['c'] or 0
            # The all-time half of each tile was `all_proposals|length` — the
            # length of a `LIMIT 50` result set. It read "30d / all" and stopped
            # counting at 50, so the day the Hub passed fifty proposals the tile
            # started under-reporting and never said so.
            cur.execute('SELECT COUNT(*) AS c FROM proposal_log')
            proposals_all = cur.fetchone()['c'] or 0
            cur.execute('SELECT COUNT(*) AS c FROM ppm_log')
            ppms_all = cur.fetchone()['c'] or 0
            cur.execute('SELECT COUNT(*) AS c FROM subscope_log')
            subscopes_all = cur.fetchone()['c'] or 0
            vault = _fetch_vault_summary(cur)
            try:
                cur.execute('SELECT COUNT(*) as cnt FROM clients')
                client_count = cur.fetchone()['cnt']
            except Exception:
                pass
            try:
                cur.execute(
                    "SELECT COUNT(*) AS c FROM knowledge_entries WHERE status = 'pending'"
                )
                ask_pps_pending_cnt = cur.fetchone()['c'] or 0
            except Exception:
                ask_pps_pending_cnt = 0
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Admin error: {e}")

    # Annotate lockouts so admin can unlock stuck teammates at a glance.
    #
    # One query for the whole roster. This was `is_login_locked(get_db, key)`
    # inside the loop, and again in the loop below — each call opens its own
    # Postgres connection, so a thirteen-person roster opened fourteen and paid
    # fourteen TLS handshakes before the table could render.
    locks = login_lock_map(get_db)
    annotated = []
    for row in rows:
        item = dict(row)
        locked, fails, mins = locks.get(item.get('user_key'), (False, 0, None))
        item['login_locked'] = locked
        item['login_failures'] = fails
        item['login_mins_left'] = mins
        item['missing_hub_row'] = False
        annotated.append(item)
    rows = annotated

    # Ensure every USERS profile appears even if hub_users row is missing (can't log in).
    present = {r.get('user_key') for r in rows}
    for key, udef in USERS.items():
        if key in present:
            continue
        locked, fails, mins = locks.get(key, (False, 0, None))
        rows.append({
            'user_key': key,
            'display_name': udef.get('display', key),
            'role': udef.get('role', ''),
            'last_login': None,
            'must_change_password': False,
            'missing_hub_row': True,
            'login_locked': locked,
            'login_failures': fails,
            'login_mins_left': mins,
        })

    # Tier is what actually governs access since the 2026-08-21 rework, and it
    # was displayed nowhere. Read it here rather than in the template so the
    # template does not need to know tiers.py exists.
    for row in rows:
        row['tier'] = user_tier(row.get('user_key'))

    rows = sorted(rows, key=lambda r: (r.get('display_name') or '').lower())

    recent_activity = merge_activity(all_proposals, all_ppms, all_subscopes, limit=30)

    return render_template('admin.html', users=rows,
                           recent_activity=recent_activity,
                           unread_feedback=unread_feedback,
                           ask_pps_pending_cnt=ask_pps_pending_cnt,
                           client_count=client_count,
                           proposals_30d=proposals_30d, ppms_30d=ppms_30d,
                           subscopes_30d=subscopes_30d,
                           proposals_all=proposals_all, ppms_all=ppms_all,
                           subscopes_all=subscopes_all,
                           vault=vault,
                           user_definitions=USERS)


@app.route('/admin/breakdown')
@require_admin
def admin_breakdown():
    """Trade and template mix. Lifted off /admin 2026-08-22.

    Four bar charts of all-time proportions is reference material, not something
    that changes between two page loads — it was occupying the middle of the
    page Thomas opens to check on people and unlock accounts. Same query, same
    charts, its own address.
    """
    breakdown = {}
    conn = None
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            breakdown = _fetch_admin_breakdown(cur)
            cur.close()
    except Exception as e:
        print(f'Admin breakdown error: {e}')
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass
    return render_template('admin_breakdown.html', breakdown=breakdown)


@app.route('/admin/training')
@require_login
def admin_training():
    """Curriculum editor. Leadership — Tony, Trey, Stephanie and Thomas.

    Shows the curriculum with DRAFTS applied, so whoever is editing sees what
    they are building rather than what is live. `enrolled_at=None` means "as
    authored today", which is the right view here: the editor is composing the
    programme a new hire will get, not previewing any particular trainee's copy.
    """
    user_key = session['user_key']
    module = (request.args.get('module') or 'psc').strip().lower()
    if module not in training_overlay.MODULES:
        module = 'psc'
    if not (can_psc_training_oversight(user_key) if module == 'psc'
            else can_pm_training_oversight(user_key)):
        return redirect(url_for('dashboard'))

    drafts = training_overlay.load_overlay(get_db, include_drafts=True).get(module)
    if module == 'psc':
        curriculum = get_training_curriculum()
    else:
        _meta, pm_weeks = get_pm_training_curriculum()
        curriculum = (dict(), pm_weeks, {}, {}, {})
    curriculum, unplaceable = training_overlay.apply(curriculum, drafts, enrolled_at=None)
    onboarding, weeks, core_values, sales_training, company_operations = curriculum

    published = set()
    for row in (drafts or {}).get('items', []):
        if row.get('published_at'):
            published.add(row['item_id'])

    return render_template(
        'admin_training.html',
        module=module,
        user=USERS.get(user_key, {}),
        onboarding=onboarding if module == 'psc' else None,
        weeks=weeks,
        sections=([
            {'key': 'core_values', 'label': 'Core Values',
             'groups': (core_values or {}).get('sections', []), 'container': 'activities'},
            {'key': 'sales_training', 'label': 'Sales Training',
             'groups': (sales_training or {}).get('modules', []), 'container': 'items'},
            {'key': 'company_operations', 'label': 'Company Operations',
             'groups': (company_operations or {}).get('modules', []), 'container': 'items'},
        ] if module == 'psc' else []),
        week_containers=training_overlay.WEEK_CONTAINERS,
        overlay_ids=published | {r['item_id'] for r in (drafts or {}).get('items', [])},
        draft_ids={r['item_id'] for r in (drafts or {}).get('items', [])
                   if not r.get('published_at')},
        edited_ids=set((drafts or {}).get('edits', {})),
        unplaceable=unplaceable,
        pending=training_overlay.pending_counts(get_db, module),
        cache_seconds=training_overlay.CACHE_TTL_SECONDS,
    )


def _training_editor_guard(module):
    """Both modules gate on their own oversight check, not a shared one.

    Trey runs PM training and Tony runs PSC; today both are leadership so the
    two resolve identically, but they are separate questions and writing them
    as one is how a future narrowing gets missed on one side.
    """
    user_key = session.get('user_key')
    if module not in training_overlay.MODULES:
        return jsonify({'error': 'Unknown module'}), 400
    ok = (can_psc_training_oversight(user_key) if module == 'psc'
          else can_pm_training_oversight(user_key))
    if not ok:
        return jsonify({'error': 'Not authorized'}), 403
    return None


@app.route('/api/training/item', methods=['POST'])
@require_login
def api_training_add_item():
    data = request.get_json(silent=True) or {}
    module = (data.get('module') or '').strip().lower()
    blocked = _training_editor_guard(module)
    if blocked:
        return blocked
    item_id = training_overlay.create_item(
        get_db, module, data.get('target') or {}, data.get('payload') or {},
        session['user_key'], sort_order=int(data.get('sort_order') or 0),
    )
    if not item_id:
        return jsonify({'error': 'Give the item a title before saving.'}), 400
    return jsonify({'success': True, 'item_id': item_id,
                    'pending': training_overlay.pending_counts(get_db, module)})


@app.route('/api/training/edit', methods=['POST'])
@require_login
def api_training_edit_item():
    data = request.get_json(silent=True) or {}
    module = (data.get('module') or '').strip().lower()
    blocked = _training_editor_guard(module)
    if blocked:
        return blocked
    hidden = data.get('hidden')
    ok = training_overlay.save_edit(
        get_db, module, (data.get('item_id') or '').strip(),
        fields=data.get('fields') or {},
        hidden=None if hidden is None else bool(hidden),
        user_key=session['user_key'],
    )
    if not ok:
        return jsonify({'error': 'Could not save that change.'}), 500
    return jsonify({'success': True,
                    'pending': training_overlay.pending_counts(get_db, module)})


@app.route('/api/training/discard', methods=['POST'])
@require_login
def api_training_discard_item():
    """Only ever removes an addition that has never published — see
    `discard_draft_item`. Anything live is hidden, not deleted."""
    data = request.get_json(silent=True) or {}
    module = (data.get('module') or '').strip().lower()
    blocked = _training_editor_guard(module)
    if blocked:
        return blocked
    ok = training_overlay.discard_draft_item(
        get_db, module, (data.get('item_id') or '').strip())
    if not ok:
        return jsonify({'error': 'That item has already published — hide it instead.'}), 400
    return jsonify({'success': True,
                    'pending': training_overlay.pending_counts(get_db, module)})


@app.route('/api/training/publish', methods=['POST'])
@require_login
def api_training_publish():
    data = request.get_json(silent=True) or {}
    module = (data.get('module') or '').strip().lower()
    blocked = _training_editor_guard(module)
    if blocked:
        return blocked
    result = training_overlay.publish(get_db, module, session['user_key'])
    if not result.get('ok'):
        return jsonify({'error': 'Could not publish.'}), 500
    return jsonify({'success': True,
                    'items': result.get('items', 0), 'edits': result.get('edits', 0),
                    'pending': training_overlay.pending_counts(get_db, module),
                    'cache_seconds': training_overlay.CACHE_TTL_SECONDS})


@app.route('/admin/pricing-defaults', methods=['GET', 'POST'])
@require_admin
def admin_pricing_defaults():
    from estimators.pricing_defaults import save_pricing_defaults, SYSTEM_DEFAULTS

    defaults = _pricing_defaults()
    message = None
    error = None
    if request.method == 'POST':
        try:
            trades = {
                'siding': {
                    'labor_per_sq': request.form.get('siding_labor_per_sq'),
                    'haul_per_sq': request.form.get('siding_haul_per_sq'),
                    'tax_pct': request.form.get('siding_tax_pct'),
                    'delivery': request.form.get('siding_delivery'),
                    'waste_pct': request.form.get('siding_waste_pct'),
                },
                'roofing': {
                    'labor_per_sq': request.form.get('roofing_labor_per_sq'),
                    'material_per_sq': request.form.get('roofing_material_per_sq'),
                    'tax_pct': request.form.get('roofing_tax_pct'),
                    'margin_pct': request.form.get('roofing_margin_pct'),
                    'waste_pct': request.form.get('roofing_waste_pct'),
                    'dump_divisor': request.form.get('roofing_dump_divisor'),
                    'dump_cost': request.form.get('roofing_dump_cost'),
                },
                'gutter': {
                    'gutter_price_per_lf': request.form.get('gutter_price_per_lf'),
                    'guard_price_per_lf': request.form.get('gutter_guard_per_lf'),
                    'labor_per_lf': request.form.get('gutter_labor_per_lf'),
                    'tax_pct': request.form.get('gutter_tax_pct'),
                    'margin_pct': request.form.get('gutter_margin_pct'),
                    'waste_pct': request.form.get('gutter_waste_pct'),
                    'downspout_lf_each': request.form.get('gutter_ds_height'),
                    'downspout_spacing_ft': request.form.get('gutter_ds_spacing'),
                },
                'painting': {
                    'labor_per_hour': request.form.get('painting_labor_per_hour'),
                    'margin_one_coat_pct': request.form.get('painting_margin_one_coat_pct'),
                    'margin_two_coat_pct': request.form.get('painting_margin_two_coat_pct'),
                    'two_coat_multiplier': request.form.get('painting_two_coat_multiplier'),
                },
            }
            defaults = save_pricing_defaults(
                get_db,
                trades,
                session['user_key'],
                session.get('display_name', ''),
            )
            message = 'Pricing defaults saved. New estimates will use these values.'
        except Exception as e:
            error = str(e)

    return render_template(
        'admin_pricing_defaults.html',
        defaults=defaults,
        system_defaults=SYSTEM_DEFAULTS,
        message=message,
        error=error,
    )


@app.route('/admin/system-state')
@require_admin
def admin_system_state():
    """One page answering what is running, where everyone stands, and whether
    the scheduled jobs fired.

    Built because on 2026-08-21 none of those were answerable from inside the
    product — which commit was live took a Render log, who still had not set a
    password took a group text, and whether Monday's recap would fire took
    someone remembering the cron service did not exist yet. /health stays thin
    and public; this is owner-only and specific enough to act on.
    """
    people = system_state.people_rows(get_db, USERS)
    conn = get_db()
    db_ok = bool(conn)
    if conn:
        try:
            conn.close()
        except Exception:
            pass
    return render_template(
        'admin_system_state.html',
        # Per-worker, so with 2 gunicorn workers a refresh can show either
        # one's numbers. That is a property of the pool, not a bug in the
        # panel — a pool shared between processes would need Redis.
        db_pool_stats=db_pool.stats(),
        service=system_state.service_rows(),
        people=people,
        summary=system_state.summarize(people),
        jobs=system_state.job_rows(get_db),
        db_ok=db_ok,
        # weekly_recap is imported at module scope; daily_digest is not (it is
        # imported inside the functions that use it), so reaching for it here
        # would NameError at request time rather than at import.
        now_label=weekly_recap.eastern_now().strftime('%b %-d, %-I:%M %p ET'),
    )


@app.route('/admin/system-health')
@require_admin
def admin_system_health():
    """Load system health async — avoids blocking /admin on outbound HTTP (single-worker deadlock)."""
    try:
        return jsonify(_run_system_health_checks())
    except Exception as e:
        print(f"System health error: {e}")
        return jsonify({
            'ok': False,
            'checks': [{'name': 'health_check', 'ok': False, 'error': str(e)}],
        })


@app.route('/admin/daily-digest-test', methods=['POST'])
@require_admin
def admin_daily_digest_test():
    """Send yesterday's digest now (admin smoke test)."""
    from daily_digest import run_daily_digest

    try:
        result = run_daily_digest(
            get_db,
            USERS,
            _format_template_label,
            _send_digest_email,
            force=True,
        )
        status = 200 if result.get('ok') else 500
        return jsonify(result), status
    except Exception as e:
        print(f'Admin digest test error: {e}')
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/admin/weekly-recap-test', methods=['POST'])
@require_admin
def admin_weekly_recap_test():
    """Send Thomas his own copy of last week's recap, on demand.

    Same code path as the cron, addressed only to him — so the format can be
    checked before it goes company-wide on a Monday. Mirrors
    /admin/daily-digest-test.
    """
    try:
        start, end = weekly_recap.last_week_bounds()
        roll_start, roll_end = weekly_recap.rolling_bounds()
        scores = weekly_recap.collect_scores(get_db, USERS, start, end)
        # The rolling pass was missing here, so build_groups fell back to
        # rolling=None and every 12-week figure rendered as 0 — lower than the
        # week beside it, which is impossible since the week is inside the
        # rolling window. The cron path always passed it; only this preview
        # did not, which is precisely the path used to check the format.
        rolling = weekly_recap.collect_scores(get_db, USERS, roll_start, roll_end)
        groups = weekly_recap.build_groups(
            USERS, scores, weekly_recap._excluded_keys(), rolling,
        )
        subject, text_body, html_body = weekly_recap.build_recap_email(
            groups, start, session.get('user_key'), USERS,
        )
        me = USERS.get(session.get('user_key'), {}).get('email', '')
        if not me:
            return jsonify({'ok': False, 'error': 'no email on your account'}), 400
        ok = _send_digest_email(f'[TEST] {subject}', text_body, html_body, [me])
        return jsonify({'ok': bool(ok), 'sent_to': me, 'week': weekly_recap.week_label(start)})
    except Exception as e:
        print(f'Weekly recap test error: {e}')
        return _api_error(e, ok=False)


@app.route('/admin/search')
@require_admin
def admin_search():
    q = request.args.get('q', '').strip()
    results = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            results = _admin_search(cur, q)
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Admin search error: {e}")
        return _api_error(e)
    return jsonify({'query': q, 'results': results, 'count': len(results)})


@app.route('/admin/vault/delete', methods=['POST'])
@require_admin
def admin_vault_delete():
    data = request.get_json(silent=True) or {}
    doc_id = data.get('document_id') or request.form.get('document_id')
    try:
        doc_id = int(doc_id)
    except (TypeError, ValueError):
        return jsonify({'error': 'Invalid document_id'}), 400
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        cur.execute('SELECT id, filename FROM documents WHERE id = %s', (doc_id,))
        row = cur.fetchone()
        if not row:
            cur.close()
            conn.close()
            return jsonify({'error': 'Document not found'}), 404
        cur.execute('UPDATE proposal_log SET document_id = NULL WHERE document_id = %s', (doc_id,))
        cur.execute('DELETE FROM documents WHERE id = %s', (doc_id,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True, 'deleted_id': doc_id})
    except Exception as e:
        print(f"Vault delete error: {e}")
        return _api_error(e)


@app.route('/admin/reset-password', methods=['POST'])
@require_admin
def admin_reset_password():
    user_key = (request.form.get('user_key') or '').strip()
    new_password = request.form.get('new_password', '')  # do not strip — match login
    if not user_key:
        return jsonify({'error': 'No team member selected. Close and open Reset again.'}), 400
    if not new_password or len(new_password) < 6:
        return jsonify({'error': 'Password must be at least 6 characters.'}), 400
    if user_key not in USERS:
        return jsonify({'error': f'Unknown user “{user_key}”.'}), 400
    try:
        ok, action = _upsert_hub_user_password(user_key, new_password, must_change=False)
        if not ok:
            reason = action or 'unknown'
            friendly = {
                'no_db': (
                    'Cannot reach the database right now. '
                    'Check Admin → System Health (database_connect), then try again in a minute.'
                ),
                'password_too_short': 'Password must be at least 6 characters.',
                'unknown_user': 'Unknown user.',
                'hash_failed': 'Could not encrypt password. Try a different password.',
            }.get(reason)
            if not friendly and reason.startswith('db_error:'):
                friendly = (
                    f'Database rejected the password save ({reason.split(":", 1)[-1]}). '
                    'Try again; if it keeps failing, check System Health.'
                )
            return jsonify({
                'error': friendly or 'Could not save password.',
                'reason': reason,
            }), 503 if reason in ('no_db',) or str(reason).startswith('db_error:') else 400
        # Unlock is best-effort — password is already saved
        unlocked = True
        try:
            clear_login_failures(get_db, user_key)
        except Exception as unlock_err:
            print(f'clear_login_failures after reset failed for {user_key}: {unlock_err}')
            unlocked = False
        return jsonify({
            'success': True,
            'action': action,
            'unlocked': unlocked,
            'user_key': user_key,
            'display': USERS[user_key].get('display'),
        })
    except Exception as e:
        _log_exception(e, 'admin_reset_password')
        return jsonify({
            'error': f'Password reset failed ({type(e).__name__}). Check System Health / database.',
        }), 500


@app.route('/admin/unlock-login', methods=['POST'])
@require_admin
def admin_unlock_login():
    """Clear failed login attempts so a teammate can sign in immediately."""
    user_key = (request.form.get('user_key') or '').strip()
    if not user_key or user_key not in USERS:
        return jsonify({'error': 'Unknown user'}), 400
    try:
        clear_login_failures(get_db, user_key)
        return jsonify({'success': True, 'user_key': user_key})
    except Exception as e:
        return _api_error(e)


@app.route('/api/internal/unlock-login', methods=['POST'])
def internal_unlock_login():
    """Ops helper: clear lockout without an admin browser session."""
    if not _internal_api_ok():
        return jsonify({'error': 'Unauthorized'}), 401
    data = request.get_json(silent=True) or {}
    user_key = (request.form.get('user_key') or data.get('user_key') or '').strip()
    if not user_key or user_key not in USERS:
        return jsonify({'error': 'Unknown user'}), 400
    try:
        clear_login_failures(get_db, user_key)
        return jsonify({
            'success': True,
            'user_key': user_key,
            'display': USERS[user_key].get('display'),
        })
    except Exception as e:
        return _api_error(e)


@app.route('/log-subscope', methods=['POST'])
def log_subscope():
    data = request.get_json()
    api_key = request.headers.get('X-API-Key', '')
    if api_key != INTERNAL_API_KEY:
        return jsonify({'error': 'Unauthorized'}), 401
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute(
                '''INSERT INTO subscope_log
                   (generated_by, property_name, pm_name, consultant_name, language,
                    property_address, po_number, consultant_key, pm_key,
                    material_provider, proposal_filename)
                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)''',
                (
                    data.get('generated_by'),
                    data.get('property_name'),
                    data.get('pm_name'),
                    data.get('consultant_name'),
                    data.get('language'),
                    data.get('property_address', ''),
                    data.get('po_number', ''),
                    data.get('consultant_key', ''),
                    data.get('pm_key', ''),
                    data.get('material_provider', ''),
                    data.get('proposal_filename', ''),
                )
            )
            conn.commit()
            cur.close()
            conn.close()
            _touch_last_active(data.get('generated_by'))
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


@app.route('/feedback', methods=['POST'])
def submit_feedback():
    if not session.get('user_key'):
        return jsonify({'error': 'Not authenticated'}), 401
    message = request.form.get('message', '').strip()
    if not message:
        return jsonify({'error': 'Message is required'}), 400
    user_key = session['user_key']
    display_name = session.get('display_name', '')
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute(
                'INSERT INTO feedback (user_key, display_name, message) VALUES (%s, %s, %s)',
                (user_key, display_name, message)
            )
            conn.commit()
            cur.close()
            conn.close()
        # Send email
        _send_feedback_email(display_name, message)
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


def _hub_notify_recipients():
    """Primary inbox(es) for hub feedback and proposal comparison submissions."""
    raw = os.environ.get('HUB_NOTIFY_EMAIL', 'thomas@purepropsolutions.com')
    return [e.strip() for e in raw.split(',') if e.strip()]


def _send_smtp_email(subject, text_body, html_body=None, recipients=None):
    """Send email via SMTP to the given recipient list."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    if not recipients:
        print(f"SMTP email (no recipients): {subject}\n{text_body}")
        return False

    smtp_host = os.environ.get('SMTP_HOST', '')
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')
    if not smtp_host:
        print(f"SMTP email:\nSubject: {subject}\nTo: {', '.join(recipients)}\n{text_body}")
        return False

    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = smtp_user
        msg['To'] = ', '.join(recipients)
        msg.attach(MIMEText(text_body, 'plain'))
        if html_body:
            msg.attach(MIMEText(html_body, 'html'))
        with smtplib.SMTP_SSL(smtp_host, 465, timeout=30) as s:
            s.login(smtp_user, smtp_pass)
            s.send_message(msg)
        return True
    except Exception as e:
        print(f"SMTP email failed: {e}")
        return False


def _send_digest_email(subject, text_body, html_body, recipients):
    """Nightly digest — SMTP first, Resend fallback (password reset uses Resend)."""
    if not recipients:
        return False
    if _send_smtp_email(subject, text_body, html_body, recipients):
        print(f'Daily digest sent via SMTP to {", ".join(recipients)}')
        return True
    if not os.environ.get('RESEND_API_KEY', '').strip():
        print('Daily digest: SMTP failed and Resend is not configured')
        return False
    print(f'Daily digest: SMTP failed, trying Resend for {", ".join(recipients)}')
    ok_all = True
    for addr in recipients:
        ok, detail = _send_resend_email(addr, subject, html_body or '', text_body)
        if ok:
            print(f'Daily digest sent via Resend to {addr} ({detail})')
        else:
            print(f'Daily digest Resend failed for {addr}: {detail}')
            ok_all = False
    return ok_all


def _send_hub_notify_email(subject, text_body, html_body=None):
    """Email hub admin when users submit feedback or voice comparisons."""
    return _send_smtp_email(subject, text_body, html_body, _hub_notify_recipients())


def _send_feedback_email(name, message):
    from html import escape

    admin_url = f"{HUB_PUBLIC_URL.rstrip('/')}/admin/feedback"
    subject = f'PPS Hub Feedback — {name}'
    text_body = f"Feedback from {name}:\n\n{message.strip()}\n\nReview in hub: {admin_url}"
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:0 auto;">
      <div style="background:#004C8C;padding:18px 22px;border-radius:8px 8px 0 0;">
        <p style="color:white;font-size:17px;font-weight:600;margin:0;">PPS Hub Feedback</p>
      </div>
      <div style="background:#f8fafc;padding:22px;border:1px solid #e2e8f0;border-top:none;border-radius:0 0 8px 8px;">
        <p style="color:#334155;font-size:14px;margin:0 0 16px;"><strong>From:</strong> {escape(name)}</p>
        <div style="background:white;border:1px solid #e2e8f0;border-radius:8px;padding:14px;color:#334155;font-size:14px;line-height:1.55;white-space:pre-wrap;">{escape(message.strip())}</div>
        <p style="margin:16px 0 0;">
          <a href="{admin_url}" style="color:#004C8C;font-weight:600;">Open feedback in Admin →</a>
        </p>
      </div>
    </div>
    """
    _send_hub_notify_email(subject, text_body, html_body)


def _save_proposal_diff(user_key, display_name, diff_analysis, voice_recommendations,
                        user_notes='', comparison_prompt=''):
    """Persist a proposal comparison submission; returns new row id or None."""
    try:
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor()
        cur.execute(
            '''INSERT INTO proposal_diffs
               (user_key, display_name, property_name, diff_analysis, user_notes,
                voice_recommendations, comparison_prompt)
               VALUES (%s, %s, %s, %s, %s, %s, %s)
               RETURNING id''',
            (user_key, display_name, '', diff_analysis, user_notes,
             voice_recommendations, comparison_prompt),
        )
        diff_id = cur.fetchone()[0]
        conn.commit()
        cur.close()
        conn.close()
        return diff_id
    except Exception as e:
        print(f'Save proposal diff error: {e}')
        return None


def _notify_proposal_diff_email(name, user_notes, diff_analysis, voice_recommendations,
                                comparison_prompt=''):
    """Send comparison notification without blocking the HTTP response."""
    import threading

    def _send():
        try:
            _send_proposal_diff_email(
                name, user_notes, diff_analysis, voice_recommendations,
                comparison_prompt=comparison_prompt,
            )
        except Exception as e:
            _log_exception(e, 'proposal-diff-email')

    threading.Thread(target=_send, daemon=True).start()


def _send_proposal_diff_email(name, user_notes, diff_analysis, voice_recommendations,
                              comparison_prompt=''):
    from html import escape

    admin_url = f"{HUB_PUBLIC_URL.rstrip('/')}/admin/diffs"
    subject = f'Proposal Comparison — {name}'
    analysis_text = (diff_analysis or '').strip()
    voice_text = (voice_recommendations or '').strip()
    prompt_text = (comparison_prompt or '').strip()
    prompt_block = f"\n\nComparison prompt:\n{prompt_text}" if prompt_text else ''
    notes_block = f"\n\nConsultant notes:\n{user_notes.strip()}" if user_notes else ''
    text_body = (
        f"New proposal comparison from {name}\n"
        f"{prompt_block}{notes_block}\n\n"
        f"What Changed:\n{analysis_text}\n\n"
        f"Voice Guide Recommendations:\n{voice_text}\n\n"
        f"Review in hub: {admin_url}"
    )
    prompt_html = ''
    if prompt_text:
        prompt_html = (
            '<div style="margin-bottom:14px;">'
            '<p style="font-size:12px;font-weight:600;color:#004C8C;text-transform:uppercase;margin:0 0 6px;">Comparison Prompt</p>'
            f'<div style="background:#FFF9E6;border:1px solid #FFE082;border-radius:8px;padding:14px;color:#334155;font-size:14px;line-height:1.55;white-space:pre-wrap;">{escape(prompt_text)}</div>'
            '</div>'
        )
    notes_html = ''
    if user_notes:
        notes_html = (
            '<div style="margin-bottom:14px;">'
            '<p style="font-size:12px;font-weight:600;color:#004C8C;text-transform:uppercase;margin:0 0 6px;">Consultant Notes</p>'
            f'<div style="background:white;border:1px solid #e2e8f0;border-radius:8px;padding:14px;color:#334155;font-size:14px;line-height:1.55;white-space:pre-wrap;">{escape(user_notes.strip())}</div>'
            '</div>'
        )
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:640px;margin:0 auto;">
      <div style="background:#004C8C;padding:18px 22px;border-radius:8px 8px 0 0;">
        <p style="color:white;font-size:17px;font-weight:600;margin:0;">Proposal Comparison</p>
      </div>
      <div style="background:#f8fafc;padding:22px;border:1px solid #e2e8f0;border-top:none;border-radius:0 0 8px 8px;">
        <p style="color:#334155;font-size:14px;margin:0 0 16px;"><strong>From:</strong> {escape(name)}</p>
        {prompt_html}
        {notes_html}
        <div style="margin-bottom:14px;">
          <p style="font-size:12px;font-weight:600;color:#004C8C;text-transform:uppercase;margin:0 0 6px;">What Changed</p>
          <div style="background:white;border:1px solid #e2e8f0;border-radius:8px;padding:14px;color:#334155;font-size:14px;line-height:1.55;white-space:pre-wrap;">{escape(analysis_text)}</div>
        </div>
        <div style="margin-bottom:14px;">
          <p style="font-size:12px;font-weight:600;color:#004C8C;text-transform:uppercase;margin:0 0 6px;">Voice Guide Recommendations</p>
          <div style="background:#FFF9E6;border:1px solid #FFE082;border-radius:8px;padding:14px;color:#334155;font-size:14px;line-height:1.55;white-space:pre-wrap;">{escape(voice_text)}</div>
        </div>
        <p style="margin:0;">
          <a href="{admin_url}" style="color:#004C8C;font-weight:600;">Open comparisons in Admin →</a>
        </p>
      </div>
    </div>
    """
    _send_hub_notify_email(subject, text_body, html_body)


@app.route('/admin/delete-activity', methods=['POST'])
def delete_activity():
    # Allow any authenticated user to delete their own entries
    # Only admin can delete anyone's entries
    if not session.get('user_key'):
        return jsonify({'error': 'Unauthorized'}), 401
    item_type = request.form.get('type')
    item_id   = request.form.get('id')
    user_key  = request.form.get('user_key')
    # Must be deleting own record OR be admin
    if user_key != session.get('user_key') and session.get('role') != 'admin':
        return jsonify({'error': 'Can only delete your own entries'}), 403
    table_map = {'proposal': 'proposal_log', 'ppm': 'ppm_log', 'subscope': 'subscope_log'}
    table = table_map.get(item_type)
    if not table:
        return jsonify({'error': 'Invalid type'}), 400
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute(f'DELETE FROM {table} WHERE id = %s AND generated_by = %s',
                        (item_id, user_key))
            conn.commit()
            cur.close()
            conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


@app.route('/admin/clear-my-data', methods=['POST'])
def clear_my_data():
    if session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            for table in ['proposal_log', 'ppm_log', 'subscope_log']:
                cur.execute(f"DELETE FROM {table} WHERE generated_by = 'thomas_ellison'")
            conn.commit()
            cur.close()
            conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


@app.route('/admin/clear-my-profile', methods=['POST'])
def clear_my_profile():
    if session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            cur.execute("DELETE FROM profile_results WHERE LOWER(name) = LOWER('Thomas Ellison')")
            conn.commit()
            cur.close()
            conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


@app.route('/admin/member/<user_key>')
def admin_member(user_key):
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    user_def = USERS.get(user_key)
    if not user_def:
        return redirect(url_for('admin'))
    proposals, ppms, subscopes, feedback_items = [], [], [], []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM proposal_log WHERE generated_by = %s ORDER BY generated_at DESC', (user_key,))
            proposals = cur.fetchall()
            cur.execute('SELECT * FROM ppm_log WHERE generated_by = %s ORDER BY generated_at DESC', (user_key,))
            ppms = cur.fetchall()
            cur.execute('SELECT * FROM subscope_log WHERE generated_by = %s ORDER BY generated_at DESC', (user_key,))
            subscopes = cur.fetchall()
            cur.execute('SELECT * FROM feedback WHERE user_key = %s ORDER BY submitted_at DESC', (user_key,))
            feedback_items = cur.fetchall()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Member detail error: {e}")
    return render_template('admin_member.html',
                           user=user_def, user_key=user_key,
                           proposals=proposals, ppms=ppms,
                           subscopes=subscopes,
                           feedback_items=feedback_items)


@app.route('/admin/feedback')
@require_login
def admin_feedback():
    """Team feedback inbox — admin only. Supports archive + permanent clear."""
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    show_archived = (request.args.get('show') or '').strip().lower() == 'archived'
    items = []
    counts = {'active': 0, 'archived': 0}
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute(
                'SELECT COUNT(*) AS c FROM feedback WHERE COALESCE(archived, FALSE) = FALSE'
            )
            counts['active'] = cur.fetchone()['c'] or 0
            cur.execute(
                'SELECT COUNT(*) AS c FROM feedback WHERE COALESCE(archived, FALSE) = TRUE'
            )
            counts['archived'] = cur.fetchone()['c'] or 0
            if show_archived:
                cur.execute(
                    '''SELECT * FROM feedback
                       WHERE COALESCE(archived, FALSE) = TRUE
                       ORDER BY COALESCE(archived_at, submitted_at) DESC'''
                )
            else:
                cur.execute(
                    '''SELECT * FROM feedback
                       WHERE COALESCE(archived, FALSE) = FALSE
                       ORDER BY submitted_at DESC'''
                )
            items = cur.fetchall()
            # Mark visible active items as read (badge clear)
            if not show_archived:
                cur.execute(
                    '''UPDATE feedback SET read_by_admin = TRUE
                       WHERE COALESCE(archived, FALSE) = FALSE'''
                )
            conn.commit()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Feedback error: {e}")
    return render_template(
        'admin_feedback.html',
        items=items,
        show_archived=show_archived,
        counts=counts,
    )


@app.route('/admin/feedback/<int:item_id>/archive', methods=['POST'])
@require_login
def admin_feedback_archive_one(item_id):
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin only'}), 403
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        cur.execute(
            '''UPDATE feedback
               SET archived = TRUE, archived_at = NOW(), read_by_admin = TRUE
               WHERE id = %s''',
            (item_id,),
        )
        conn.commit()
        cur.close()
        conn.close()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or (
            request.accept_mimetypes.best == 'application/json'
        ):
            return jsonify({'success': True})
        return redirect(url_for('admin_feedback'))
    except Exception as e:
        print(f'Feedback archive one error: {e}')
        return jsonify({'error': 'Could not archive'}), 500


@app.route('/admin/feedback/<int:item_id>/unarchive', methods=['POST'])
@require_login
def admin_feedback_unarchive_one(item_id):
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin only'}), 403
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        cur.execute(
            '''UPDATE feedback
               SET archived = FALSE, archived_at = NULL
               WHERE id = %s''',
            (item_id,),
        )
        conn.commit()
        cur.close()
        conn.close()
        return redirect(url_for('admin_feedback', show='archived'))
    except Exception as e:
        print(f'Feedback unarchive error: {e}')
        return redirect(url_for('admin_feedback', show='archived'))


@app.route('/admin/feedback/<int:item_id>/delete', methods=['POST'])
@require_login
def admin_feedback_delete_one(item_id):
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin only'}), 403
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        cur.execute('DELETE FROM feedback WHERE id = %s', (item_id,))
        conn.commit()
        cur.close()
        conn.close()
        next_show = (request.form.get('show') or request.args.get('show') or '').strip()
        if next_show == 'archived':
            return redirect(url_for('admin_feedback', show='archived'))
        return redirect(url_for('admin_feedback'))
    except Exception as e:
        print(f'Feedback delete one error: {e}')
        return redirect(url_for('admin_feedback'))


@app.route('/admin/feedback/archive-all', methods=['POST'])
@require_login
def admin_feedback_archive_all():
    """Hide all active feedback from the main inbox (soft clear)."""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin only'}), 403
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        cur.execute(
            '''UPDATE feedback
               SET archived = TRUE, archived_at = NOW(), read_by_admin = TRUE
               WHERE COALESCE(archived, FALSE) = FALSE'''
        )
        n = cur.rowcount
        conn.commit()
        cur.close()
        conn.close()
        return redirect(url_for('admin_feedback'))
    except Exception as e:
        print(f'Feedback archive all error: {e}')
        return redirect(url_for('admin_feedback'))


@app.route('/admin/feedback/delete-all', methods=['POST'])
@require_login
def admin_feedback_delete_all():
    """Permanently delete feedback. scope=active|archived|all."""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin only'}), 403
    scope = (request.form.get('scope') or 'active').strip().lower()
    if scope not in ('active', 'archived', 'all'):
        scope = 'active'
    # Require typed confirmation for permanent delete
    confirm = (request.form.get('confirm') or '').strip().upper()
    if confirm != 'DELETE':
        return redirect(url_for('admin_feedback', err='confirm'))
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 503
        cur = conn.cursor()
        if scope == 'all':
            cur.execute('DELETE FROM feedback')
        elif scope == 'archived':
            cur.execute('DELETE FROM feedback WHERE COALESCE(archived, FALSE) = TRUE')
        else:
            cur.execute('DELETE FROM feedback WHERE COALESCE(archived, FALSE) = FALSE')
        conn.commit()
        cur.close()
        conn.close()
        if scope == 'archived':
            return redirect(url_for('admin_feedback', show='archived'))
        return redirect(url_for('admin_feedback'))
    except Exception as e:
        print(f'Feedback delete all error: {e}')
        return redirect(url_for('admin_feedback'))


# Per-table ceiling on the rows Team View serialises into the page. Twelve
# weeks across thirteen people has never come close to this; it exists so a
# bulk import or a runaway integration cannot put a 40MB page in front of
# everyone. The template reports when it bites rather than silently truncating.
TEAM_VIEW_ROW_CAP = 2000


def _serialize_log_rows(rows):
    """Convert DB rows to JSON-safe dicts for templates."""
    out = []
    for row in rows or []:
        d = dict(row)
        for k, v in d.items():
            if hasattr(v, 'isoformat'):
                d[k] = v.isoformat()
        out.append(d)
    return out


@app.route('/my-proposals')
@require_login
def my_proposals():
    """Proposal history: consultants see their book; PMs see own + paired consultants'."""
    user_key = session['user_key']
    user = USERS.get(user_key, {})
    role = user.get('role')
    # Consultants, PMs (incl. Trey), admin — anyone who can generate proposals
    if role not in ('consultant', 'pm', 'admin'):
        return redirect(url_for('dashboard'))

    rows = []
    stats = {'total': 0, 'last_30': 0, 'with_file': 0, 'mine': 0, 'listed': 0}
    access = get_user_proposal_access(user_key)
    # Cap list size so PM "all" access (Trey) cannot hang the page / browser
    list_limit = 300
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            if role == 'admin':
                cur.execute(
                    'SELECT * FROM proposal_log ORDER BY generated_at DESC LIMIT %s',
                    (list_limit,),
                )
                cur2_sql = '''SELECT COUNT(*) AS c FROM proposal_log
                              WHERE generated_at >= NOW() - INTERVAL '30 days' '''
                cur2_args = ()
                total_sql = 'SELECT COUNT(*) AS c FROM proposal_log'
                total_args = ()
            elif access:
                cur.execute(
                    '''SELECT * FROM proposal_log
                       WHERE generated_by = %s OR consultant_key = ANY(%s)
                       ORDER BY generated_at DESC
                       LIMIT %s''',
                    (user_key, list(access), list_limit),
                )
                cur2_sql = '''SELECT COUNT(*) AS c FROM proposal_log
                              WHERE (generated_by = %s OR consultant_key = ANY(%s))
                              AND generated_at >= NOW() - INTERVAL '30 days' '''
                cur2_args = (user_key, list(access))
                total_sql = '''SELECT COUNT(*) AS c FROM proposal_log
                               WHERE generated_by = %s OR consultant_key = ANY(%s)'''
                total_args = (user_key, list(access))
            else:
                cur.execute(
                    '''SELECT * FROM proposal_log
                       WHERE generated_by = %s
                       ORDER BY generated_at DESC LIMIT %s''',
                    (user_key, list_limit),
                )
                cur2_sql = '''SELECT COUNT(*) AS c FROM proposal_log
                              WHERE generated_by = %s
                              AND generated_at >= NOW() - INTERVAL '30 days' '''
                cur2_args = (user_key,)
                total_sql = 'SELECT COUNT(*) AS c FROM proposal_log WHERE generated_by = %s'
                total_args = (user_key,)

            rows = cur.fetchall()
            stats['listed'] = len(rows)
            stats['with_file'] = sum(1 for r in rows if r.get('document_id'))
            stats['mine'] = sum(1 for r in rows if r.get('generated_by') == user_key)

            cur.execute(total_sql, total_args)
            stats['total'] = cur.fetchone()['c'] or 0
            cur.execute(cur2_sql, cur2_args)
            stats['last_30'] = cur.fetchone()['c'] or 0
            cur.close()
            conn.close()
    except Exception as e:
        print(f"My proposals error: {e}")
        import traceback
        traceback.print_exc()

    if role == 'pm':
        page_sub = (
            f"{stats['total']} proposal{'s' if stats['total'] != 1 else ''} "
            f"in your access (you + consultants you support)"
            f" · showing latest {stats['listed']}"
            f" · {stats['mine']} generated by you"
        )
    elif role == 'admin':
        page_sub = f"{stats['total']} proposals company-wide · showing latest {stats['listed']}"
    else:
        page_sub = (
            f"{stats['total']} proposal{'s' if stats['total'] != 1 else ''} you've generated "
            f"· re-download or regenerate anytime"
        )

    return render_template(
        'my_proposals.html',
        user=user,
        user_key=user_key,
        rows=rows,
        stats=stats,
        page_sub=page_sub,
        proposal_url=os.environ.get('PROPOSAL_URL', 'https://pps-proposal-tool.onrender.com'),
    )


@app.route('/my-ppms')
@require_login
def my_ppms():
    user_key = session['user_key']
    user = USERS.get(user_key, {})
    if user.get('role') not in ('pm', 'admin'):
        return redirect(url_for('dashboard'))

    rows = []
    stats = {'total': 0, 'last_30': 0, 'as_pm': 0}
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute(
                '''SELECT * FROM ppm_log
                   WHERE generated_by = %s OR pm_key = %s
                   ORDER BY generated_at DESC''',
                (user_key, user_key)
            )
            rows = cur.fetchall()
            stats['total'] = len(rows)
            stats['as_pm'] = sum(1 for r in rows if r.get('pm_key') == user_key)
            cur.execute(
                '''SELECT COUNT(*) AS c FROM ppm_log
                   WHERE (generated_by = %s OR pm_key = %s)
                   AND generated_at >= NOW() - INTERVAL '30 days' ''',
                (user_key, user_key)
            )
            stats['last_30'] = cur.fetchone()['c'] or 0
            cur.close()
            conn.close()
    except Exception as e:
        print(f"My PPMs error: {e}")

    return render_template(
        'my_ppms.html',
        user=user,
        user_key=user_key,
        rows=rows,
        stats=stats,
        proposal_url=os.environ.get('PROPOSAL_URL', 'https://pps-proposal-tool.onrender.com'),
    )


@app.route('/my-tpscopes')
@require_login
def my_tpscopes():
    """Trade Partner Scope history: anything you generated or are listed as PM on."""
    user_key = session['user_key']
    user = USERS.get(user_key, {})
    role = user.get('role')
    # Anyone who can run TPS / production tools (PMs, consultants, admin)
    if role not in ('pm', 'consultant', 'admin'):
        return redirect(url_for('dashboard'))

    rows = []
    stats = {'total': 0, 'last_30': 0, 'mine': 0, 'as_pm': 0}
    list_limit = 300
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            if role == 'admin':
                cur.execute(
                    'SELECT * FROM subscope_log ORDER BY generated_at DESC LIMIT %s',
                    (list_limit,),
                )
                rows = cur.fetchall()
                cur.execute('SELECT COUNT(*) AS c FROM subscope_log')
                stats['total'] = cur.fetchone()['c'] or 0
                cur.execute(
                    '''SELECT COUNT(*) AS c FROM subscope_log
                       WHERE generated_at >= NOW() - INTERVAL '30 days' '''
                )
                stats['last_30'] = cur.fetchone()['c'] or 0
            else:
                cur.execute(
                    '''SELECT * FROM subscope_log
                       WHERE generated_by = %s OR pm_key = %s
                       ORDER BY generated_at DESC
                       LIMIT %s''',
                    (user_key, user_key, list_limit),
                )
                rows = cur.fetchall()
                cur.execute(
                    '''SELECT COUNT(*) AS c FROM subscope_log
                       WHERE generated_by = %s OR pm_key = %s''',
                    (user_key, user_key),
                )
                stats['total'] = cur.fetchone()['c'] or 0
                cur.execute(
                    '''SELECT COUNT(*) AS c FROM subscope_log
                       WHERE (generated_by = %s OR pm_key = %s)
                       AND generated_at >= NOW() - INTERVAL '30 days' ''',
                    (user_key, user_key),
                )
                stats['last_30'] = cur.fetchone()['c'] or 0
            stats['mine'] = sum(1 for r in rows if r.get('generated_by') == user_key)
            stats['as_pm'] = sum(1 for r in rows if r.get('pm_key') == user_key)
            cur.close()
            conn.close()
    except Exception as e:
        print(f"My TPS error: {e}")
        import traceback
        traceback.print_exc()

    if role == 'admin':
        page_sub = (
            f"{stats['total']} TPS company-wide · showing latest {len(rows)} · "
            f"filter with search"
        )
    else:
        page_sub = (
            f"{stats['total']} Trade Partner Scope{'s' if stats['total'] != 1 else ''} "
            f"you generated or are listed as PM on · showing latest {len(rows)}"
        )

    return render_template(
        'my_tpscopes.html',
        user=user,
        user_key=user_key,
        rows=rows,
        rows_json=_serialize_log_rows(rows),
        stats=stats,
        page_sub=page_sub,
        proposal_url=os.environ.get('PROPOSAL_URL', 'https://pps-proposal-tool.onrender.com'),
    )


@app.route('/admin/proposals')
def admin_proposals():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    from datetime import datetime as _dt, timedelta
    consultant_filter = request.args.get('consultant', '')
    period = request.args.get('period', 'all')
    rows = []
    counts = {}
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            # Get all proposals filtered by consultant_key (for)
            if consultant_filter:
                cur.execute(
                    'SELECT * FROM proposal_log WHERE consultant_key = %s ORDER BY generated_at DESC',
                    (consultant_filter,)
                )
            else:
                cur.execute('SELECT * FROM proposal_log ORDER BY generated_at DESC')
            rows = cur.fetchall()
            # Get counts per consultant
            cur.execute('''
                SELECT consultant_key, consultant_name,
                    COUNT(*) as total,
                    COUNT(*) FILTER (WHERE generated_at >= NOW() - INTERVAL '30 days') as last_30
                FROM proposal_log GROUP BY consultant_key, consultant_name ORDER BY consultant_name
            ''')
            counts = {r['consultant_key']: r for r in cur.fetchall()}
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Admin proposals error: {e}")
    return render_template('admin_proposals.html',
                           rows=rows, counts=counts,
                           consultant_filter=consultant_filter,
                           consultants=CONSULTANTS)


@app.route('/admin/stats')
def admin_stats():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    from datetime import datetime as _dt
    stats = {}
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            for table, key in [
                ('proposal_log', 'proposals'),
                ('ppm_log', 'ppms'),
                ('subscope_log', 'subscopes'),
                ('profile_results', 'profiles'),
            ]:
                try:
                    cur.execute(f'''
                        SELECT COUNT(*) as total,
                        COUNT(*) FILTER (WHERE generated_at >= NOW() - INTERVAL '30 days') as last_30
                        FROM {table}
                    ''')
                    stats[key] = cur.fetchone()
                except:
                    stats[key] = {'total': 0, 'last_30': 0}
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Stats error: {e}")
    return jsonify(stats)


@app.route('/generate-code', methods=['POST'])
@require_login
def generate_code():
    """Called by hub dashboard JS — returns a one-time SSO code for satellite tools."""
    user_key = session['user_key']
    code = generate_sso_code(get_db, user_key, session.get('display_name', ''), session.get('role', 'user'))
    if not code:
        return jsonify({'error': 'Could not create sign-in code. Check database connection.'}), 500
    return jsonify({'code': code})


@app.route('/generate-token', methods=['POST'])
@require_login
def generate_token():
    """Legacy alias — returns a code (not a long-lived token)."""
    user_key = session['user_key']
    code = generate_sso_code(get_db, user_key, session.get('display_name', ''), session.get('role', 'user'))
    if not code:
        return jsonify({'error': 'Could not create sign-in code.'}), 500
    return jsonify({'code': code, 'token': code})


@app.route('/exchange-code', methods=['POST'])
def exchange_code_route():
    """Server-to-server: satellite tools exchange a one-time code for user info."""
    api_key = request.headers.get('X-API-Key', '')
    if api_key != INTERNAL_API_KEY:
        return jsonify({'error': 'Unauthorized'}), 401
    data = request.get_json(silent=True) or {}
    code = (data.get('code') or '').strip()
    if not code:
        return jsonify({'valid': False, 'reason': 'code required'}), 400
    row = exchange_sso_code(get_db, code)
    if not row:
        return jsonify({'valid': False, 'reason': 'Code invalid or expired'})
    _touch_last_active(row['user_key'], force=True)
    return jsonify({'valid': True, **row})


@app.route('/validate-token', methods=['POST'])
def validate_token():
    """Legacy token validation — deprecated; use /exchange-code."""
    api_key = request.headers.get('X-API-Key', '')
    if api_key != INTERNAL_API_KEY:
        return jsonify({'error': 'Unauthorized'}), 401
    data = request.get_json(silent=True) or {}
    token = (data.get('token') or '').strip()
    if not token:
        return jsonify({'valid': False}), 400
    row = exchange_sso_code(get_db, token)
    if row:
        _touch_last_active(row['user_key'], force=True)
        return jsonify({'valid': True, **row})
    return jsonify({'valid': False, 'reason': 'Token invalid or expired'})


def _upload_size_bytes(file_storage):
    pos = file_storage.tell()
    file_storage.seek(0, os.SEEK_END)
    size = file_storage.tell()
    file_storage.seek(pos)
    return size


def _extract_upload_text(file_storage, label='file'):
    """Extract plain text from an uploaded proposal (.docx, .txt; .pdf if pdftotext available)."""
    if not file_storage or not file_storage.filename:
        raise ValueError(f'Missing {label}')
    size = _upload_size_bytes(file_storage)
    if size > MAX_DOCUMENT_BYTES:
        limit_mb = MAX_DOCUMENT_BYTES // (1024 * 1024)
        raise ValueError(f'{label} exceeds {limit_mb} MB — use a smaller file.')
    filename = file_storage.filename.lower()
    if filename.endswith('.docx'):
        from docx import Document as DocxDoc
        data = file_storage.read()
        doc = DocxDoc(BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return '\n'.join(parts)
    if filename.endswith('.txt'):
        return file_storage.read().decode('utf-8', errors='ignore')
    if filename.endswith('.pdf'):
        import subprocess
        import tempfile
        data = file_storage.read()
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            result = subprocess.run(
                ['pdftotext', tmp_path, '-'],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
        raise ValueError('Could not extract text from PDF — upload a Word (.docx) or text file instead.')
    raise ValueError('Unsupported file type — use .docx or .txt')


@app.route('/analyze-diff', methods=['POST'])
def analyze_diff():
    """Compare original vs edited proposal; return Claude analysis for voice guide updates."""
    if not session.get('user_key'):
        return jsonify({'error': 'Not authenticated'}), 401
    if not CLAUDE_API_KEY:
        return jsonify({'error': 'Claude API key not configured on hub (CLAUDE_API_KEY).'}), 503

    original_file = request.files.get('original')
    edited_file = request.files.get('edited')

    try:
        original_text = _extract_upload_text(original_file, 'original proposal')
        edited_text = _extract_upload_text(edited_file, 'edited proposal')
    except ValueError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
    except Exception as e:
        _log_exception(e, 'analyze-diff')
        return jsonify({'success': False, 'error': 'Could not read one or both files.'}), 400

    if not original_text.strip() or not edited_text.strip():
        return jsonify({'success': False, 'error': 'Could not extract text from one or both files.'}), 400

    comparison_prompt = (request.form.get('comparison_prompt') or '').strip()[:2000]

    prompt = f"""You are helping improve the PPS (Pure Property Solutions) construction proposal voice guide.

A consultant generated a proposal with AI, then edited it before sending to the client.
Compare the ORIGINAL and EDITED versions. Focus on meaningful changes to tone, structure,
wording, scope language, and client-facing phrasing — not trivial formatting."""

    if comparison_prompt:
        prompt += f"""

ADDITIONAL INSTRUCTIONS FROM THE CONSULTANT (apply on top of the standard comparison above — do not ignore the base task):
{comparison_prompt}"""

    prompt += f"""

ORIGINAL (AI-generated):
{original_text[:14000]}

EDITED (consultant's final version):
{edited_text[:14000]}

Respond with ONLY valid JSON (no markdown fences) using exactly these keys:
{{
  "diff_analysis": "Bullet-style summary of what changed and why it matters for client-facing proposals",
  "voice_recommendations": "Specific, actionable updates to recommend for pps_voice.txt — phrasing rules, tone shifts, trade language, or sections to add"
}}"""

    try:
        raw = _claude_roleplay_call(
            'You compare proposal drafts and return strict JSON only.',
            [{'role': 'user', 'content': prompt}],
            max_tokens=2500,
            timeout=90.0,
        )
        if raw.startswith('```'):
            raw = raw.split('\n', 1)[-1]
            if raw.endswith('```'):
                raw = raw.rsplit('```', 1)[0]
            raw = raw.strip()
        parsed = json.loads(raw)
        diff_analysis = (parsed.get('diff_analysis') or '').strip()
        voice_recommendations = (parsed.get('voice_recommendations') or '').strip()
        if not diff_analysis and not voice_recommendations:
            return jsonify({'success': False, 'error': 'Claude returned an empty analysis.'}), 500

        user_key = session['user_key']
        display_name = session.get('display_name', '')
        diff_id = _save_proposal_diff(
            user_key, display_name, diff_analysis, voice_recommendations,
            comparison_prompt=comparison_prompt,
        )
        _notify_proposal_diff_email(
            display_name, '', diff_analysis, voice_recommendations,
            comparison_prompt=comparison_prompt,
        )

        return jsonify({
            'success': True,
            'diff_analysis': diff_analysis,
            'voice_recommendations': voice_recommendations,
            'diff_id': diff_id,
            'shared_with_admin': bool(diff_id),
        })
    except json.JSONDecodeError:
        return jsonify({'success': False, 'error': 'Could not parse Claude response. Try again.'}), 500
    except Exception as e:
        _log_exception(e, 'analyze-diff')
        err_name = type(e).__name__
        if err_name in ('APITimeoutError', 'TimeoutError') or 'timeout' in str(e).lower():
            return jsonify({
                'success': False,
                'error': 'Analysis timed out. Try again with smaller files or a shorter comparison prompt.',
            }), 504
        return jsonify({'success': False, 'error': 'Analysis failed. Please try again.'}), 500


@app.route('/submit-diff', methods=['POST'])
def submit_diff():
    """Optional follow-up: attach consultant notes to an existing comparison."""
    if not session.get('user_key'):
        return jsonify({'error': 'Not authenticated'}), 401
    diff_id = request.form.get('diff_id', '').strip()
    user_notes = request.form.get('user_notes', '').strip()
    if not user_notes:
        return jsonify({'success': False, 'error': 'No notes to save.'}), 400
    user_key = session['user_key']
    display_name = session.get('display_name', '')
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 500
        cur = conn.cursor(cursor_factory=RealDictCursor)
        comparison_prompt = ''
        if diff_id:
            cur.execute(
                '''SELECT id, diff_analysis, voice_recommendations, comparison_prompt
                   FROM proposal_diffs WHERE id = %s AND user_key = %s''',
                (diff_id, user_key),
            )
            row = cur.fetchone()
            if not row:
                cur.close()
                conn.close()
                return jsonify({'error': 'Comparison not found.'}), 404
            cur.execute(
                'UPDATE proposal_diffs SET user_notes = %s WHERE id = %s',
                (user_notes, diff_id),
            )
            diff_analysis = row.get('diff_analysis') or ''
            voice_recommendations = row.get('voice_recommendations') or ''
            comparison_prompt = row.get('comparison_prompt') or ''
        else:
            diff_analysis = request.form.get('diff_analysis', '').strip()
            voice_recommendations = request.form.get('voice_recommendations', '').strip()
            cur.execute(
                '''INSERT INTO proposal_diffs
                   (user_key, display_name, property_name, diff_analysis, user_notes, voice_recommendations)
                   VALUES (%s, %s, %s, %s, %s, %s)''',
                (user_key, display_name, '', diff_analysis, user_notes, voice_recommendations),
            )
        conn.commit()
        cur.close()
        conn.close()
        _notify_proposal_diff_email(
            display_name, user_notes, diff_analysis, voice_recommendations,
            comparison_prompt=comparison_prompt,
        )
        return jsonify({'success': True})
    except Exception as e:
        return _api_error(e)


@app.route('/admin/diffs')
def admin_diffs():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    rows = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM proposal_diffs ORDER BY submitted_at DESC')
            rows = cur.fetchall()
            cur.execute('UPDATE proposal_diffs SET reviewed_by_admin = TRUE')
            conn.commit()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Admin diffs error: {e}")
    return render_template('admin_diffs.html', rows=rows)


@app.route('/my-diffs')
def my_diffs():
    if not session.get('user_key'):
        return redirect(url_for('login'))
    rows = []
    user_key = session['user_key']
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            if session.get('role') == 'admin':
                cur.execute('SELECT * FROM proposal_diffs ORDER BY submitted_at DESC')
            else:
                cur.execute(
                    'SELECT * FROM proposal_diffs WHERE user_key = %s ORDER BY submitted_at DESC',
                    (user_key,)
                )
            rows = cur.fetchall()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"My diffs error: {e}")
    return render_template('proposal_diff.html', rows=rows,
                           is_admin=session.get('role') == 'admin')


@app.route('/team-view')
@require_login
def team_view():
    """Per-person activity, scored the same way the Monday email scores it.

    Two things were wrong here before 2026-08-23:

    **It disagreed with the recap about the same people.** The cards counted
    every row a person had ever generated, with no window and no activity cap,
    while the email counts one week plus a rolling twelve with pipeline capped
    at five. Andy's number in his inbox and Andy's number on this page were
    different numbers with his name on both. Both now come from
    `weekly_recap.collect_scores` — one scoring function, not two.

    **It grew without bound.** Three `SELECT *` per person with no date filter
    and no LIMIT — thirty-nine queries fetching every row ever written, all
    serialised into the page as JSON. Now three windowed queries for the whole
    roster, plus grouped counts for the all-time figures, which is what people
    actually want the lifetime number for: a number, not five hundred rows.
    """
    user_key = session['user_key']
    user = USERS.get(user_key, {})
    member_keys = list(USERS.keys())
    members = [
        {'key': k, 'display': u['display'], 'title': u.get('title', ''), 'role': u['role']}
        for k, u in USERS.items() if u
    ]

    week_start, week_end = weekly_recap.last_week_bounds()
    roll_start, roll_end = weekly_recap.rolling_bounds()
    week_scores = weekly_recap.collect_scores(get_db, USERS, week_start, week_end)
    roll_scores = weekly_recap.collect_scores(get_db, USERS, roll_start, roll_end)

    member_data = {k: {'proposals': [], 'ppms': [], 'tpscopes': []} for k in member_keys}
    lifetime = {k: {'proposals': 0, 'ppms': 0, 'tpscopes': 0} for k in member_keys}

    # Detail rows are windowed to the same twelve weeks the rolling score covers,
    # so the lists and the number above them describe the same period. Older work
    # is still reachable in full from /my-proposals, /admin/proposals and the
    # Admin search — this page is a scoreboard, not an archive.
    DETAIL = (
        ('proposals', 'proposal_log'),
        ('ppms', 'ppm_log'),
        ('tpscopes', 'subscope_log'),
    )
    conn = None
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            for field, table in DETAIL:
                try:
                    cur.execute(
                        f'SELECT * FROM {table} '
                        f'WHERE generated_by = ANY(%s) AND generated_at >= %s '
                        f'ORDER BY generated_at DESC LIMIT %s',
                        (member_keys, roll_start, TEAM_VIEW_ROW_CAP),
                    )
                    for row in cur.fetchall():
                        owner = row.get('generated_by')
                        if owner in member_data:
                            member_data[owner][field].append(row)
                    for key in member_keys:
                        member_data[key][field] = _serialize_log_rows(member_data[key][field])
                    cur.execute(
                        f'SELECT generated_by, COUNT(*) AS c FROM {table} GROUP BY 1'
                    )
                    for row in cur.fetchall():
                        if row['generated_by'] in lifetime:
                            lifetime[row['generated_by']][field] = row['c'] or 0
                except Exception as e:
                    # A missing table is survivable — several are created lazily
                    # on first use. Roll back so the next query can still run.
                    conn.rollback()
                    print(f'team view: skipped {table} ({e})')
            cur.close()
    except Exception as e:
        print(f"Team view error: {e}")
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass

    cards = []
    for m in members:
        key = m['key']
        week_break = week_scores.get(key, {})
        roll_break = roll_scores.get(key, {})
        cards.append({
            **m,
            'week_score': weekly_recap.score_total(week_break, weeks=1),
            'rolling_score': weekly_recap.score_total(
                roll_break, weeks=weekly_recap.ROLLING_WEEKS),
            'counts': {
                'proposals': len(member_data[key]['proposals']),
                'ppms': len(member_data[key]['ppms']),
                'tpscopes': len(member_data[key]['tpscopes']),
            },
            'lifetime': lifetime[key],
        })
    cards.sort(key=lambda c: (-c['week_score'], c['display'].lower()))

    return render_template(
        'team_view.html',
        user=user,
        user_key=user_key,
        members=cards,
        member_data=member_data,
        week_label=weekly_recap.week_label(week_start),
        rolling_weeks=weekly_recap.ROLLING_WEEKS,
        row_cap=TEAM_VIEW_ROW_CAP,
        is_admin=session.get('role') == 'admin',
        proposal_url=os.environ.get('PROPOSAL_URL', 'https://pps-proposal-tool.onrender.com'),
    )


@app.route('/admin/tpscopes')
def admin_tpscopes():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    rows = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM subscope_log ORDER BY generated_at DESC')
            rows = cur.fetchall()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"Admin TPS error: {e}")
    return render_template('admin_tpscopes.html', rows=rows)


# /admin/reset-team-view removed 2026-08-21. It toggled the per-user `team_view`
# field, which no longer exists — Team View is open to everyone. It was also
# quietly broken: it mutated the in-process USERS dict, so a change only ever
# applied to whichever of the two gunicorn workers served that request and was
# lost on the next deploy. No template linked it.


@app.route('/site-visit')
def site_visit():
    if not session.get('user_key'):
        return redirect(url_for('login'))
    user_key = session['user_key']
    return render_template(
        'site_visit.html',
        user_key=user_key,
        display_name=session.get('display_name', ''),
        user_email=USERS.get(user_key, {}).get('email', session.get('user_email', '')),
    )


@app.route('/site-visit/generate', methods=['POST'])
def site_visit_generate():
    if not session.get('user_key'):
        return jsonify({'error': 'Not authenticated'}), 401

    import json as _json
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data received'}), 400

    user_key     = session['user_key']
    display_name = session.get('display_name', '')
    data['visited_by'] = data.get('visited_by') or display_name

    action = data.get('action', 'both')  # 'download', 'save', 'both'

    # Save to DB if requested
    saved_id = None
    if action in ('save', 'both'):
        try:
            conn = get_db()
            if conn:
                cur = conn.cursor()
                cur.execute('''
                    INSERT INTO site_visit_log (
                        generated_by, display_name, property_name, property_address,
                        visit_date, visit_time, po_number,
                        trade_partner_present, trade_partner_company, crew_lead, crew_count,
                        met_with_staff, staff_contact, topics_discussed,
                        complaints_received, complaint_details,
                        checklist, overall_status, quality_status, schedule_status,
                        observations, photos_taken, next_visit_date
                    ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    RETURNING id
                ''', (
                    user_key, display_name,
                    data.get('property_name'), data.get('property_address'),
                    data.get('visit_date'), data.get('visit_time'), data.get('po_number'),
                    data.get('trade_partner_present'), data.get('trade_partner_company'),
                    data.get('crew_lead'), data.get('crew_count'),
                    data.get('met_with_staff'), data.get('staff_contact'),
                    data.get('topics_discussed'), data.get('complaints_received'),
                    data.get('complaint_details'),
                    _json.dumps(data.get('checklist', [])),
                    data.get('overall_status'), data.get('quality_status'),
                    data.get('schedule_status'), data.get('observations'),
                    bool(data.get('photos_taken')), data.get('next_visit_date'),
                ))
                saved_id = cur.fetchone()[0]
                conn.commit()
                cur.close()
                conn.close()
        except Exception as e:
            print(f"Site visit save error: {e}")

    if action == 'save':
        return jsonify({'success': True, 'saved_id': saved_id})

    # Generate Word doc
    try:
        import sys, os
        sys.path.insert(0, os.path.dirname(__file__))
        from site_visit_builder import build_site_visit
        buf = build_site_visit(data)
        from flask import send_file
        prop = data.get('property_name', 'Site_Visit').replace(' ', '_')
        date = data.get('visit_date', '').replace('/', '-').replace(' ', '_')
        filename = f"PPS_Site_Visit_{prop}_{date}.docx"
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        import traceback; traceback.print_exc()
        return _api_error(e)


@app.route('/site-visit/download/<int:visit_id>')
def site_visit_download(visit_id):
    user_key = session.get('user_key')
    if not user_key:
        return redirect(url_for('login'))
    role = session.get('role', '')
    row = None
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM site_visit_log WHERE id = %s', (visit_id,))
            row = cur.fetchone()
            cur.close(); conn.close()
        if not row:
            return "Not found", 404
        if not _can_view_activity_row(user_key, role, 'site_visit', row):
            return "Forbidden", 403
        import json as _json, sys, os
        data = dict(row)
        if isinstance(data.get('checklist'), str):
            data['checklist'] = _json.loads(data['checklist'])
        sys.path.insert(0, os.path.dirname(__file__))
        from site_visit_builder import build_site_visit
        buf = build_site_visit(data)
        from flask import send_file
        prop = (data.get('property_name') or 'Site_Visit').replace(' ','_')
        filename = f"PPS_Site_Visit_{prop}.docx"
        return send_file(buf, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        _log_exception(e)
        return GENERIC_DOWNLOAD_ERROR, 500


@app.route('/admin/site-visits')
def admin_site_visits():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    rows = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM site_visit_log ORDER BY generated_at DESC')
            rows = cur.fetchall()
            cur.close(); conn.close()
    except Exception as e:
        print(f"Admin site visits error: {e}")
    return render_template('admin_site_visits.html', rows=rows)


@app.route('/pm-training')
@require_login
def pm_training():
    """PM onboarding module — open to all logged-in users while under construction."""
    user_key = session['user_key']
    user = USERS.get(user_key, {})
    enrollment = get_pm_enrollment(user_key) or {}
    manager = USERS.get(enrollment.get('manager_key') or PM_TRAINING_MANAGER, {})
    meta, weeks = get_pm_training_curriculum()
    # Same overlay contract as PSC. The PM curriculum has no standalone
    # sections, so only the week containers are addressable — `apply` handles
    # the empty ones by ignoring them rather than inventing keys.
    (_ob, weeks, _cv, _st, _co), added_since = training_overlay.apply(
        (dict(), weeks, {}, {}, {}),
        training_overlay.load_overlay(get_db).get('pm'),
        enrolled_at=enrollment.get('enrolled_at'),
    )
    progress = get_pm_training_progress(user_key)
    notes = get_pm_training_notes(user_key)
    stats = compute_pm_training_stats(user_key)
    week_status = {wp['week']: wp for wp in stats['week_pcts']}
    return render_template(
        'pm_training.html',
        added_since=added_since,
        added_since_done=sum(
            1 for i in training_overlay.added_since_item_ids(added_since) if progress.get(i)
        ),
        meta=meta,
        weeks=weeks,
        total_items=count_pm_trackable_items(),
        progress_json=json.dumps(progress),
        notes_json=json.dumps(notes),
        enrollment=enrollment,
        manager=manager,
        stats=stats,
        week_status=week_status,
        user=user,
        under_construction=True,
    )


@app.route('/api/pm-training/progress', methods=['GET', 'POST'])
@require_login
def pm_training_progress_api():
    user_key = session['user_key']
    if request.method == 'GET':
        return jsonify({'progress': get_pm_training_progress(user_key)})
    data = request.get_json(silent=True) or {}
    progress = data.get('progress', {})
    if not isinstance(progress, dict):
        return jsonify({'error': 'Invalid progress data'}), 400
    ok = save_pm_training_progress(user_key, progress)
    if not ok:
        return jsonify({'error': 'Could not save progress'}), 500
    return jsonify({'success': True, 'stats': compute_pm_training_stats(user_key)})


@app.route('/api/pm-training/notes', methods=['GET', 'POST'])
@require_login
def pm_training_notes_api():
    user_key = session['user_key']
    if request.method == 'GET':
        return jsonify({'notes': get_pm_training_notes(user_key)})
    data = request.get_json(silent=True) or {}
    week_num = data.get('week')
    if week_num is None:
        return jsonify({'error': 'week required'}), 400
    notes_text = data.get('notes', '')
    ok = save_pm_training_notes(user_key, week_num, notes_text)
    if not ok:
        return jsonify({'error': 'Could not save notes'}), 500
    return jsonify({'success': True})


@app.route('/api/pm-training/feedback', methods=['POST'])
@require_login
def pm_training_feedback_api():
    user_key = session['user_key']
    user = USERS.get(user_key, {})
    data = request.get_json(silent=True) or {}
    message = (data.get('message') or '').strip()
    if not message:
        return jsonify({'error': 'Message required'}), 400
    week_num = data.get('week')
    if week_num is not None and week_num != '':
        try:
            week_num = int(week_num)
        except (TypeError, ValueError):
            week_num = None
    else:
        week_num = None
    feedback_type = (data.get('type') or 'improvement').strip()[:50]
    ok = submit_pm_training_feedback(
        user_key,
        user.get('display', user_key),
        message,
        week_num=week_num,
        feedback_type=feedback_type,
    )
    if not ok:
        return jsonify({'error': 'Could not submit feedback'}), 500
    return jsonify({'success': True})


@app.route('/pm-training/oversight')
@require_login
def pm_training_oversight():
    user_key = session['user_key']
    if not can_pm_training_oversight(user_key):
        return redirect(url_for('dashboard'))
    trainees = list_pm_enrolled_trainees()
    trainee_rows = []
    for t in trainees:
        key = t['user_key']
        stats = compute_pm_training_stats(key)
        trainee_rows.append({**t, 'stats': stats})
    # Everyone can open the module; also list recent progress users not formally enrolled
    enrollable = sorted(
        [{'key': k, 'display': v.get('display', k), 'role': v.get('role', '')}
         for k, v in USERS.items()],
        key=lambda u: u['display'],
    )
    return render_template(
        'pm_training_oversight.html',
        trainees=trainee_rows,
        enrollable=enrollable,
        feedback=list_pm_training_feedback(),
        manager_name=USERS.get(PM_TRAINING_MANAGER, {}).get('display', 'Production Manager'),
        is_admin=(USERS.get(user_key, {}).get('role') == 'admin'),
        under_construction=True,
    )


@app.route('/admin/pm-training')
@require_login
def admin_pm_training():
    """Old bookmark. See the note on admin_psc_training."""
    if not can_pm_training_oversight(session['user_key']):
        return redirect(url_for('dashboard'))
    return redirect(url_for('pm_training_oversight'))


@app.route('/api/pm-training/signoff', methods=['POST'])
@require_login
def pm_training_signoff_api():
    user_key = session['user_key']
    if not can_pm_training_oversight(user_key):
        return jsonify({'error': 'Not authorized'}), 403
    data = request.get_json(silent=True) or {}
    trainee_key = (data.get('user_key') or '').strip()
    try:
        week_num = int(data.get('week'))
    except (TypeError, ValueError):
        return jsonify({'error': 'week required'}), 400
    if not trainee_key:
        return jsonify({'error': 'user_key required'}), 400
    try:
        conn = get_db()
        if not conn:
            return jsonify({'error': 'Database unavailable'}), 500
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO pm_training_manager_signoffs (user_key, week_num, signed_by, signed_at)
            VALUES (%s, %s, %s, NOW())
            ON CONFLICT (user_key, week_num)
            DO UPDATE SET signed_by = EXCLUDED.signed_by, signed_at = NOW()
        ''', (trainee_key, week_num, user_key))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"PM signoff error: {e}")
        return jsonify({'error': 'Could not sign off'}), 500
    return jsonify({'success': True, 'stats': compute_pm_training_stats(trainee_key)})


@app.route('/api/pm-training/enroll', methods=['POST'])
@require_login
def pm_training_enroll_api():
    user_key = session['user_key']
    if not can_pm_training_oversight(user_key):
        return jsonify({'error': 'Not authorized'}), 403
    data = request.get_json(silent=True) or {}
    target_key = (data.get('user_key') or '').strip()
    action = (data.get('action') or 'enroll').strip()
    if not target_key:
        return jsonify({'error': 'user_key required'}), 400
    if action == 'unenroll':
        ok = unenroll_pm_trainee(target_key)
        return jsonify({'success': bool(ok)})
    if action == 'graduate':
        ok = graduate_pm_trainee(target_key)
        return jsonify({'success': bool(ok)})
    ok, err = enroll_pm_trainee(target_key, user_key, data.get('manager_key'))
    if not ok:
        return jsonify({'error': err or 'Could not enroll'}), 400
    return jsonify({'success': True})



@app.route('/psc-training')
@require_login
def psc_training():
    user_key = session['user_key']
    if not is_psc_training_enrolled(user_key):
        return redirect(url_for('dashboard'))
    user = USERS.get(user_key, {})
    enrollment = get_psc_enrollment(user_key) or {}
    manager = USERS.get(enrollment.get('manager_key') or PSC_TRAINING_MANAGER, {})
    curriculum = get_training_curriculum()
    # Overlay: leadership's published edits and additions. `enrolled_at` decides
    # placement — items published after this person started are handed back in
    # `added_since` and rendered at the end, so a week they already closed does
    # not reopen. Empty tables make this a no-op; see tests/test_training_overlay.
    curriculum, added_since = training_overlay.apply(
        curriculum,
        training_overlay.load_overlay(get_db).get('psc'),
        enrolled_at=enrollment.get('enrolled_at'),
    )
    onboarding, weeks, core_values, sales_training, company_operations = curriculum
    progress = get_psc_training_progress(user_key)
    notes = get_psc_training_notes(user_key)
    stats = compute_psc_training_stats(user_key)
    week_status = {wp['week']: wp for wp in stats['week_pcts']}
    return render_template(
        'psc_training.html',
        added_since=added_since,
        added_since_done=sum(
            1 for i in training_overlay.added_since_item_ids(added_since) if progress.get(i)
        ),
        meta=PSC_TRAINING_META,
        onboarding=onboarding,
        weeks=weeks,
        core_values=core_values,
        sales_training=sales_training,
        company_operations=company_operations,
        total_items=count_trackable_items(),
        progress_json=json.dumps(progress),
        notes_json=json.dumps(notes),
        enrollment=enrollment,
        manager=manager,
        stats=stats,
        week_status=week_status,
        user=user,
        roleplay_by_week=get_roleplay_week_links(),
        roleplay_sales_links=get_roleplay_sales_links(),
    )


@app.route('/api/psc-training/progress', methods=['GET', 'POST'])
@require_login
def psc_training_progress_api():
    user_key = session['user_key']
    if not is_psc_training_enrolled(user_key):
        return jsonify({'error': 'Not enrolled in PSC training'}), 403
    if request.method == 'GET':
        return jsonify({'progress': get_psc_training_progress(user_key)})
    data = request.get_json(silent=True) or {}
    progress = data.get('progress', {})
    if not isinstance(progress, dict):
        return jsonify({'error': 'Invalid progress data'}), 400
    ok = save_psc_training_progress(user_key, progress)
    if not ok:
        return jsonify({'error': 'Could not save progress'}), 500
    return jsonify({'success': True, 'stats': compute_psc_training_stats(user_key)})


@app.route('/api/psc-training/notes', methods=['GET', 'POST'])
@require_login
def psc_training_notes_api():
    user_key = session['user_key']
    if not is_psc_training_enrolled(user_key):
        return jsonify({'error': 'Not enrolled in PSC training'}), 403
    if request.method == 'GET':
        return jsonify({'notes': get_psc_training_notes(user_key)})
    data = request.get_json(silent=True) or {}
    week_num = data.get('week')
    if week_num is None:
        return jsonify({'error': 'week required'}), 400
    notes_text = data.get('notes', '')
    ok = save_psc_training_notes(user_key, week_num, notes_text)
    if not ok:
        return jsonify({'error': 'Could not save notes'}), 500
    return jsonify({'success': True})


@app.route('/api/psc-training/feedback', methods=['POST'])
@require_login
def psc_training_feedback_api():
    user_key = session['user_key']
    if not is_psc_training_enrolled(user_key):
        return jsonify({'error': 'Not enrolled in PSC training'}), 403
    user = USERS.get(user_key, {})
    data = request.get_json(silent=True) or {}
    message = (data.get('message') or '').strip()
    if not message:
        return jsonify({'error': 'Message required'}), 400
    week_num = data.get('week')
    if week_num is not None and week_num != '':
        try:
            week_num = int(week_num)
        except (TypeError, ValueError):
            week_num = None
    else:
        week_num = None
    feedback_type = (data.get('type') or 'improvement').strip()[:50]
    ok = submit_psc_training_feedback(
        user_key,
        user.get('display', user_key),
        message,
        week_num=week_num,
        feedback_type=feedback_type,
    )
    if not ok:
        return jsonify({'error': 'Could not submit feedback'}), 500
    return jsonify({'success': True})


@app.route('/psc-training/roleplay')
@require_login
def psc_roleplay_page():
    user_key = session['user_key']
    if not can_access_psc_roleplay(user_key):
        return redirect(url_for('dashboard'))
    user = USERS.get(user_key, {})
    stats = compute_psc_training_stats(user_key) if is_psc_training_enrolled(user_key) else None
    week_pcts = stats['week_pcts'] if stats else []
    suggested_ids = get_suggested_roleplay_ids(week_pcts)
    user_stats = get_roleplay_user_stats(user_key)
    scenarios = []
    for sc in PSC_ROLEPLAY_SCENARIOS:
        st = user_stats.get(sc['id'], {})
        scenarios.append({
            'id': sc['id'],
            'title': sc['title'],
            'segment': sc['segment'],
            'difficulty': sc['difficulty'],
            'trainee_brief': sc['trainee_brief'],
            'objectives': sc['objectives'],
            'opening_line': sc['opening_line'],
            'max_turns': sc.get('max_turns', 12),
            'segment_color': segment_color(sc['segment']),
            'attempts': st.get('attempts', 0),
            'best_overall': st.get('best_overall'),
            'best_result': st.get('best_result'),
            'suggested': sc['id'] in suggested_ids,
        })
    return render_template(
        'psc_roleplay.html',
        meta=PSC_TRAINING_META,
        scenarios=scenarios,
        scenarios_json=json.dumps(scenarios),
        suggested_ids=suggested_ids,
        display_name=user.get('display', user_key),
        segments=['Apartments', 'Condos', 'Hospitality / Commercial', 'Any'],
    )


@app.route('/api/psc-training/roleplay/turn', methods=['POST'])
@require_login
def psc_roleplay_turn_api():
    user_key = session['user_key']
    if not can_access_psc_roleplay(user_key):
        return jsonify({'error': 'Not authorized for role-play practice'}), 403
    if not CLAUDE_API_KEY:
        return jsonify({'error': 'Practice partner unavailable — try again later.'}), 503

    data = request.get_json(silent=True) or {}
    scenario_id = (data.get('scenario_id') or '').strip()
    scenario = get_roleplay_scenario(scenario_id)
    if not scenario:
        return jsonify({'error': 'Unknown scenario'}), 400

    max_turns = scenario.get('max_turns', 12)
    messages, err = _validate_roleplay_messages(data.get('messages'), max_turns)
    if err:
        return jsonify({'error': err}), 400

    trainee_turns = _count_trainee_turns(messages)
    if trainee_turns > max_turns:
        return jsonify({'error': f'Turn limit reached ({max_turns} trainee messages)'}), 400

    turn_count, grade_count = _roleplay_usage_today(user_key)
    if turn_count >= ROLEPLAY_DAILY_TURN_LIMIT:
        return jsonify({
            'error': f'Daily practice limit reached ({ROLEPLAY_DAILY_TURN_LIMIT} turns). Try again tomorrow.',
        }), 429

    if not messages or messages[-1].get('role') != 'user':
        return jsonify({'error': 'Last message must be from the trainee'}), 400

    system = _ROLEPLAY_PERSONA_SYSTEM.format(persona=scenario['persona'])
    api_messages = [{'role': m['role'], 'content': m['content']} for m in messages]

    try:
        reply = _claude_roleplay_call(system, api_messages, max_tokens=400)
    except Exception:
        return jsonify({'error': 'The practice partner is unavailable — try again in a minute.'}), 503

    _increment_roleplay_usage(user_key, turn_delta=1)
    turns_left = max(0, max_turns - trainee_turns)
    return jsonify({'success': True, 'reply': reply, 'turns_left': turns_left})


@app.route('/api/psc-training/roleplay/finish', methods=['POST'])
@require_login
def psc_roleplay_finish_api():
    user_key = session['user_key']
    if not can_access_psc_roleplay(user_key):
        return jsonify({'error': 'Not authorized for role-play practice'}), 403
    if not CLAUDE_API_KEY:
        return jsonify({'error': 'Practice partner unavailable — try again later.'}), 503

    data = request.get_json(silent=True) or {}
    scenario_id = (data.get('scenario_id') or '').strip()
    scenario = get_roleplay_scenario(scenario_id)
    if not scenario:
        return jsonify({'error': 'Unknown scenario'}), 400

    max_turns = scenario.get('max_turns', 12)
    messages, err = _validate_roleplay_messages(data.get('messages'), max_turns)
    if err:
        return jsonify({'error': err}), 400

    trainee_turns = _count_trainee_turns(messages)
    if trainee_turns < 3:
        return jsonify({
            'error': 'Have a real conversation first — send at least three messages before requesting feedback.',
        }), 400

    turn_count, grade_count = _roleplay_usage_today(user_key)
    if grade_count >= ROLEPLAY_DAILY_GRADE_LIMIT:
        return jsonify({
            'error': f'Daily feedback limit reached ({ROLEPLAY_DAILY_GRADE_LIMIT} sessions). Try again tomorrow.',
        }), 429

    feedback, err = _grade_roleplay_session(scenario, messages)
    if err:
        return jsonify({'error': err}), 500

    session_id = _save_roleplay_session(user_key, scenario_id, messages, feedback, trainee_turns)
    if not session_id:
        return jsonify({'error': 'Could not save session — try again.'}), 500

    _increment_roleplay_usage(user_key, grade_delta=1)
    return jsonify({'success': True, 'session_id': session_id, **feedback})


@app.route('/api/psc-training/roleplay/history')
@require_login
def psc_roleplay_history_api():
    user_key = session['user_key']
    target_key = (request.args.get('user_key') or '').strip() or user_key
    if target_key != user_key:
        if not can_psc_training_oversight(user_key):
            return jsonify({'error': 'Unauthorized'}), 403
    elif not can_access_psc_roleplay(user_key):
        return jsonify({'error': 'Not authorized'}), 403

    rows = get_roleplay_history_rows(target_key)
    history = []
    for row in rows:
        history.append({
            'scenario_id': row['scenario_id'],
            'overall': float(row['overall']) if row['overall'] is not None else None,
            'result': row['result'],
            'created_at': row['created_at'].isoformat() if row.get('created_at') else None,
        })
    return jsonify({'history': history, 'summary': get_roleplay_summary_for_oversight(target_key)})


def _psc_training_oversight_data(mark_read=True):
    enrolled = list_psc_enrolled_trainees()
    trainees = []
    for row in enrolled:
        key = row['user_key']
        stats = compute_psc_training_stats(key)
        trainees.append({
            'user_key': key,
            'display': row.get('display', key),
            'enrolled_at': row.get('enrolled_at'),
            'last_activity_at': row.get('last_activity_at'),
            'manager_key': row.get('manager_key'),
            'manager_display': USERS.get(row.get('manager_key', ''), {}).get('display', ''),
            'done': stats['done'],
            'pct': stats['pct'],
            'week_pcts': stats['week_pcts'],
            'roleplay': get_roleplay_summary_for_oversight(key),
        })
    trainees.sort(key=lambda t: (-t['pct'], t['display']))
    all_consultants = []
    for k, u in USERS.items():
        if u.get('role') != 'consultant':
            continue
        enr = get_psc_enrollment(k)
        all_consultants.append({
            'user_key': k,
            'display': u.get('display', k),
            'enrolled': bool(enr and enr.get('active') and not enr.get('graduated_at')),
        })
    feedback_items = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM psc_training_feedback ORDER BY submitted_at DESC LIMIT 100')
            feedback_items = cur.fetchall()
            if mark_read:
                cur.execute('UPDATE psc_training_feedback SET read_by_admin = TRUE WHERE read_by_admin = FALSE')
            conn.commit()
            cur.close()
            conn.close()
    except Exception as e:
        print(f"PSC training feedback error: {e}")
    return trainees, all_consultants, feedback_items


@app.route('/psc-training/oversight')
@require_login
def psc_training_oversight():
    user_key = session['user_key']
    if not can_psc_training_oversight(user_key):
        return redirect(url_for('dashboard'))
    trainees, all_consultants, feedback_items = _psc_training_oversight_data(mark_read=True)
    return render_template(
        'psc_training_oversight.html',
        trainees=trainees,
        all_consultants=all_consultants,
        feedback_items=feedback_items,
        total_items=count_trackable_items(),
        is_admin=(session.get('role') == 'admin'),
        # Enrolment is leadership as of 2026-08-23; `is_admin` now only drives
        # the back-link. Anyone past the guard above is leadership, so this is
        # True here — kept explicit so the template reads honestly and so
        # narrowing it later is a one-line change.
        can_enroll=can_psc_training_oversight(user_key),
        can_signoff=True,
        current_user_key=user_key,
        self_enrolled_as_trainee=is_psc_training_enrolled(user_key),
        manager_name=USERS.get(PSC_TRAINING_MANAGER, {}).get('display', 'VP Sales'),
    )


@app.route('/admin/psc-training')
@require_login
def admin_psc_training():
    """Old bookmark. Guarded by the page it lands on, and checked here too.

    Was `@require_admin`, which bounced Tony off a URL whose destination he can
    open directly — the oversight page itself is leadership. Both training
    redirects now check the same thing their target checks, rather than relying
    on the target to do it: a redirect that guards differently from its
    destination is either a locked door to an open room, or the reverse.
    """
    if not can_psc_training_oversight(session['user_key']):
        return redirect(url_for('dashboard'))
    return redirect(url_for('psc_training_oversight'))


@app.route('/api/psc-training/signoff', methods=['POST'])
@require_login
def psc_training_signoff_api():
    user_key = session['user_key']
    if not can_psc_training_oversight(user_key):
        return jsonify({'error': 'Unauthorized'}), 403
    data = request.get_json(silent=True) or {}
    trainee_key = (data.get('user_key') or '').strip()
    action = (data.get('action') or 'signoff').strip()
    if not trainee_key:
        return jsonify({'error': 'user_key required'}), 400
    try:
        week_num = int(data.get('week'))
    except (TypeError, ValueError):
        return jsonify({'error': 'week required'}), 400
    if action == 'revoke':
        if session.get('role') != 'admin':
            return jsonify({'error': 'Only admin can revoke sign-offs'}), 403
        ok = revoke_psc_manager_signoff(trainee_key, week_num)
        return jsonify({'success': ok}) if ok else (jsonify({'error': 'Could not revoke'}), 500)
    ok, err = manager_signoff_psc_week(trainee_key, week_num, user_key)
    if not ok:
        return jsonify({'error': err or 'Could not sign off'}), 400
    return jsonify({'success': True, 'stats': compute_psc_training_stats(trainee_key)})


@app.route('/api/psc-training/enroll', methods=['POST'])
@require_login
def psc_training_enroll_api():
    """Enrol, graduate or remove a PSC trainee. Leadership (2026-08-23).

    Was owner-only, which left Tony — the VP of Sales who actually runs PSC
    onboarding and who could already sign off every week — unable to add the
    hire he was about to sign off. The PM module never had that split: Trey has
    enrolled and removed PM trainees since it shipped. This makes the two
    modules agree, and Thomas asked for Stephanie on both, which the leadership
    tier already covers.

    Revoking a sign-off stays owner-only on purpose — see the branch in
    `psc_training_signoff_api`. Enrolling is a roster decision; erasing a
    completed record is an audit one.
    """
    if not can_psc_training_oversight(session['user_key']):
        return jsonify({'error': 'Not authorized'}), 403
    data = request.get_json(silent=True) or {}
    target_key = (data.get('user_key') or '').strip()
    action = (data.get('action') or 'enroll').strip()
    if not target_key:
        return jsonify({'error': 'user_key required'}), 400
    if action == 'graduate':
        ok = graduate_psc_trainee(target_key)
        return jsonify({'success': ok}) if ok else (jsonify({'error': 'Could not graduate'}), 500)
    if action == 'unenroll':
        ok = unenroll_psc_trainee(target_key)
        return jsonify({'success': ok}) if ok else (jsonify({'error': 'Could not unenroll'}), 500)
    ok, err = enroll_psc_trainee(target_key, session.get('user_key'), data.get('manager_key'))
    if not ok:
        return jsonify({'error': err or 'Could not enroll'}), 400
    return jsonify({'success': True})


def _send_resend_document(to_email, doc_type, filename, property_name, doc_b64, sender_name):
    """Send a document attachment via Resend (internal — self or coworker, not client delivery)."""
    import urllib.request as _ur
    import urllib.error as _ur_err
    import json as _json
    from pps_brand import EMAIL_INTERNAL_NOTICE

    resend_key = os.environ.get('RESEND_API_KEY', '')
    from_email = os.environ.get('RESEND_FROM', 'noreply@purepropsolutions.com')
    reply_to = os.environ.get('RESEND_REPLY_TO', '').strip() or session.get('user_email', '')

    if not resend_key:
        return jsonify({'error': 'Email service not configured. Contact Thomas.'}), 500

    type_labels = {
        'proposal': 'Proposal',
        'ppm': 'PPM Checklist',
        'tps': 'Trade Partner Scope',
        'site_visit': 'Site Visit Report',
        'siding_estimate': 'Siding Estimate',
        'roofing_estimate': 'Roofing Estimate',
        'gutter_estimate': 'Gutter Estimate',
        'painting_estimate': 'Exterior Painting Estimate',
    }
    label = type_labels.get(doc_type, 'Document')
    subject = f'PPS {label} — {property_name}' if property_name else f'PPS {label}'
    prop_line = f'Property: {property_name}' if property_name else ''
    text_body = (
        f'PPS {label} attached — internal use.\n\n'
        f'{EMAIL_INTERNAL_NOTICE}\n\n'
        f'{prop_line + chr(10) if prop_line else ""}'
        f'Generated by: {sender_name}\n\n'
        'The Pure Way: Trust. Quality. Results.'
    )
    html_body = f"""
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:0 auto;">
      <div style="background:#004C8C;padding:20px 24px;border-radius:8px 8px 0 0;">
        <p style="color:white;font-size:18px;font-weight:600;margin:0;">Pure Property Solutions</p>
      </div>
      <div style="background:#f8fafc;padding:24px;border:1px solid #e2e8f0;border-top:none;border-radius:0 0 8px 8px;">
        <p style="color:#334155;font-size:15px;"><strong>PPS {label}</strong> attached — <em>internal use only</em>.</p>
        <p style="color:#7c5e10;font-size:13px;background:#fffbeb;border:1px solid #fde68a;border-radius:6px;padding:10px 12px;">
          {EMAIL_INTERNAL_NOTICE}
        </p>
        {'<p style="color:#64748b;font-size:14px;">Property: ' + property_name + '</p>' if property_name else ''}
        <p style="color:#64748b;font-size:14px;">Generated by: {sender_name}</p>
        <hr style="border:none;border-top:1px solid #e2e8f0;margin:20px 0;">
        <p style="color:#94a3b8;font-size:12px;font-style:italic;">
          The Pure Way: Trust. Quality. Results.™
        </p>
      </div>
    </div>
    """
    email_payload = {
        'from': f'Pure Property Solutions <{from_email}>',
        'to': [to_email],
        'subject': subject,
        'html': html_body,
        'text': text_body,
        'tags': [{'name': 'source', 'value': doc_type[:50]}],
        'attachments': [{'filename': filename, 'content': doc_b64}],
    }
    if reply_to:
        email_payload['reply_to'] = [reply_to]

    payload = _json.dumps(email_payload).encode('utf-8')
    try:
        req = _ur.Request(
            'https://api.resend.com/emails',
            data=payload,
            headers={
                'Authorization': f'Bearer {resend_key}',
                'Content-Type': 'application/json',
                'User-Agent': 'PPS-Hub/1.0',
            },
            method='POST',
        )
        resp = _ur.urlopen(req, timeout=30)
        result = _json.loads(resp.read().decode('utf-8'))
        resend_id = result.get('id', '')
        if not resend_id:
            return jsonify({'error': 'Email service accepted the request but returned no message ID.'}), 500
        return jsonify({'success': True, 'sent_to': to_email, 'resend_id': resend_id, 'from_email': from_email})
    except _ur_err.HTTPError as e:
        body = e.read().decode('utf-8', errors='replace')
        try:
            msg = _json.loads(body).get('message', body)
        except Exception:
            msg = body.strip() or str(e)
        return jsonify({'error': msg}), 500
    except Exception as e:
        return _api_error(e)


@app.route('/email-doc', methods=['POST'])
def email_doc():
    """Email a generated document via Resend API."""
    api_key = request.headers.get('X-API-Key', '')
    internal_ok = api_key == INTERNAL_API_KEY
    if not session.get('user_key') and not internal_ok:
        return jsonify({'error': 'Not authenticated'}), 401

    import base64, urllib.request as _ur, urllib.error as _ur_err, json as _json

    if request.is_json:
        data = request.get_json(silent=True) or {}
        to_email  = (data.get('to_email') or '').strip()
        doc_type  = data.get('doc_type', 'document')
        filename  = data.get('filename', 'PPS_Document.docx')
        prop_name = data.get('property_name', '')
        doc_b64   = data.get('doc_base64', '')
        sender    = data.get('sender_name', 'PPS Proposal Tool')
    else:
        to_email  = request.form.get('to_email', '').strip()
        doc_type  = request.form.get('doc_type', 'document')
        filename  = request.form.get('filename', 'PPS_Document.docx')
        prop_name = request.form.get('property_name', '')
        doc_b64   = request.form.get('doc_base64', '')
        sender    = session.get('display_name', 'PPS Hub')

    if not to_email:
        to_email = session.get('user_email', '')
    if not to_email:
        return jsonify({'error': 'No email address available'}), 400
    if not doc_b64:
        return jsonify({'error': 'No document content'}), 400

    return _send_resend_document(
        to_email=to_email,
        doc_type=doc_type,
        filename=filename,
        property_name=prop_name,
        doc_b64=doc_b64,
        sender_name=sender,
    )


# ── CLIENT DATABASE ─────────────────────────────────────────────────────────

@app.route('/api/clients/search')
def clients_search():
    """Search clients by name or company — returns top 10 matches."""
    # Allow cross-origin from proposal tool (server-to-server or CORS)
    api_key = request.headers.get('X-API-Key', '')
    session_ok = session.get('user_key')
    internal_ok = api_key == INTERNAL_API_KEY
    if not session_ok and not internal_ok:
        resp = jsonify({'error': 'Not authenticated'})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 401
    q = request.args.get('q', '').strip()
    if len(q) < 2:
        return jsonify([])
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('''
                SELECT id, name, email, company, property_name, address
                FROM clients
                WHERE LOWER(name) LIKE %s
                   OR LOWER(company) LIKE %s
                   OR LOWER(property_name) LIKE %s
                ORDER BY
                    CASE WHEN LOWER(name) LIKE %s THEN 0 ELSE 1 END,
                    name
                LIMIT 10
            ''', (f'%{q.lower()}%', f'%{q.lower()}%', f'%{q.lower()}%', f'{q.lower()}%'))
            rows = cur.fetchall()
            cur.close(); conn.close()
            resp = jsonify([dict(r) for r in rows])
            resp.headers['Access-Control-Allow-Origin'] = '*'
            resp.headers['Access-Control-Allow-Headers'] = 'X-API-Key, Content-Type'
            return resp
    except Exception as e:
        _log_exception(e, 'clients/search')
        resp = jsonify({'error': GENERIC_API_ERROR})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500
    resp = jsonify([])
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp


@app.route('/api/clients/search', methods=['OPTIONS'])
def clients_search_options():
    resp = jsonify({})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Headers'] = 'X-API-Key, Content-Type'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
    return resp


@app.route('/api/clients/save', methods=['POST', 'OPTIONS'])
def clients_save():
    """Create or update a client record."""
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Headers'] = 'X-API-Key, Content-Type'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        return resp

    data = request.get_json() or {}
    api_key = request.headers.get('X-API-Key', '')
    internal_ok = api_key == INTERNAL_API_KEY

    user_key = session.get('user_key')
    if internal_ok and data.get('user_key'):
        user_key = data.get('user_key')

    if not user_key:
        resp = jsonify({'error': 'Not authenticated'})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 401

    if not can_manage_contacts(user_key):
        resp = jsonify({'error': 'Permission denied'})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 403
    client_id = data.get('id')
    name    = (data.get('name') or '').strip()
    email   = (data.get('email') or '').strip()
    company = (data.get('company') or '').strip()
    prop    = (data.get('property_name') or '').strip()
    address = (data.get('address') or '').strip()
    notes   = (data.get('notes') or '').strip()

    if not name:
        resp = jsonify({'error': 'Name is required'})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 400

    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            if client_id:
                cur.execute('''
                    UPDATE clients SET name=%s, email=%s, company=%s,
                    property_name=%s, address=%s, notes=%s, updated_at=NOW()
                    WHERE id=%s RETURNING id
                ''', (name, email, company, prop, address, notes, client_id))
            else:
                cur.execute('''
                    INSERT INTO clients (name, email, company, property_name, address, notes, added_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id
                ''', (name, email, company, prop, address, notes, user_key))
            row = cur.fetchone()
            conn.commit(); cur.close(); conn.close()
            resp = jsonify({'success': True, 'id': row['id']})
            resp.headers['Access-Control-Allow-Origin'] = '*'
            return resp
    except Exception as e:
        _log_exception(e, 'clients/save')
        resp = jsonify({'error': GENERIC_API_ERROR})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500


@app.route('/api/clients/seed', methods=['POST'])
def clients_seed():
    """Seed clients from JSON — admin only, run once."""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin only'}), 403
    data = request.get_json()
    clients_data = data.get('clients', [])
    inserted = 0
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor()
            for c in clients_data:
                name = (c.get('name') or '').strip()
                if not name: continue
                # Skip if name already exists
                cur.execute('SELECT id FROM clients WHERE LOWER(name) = LOWER(%s)', (name,))
                if cur.fetchone(): continue
                cur.execute('''
                    INSERT INTO clients (name, email, company, property_name, address, added_by)
                    VALUES (%s, %s, %s, %s, %s, %s)
                ''', (name, c.get('email',''), c.get('company',''),
                      c.get('property_name',''), c.get('address',''), 'seed'))
                inserted += 1
            conn.commit(); cur.close(); conn.close()
        return jsonify({'success': True, 'inserted': inserted})
    except Exception as e:
        return _api_error(e)


@app.route('/clients')
@require_login
def clients_page():
    """Client / contacts database — admin, consultants, office, and all PMs."""
    user_key = session.get('user_key')
    if not can_manage_contacts(user_key):
        return redirect(url_for('dashboard'))
    rows = []
    try:
        conn = get_db()
        if conn:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute('SELECT * FROM clients ORDER BY name')
            rows = cur.fetchall()
            cur.close(); conn.close()
    except Exception as e:
        print(f"Clients page error: {e}")
    return render_template('clients.html', rows=rows, can_edit=True)


@app.route('/admin/seed-clients')
def seed_clients_page():
    if session.get('role') != 'admin':
        return redirect(url_for('login'))
    return render_template('seed_clients.html')


@app.route('/test-token')
@require_login
def test_token():
    """Debug endpoint - shows session state and SSO code generation."""
    code = generate_sso_code(
        get_db,
        session['user_key'],
        session.get('display_name', ''),
        session.get('role', 'user'),
    )
    return jsonify({
        'session': 'active',
        'user_key': session['user_key'],
        'code_generated': code is not None,
        'code_preview': code[:8] + '...' if code else None,
    })


def _send_resend_email(to_email, subject, html_body, text_body):
    import urllib.request as _ur
    import urllib.error as _ur_err
    resend_key = os.environ.get('RESEND_API_KEY', '')
    from_email = os.environ.get('RESEND_FROM', 'noreply@purepropsolutions.com')
    if not resend_key:
        return False, 'Email service not configured'
    payload = json.dumps({
        'from': f'Pure Property Solutions <{from_email}>',
        'to': [to_email],
        'subject': subject,
        'html': html_body,
        'text': text_body,
    }).encode('utf-8')
    try:
        req = _ur.Request(
            'https://api.resend.com/emails',
            data=payload,
            headers={
                'Authorization': f'Bearer {resend_key}',
                'Content-Type': 'application/json',
                'User-Agent': 'PPS-Hub/1.0',
            },
            method='POST',
        )
        resp = _ur.urlopen(req, timeout=30)
        result = json.loads(resp.read().decode('utf-8'))
        if result.get('id'):
            return True, result['id']
        return False, 'No message ID returned'
    except _ur_err.HTTPError as e:
        body = e.read().decode('utf-8', errors='replace')
        return False, body
    except Exception as e:
        return False, str(e)


@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    message = None
    error = None
    if request.method == 'POST':
        user_key = _resolve_login_user_key(request.form.get('user_key', ''))
        user_def = USERS.get(user_key)
        if not user_def:
            error = 'Tap your name in the list first.'
        else:
            token = create_password_reset_token(get_db, user_key)
            to_email = user_def.get('email', '')
            if not token or not to_email:
                error = 'Could not start reset. Contact Thomas or Stephanie.'
            else:
                link = reset_url_for_token(token)
                ok, detail = _send_resend_email(
                    to_email.strip().lower(),
                    'Reset your PPS Hub password',
                    f'<p>Hi {user_def["display"].split()[0]},</p>'
                    f'<p><a href="{link}">Click here to reset your PPS Hub password</a>. '
                    f'This link expires in 1 hour.</p>'
                    f'<p>If you did not request this, ignore this email.</p>',
                    f'Reset your PPS password: {link}\n\nThis link expires in 1 hour.',
                )
                if ok:
                    # Reset email path also clears lockout so they aren't blocked
                    # while waiting for the link.
                    clear_login_failures(get_db, user_key)
                    message = (
                        f'If {to_email} is on file, a reset link was sent. '
                        f'Open it on this phone. Any temporary lockout was cleared.'
                    )
                else:
                    print(f'Password reset email failed: {detail}')
                    error = 'Could not send email. Contact Thomas or Stephanie.'
    return _no_store_html(
        'forgot_password.html',
        users=sorted(USERS.items(), key=lambda x: x[1]['display']),
        message=message,
        error=error,
    )


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password_with_token(token):
    error = None
    if request.method == 'GET':
        if not peek_password_reset_token(get_db, token):
            error = 'This reset link is invalid or has expired.'
            return _no_store_html('reset_password.html', error=error, token=None)
        return _no_store_html('reset_password.html', error=None, token=token)

    token = request.form.get('token', '').strip()
    new_password = request.form.get('new_password', '') or ''
    confirm = request.form.get('confirm_password', '') or ''
    user_key = consume_password_reset_token(get_db, token)
    if not user_key:
        error = 'This reset link is invalid or has expired.'
    elif len(new_password) < 8:
        error = 'Password must be at least 8 characters.'
    elif new_password != confirm:
        error = 'Passwords do not match.'
    else:
        try:
            ok, action = _upsert_hub_user_password(user_key, new_password, must_change=False)
            if ok:
                clear_login_failures(get_db, user_key)
                print(f'Password reset via token for {user_key} ({action})')
                session['login_flash_user'] = user_key
                session['login_flash_success'] = (
                    'Password updated. Tap your name and sign in with the new password.'
                )
                session.pop('login_flash_error', None)
                resp = redirect(url_for('login'))
                resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, private'
                return resp
            if action == 'no_db' or str(action).startswith('db_error:'):
                error = 'Database is temporarily unavailable. Wait a minute and try the link again, or contact Thomas.'
            elif action == 'password_too_short':
                error = 'Password must be at least 8 characters.'
            else:
                error = 'Could not save the new password. Contact Thomas or Stephanie.'
            print(f'reset password upsert failed for {user_key}: {action}')
        except Exception as e:
            print(f'reset password error for {user_key}: {e}')
            error = 'Something went wrong saving the password. Try again or contact Thomas.'
    return _no_store_html('reset_password.html', error=error, token=token)


@app.route('/change-password', methods=['GET', 'POST'])
@require_login
def change_password():
    error = None
    if request.method == 'POST':
        current = request.form.get('current_password', '').strip()
        new_password = request.form.get('new_password', '').strip()
        confirm = request.form.get('confirm_password', '').strip()
        user_key = session['user_key']
        if len(new_password) < 8:
            error = 'New password must be at least 8 characters.'
        elif new_password != confirm:
            error = 'Passwords do not match.'
        else:
            try:
                conn = get_db()
                if conn:
                    cur = conn.cursor(cursor_factory=RealDictCursor)
                    cur.execute('SELECT password_hash FROM hub_users WHERE user_key = %s', (user_key,))
                    row = cur.fetchone()
                    if not row or not check_password_hash(row['password_hash'], current):
                        error = 'Current password is incorrect.'
                    else:
                        hashed = generate_password_hash(new_password)
                        cur.execute(
                            'UPDATE hub_users SET password_hash = %s, must_change_password = FALSE, '
                            'password_epoch = COALESCE(password_epoch, 0) + 1 WHERE user_key = %s',
                            (hashed, user_key),
                        )
                        conn.commit()
                        session.pop('must_change_password', None)
                        cur.close()
                        conn.close()
                        return redirect(url_for('dashboard'))
                    cur.close()
                    conn.close()
            except Exception as e:
                print(f'change password error: {e}')
                error = 'Something went wrong. Please try again.'
    return render_template('change_password.html', error=error)


def _parse_siding_pricing_upload(pricing_file):
    from estimators.siding.pricing_parser import parse_pricing_upload
    return parse_pricing_upload(pricing_file)


def _load_siding_estimate_row(estimate_id, user_key=None):
    conn = get_db()
    if not conn:
        return None
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute('SELECT * FROM siding_estimate_log WHERE id = %s', (estimate_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    if not row:
        return None
    if user_key and row.get('generated_by') != user_key and session.get('role') != 'admin':
        return None
    return row


def _siding_job_data_from_row(row):
    data = row.get('job_data') or {}
    if isinstance(data, str):
        data = json.loads(data)
    return data


def _build_siding_excel_from_row(row):
    from estimators.siding import build_estimate_excel
    data = _siding_job_data_from_row(row)
    return build_estimate_excel(
        data.get('job', {}),
        data.get('buildings', []),
        data.get('inputs', {}),
        data.get('pricing', {}),
        library_rows=data.get('library', []),
        confidence=data.get('confidence'),
    )


def _siding_filename(job):
    prop = (job.get('property_name') or 'Property').replace(' ', '_')
    return f'PPS_Siding_Estimate_{prop}.xlsx'


def _siding_preview_context(row):
    from estimators.siding import calculate_quantities, aggregate_building_quantities, compute_price_stack
    from estimators.siding.excel_builder import SOURCE_LABELS

    data = _siding_job_data_from_row(row)
    job = data.get('job', {})
    inputs = data.get('inputs', {})
    buildings = data.get('buildings', [])
    building_rows = []
    for b in buildings:
        qty = max(int(b.get('qty') or 1), 1)
        q = calculate_quantities(b.get('measurements') or {}, inputs, qty=qty)
        building_rows.append({
            'label': b.get('label') or 'Building',
            'building_type': b.get('building_type') or 'A',
            'qty': qty,
            'source_label': SOURCE_LABELS.get(b.get('source'), b.get('source', '')),
            'quantities': q,
        })
    totals = aggregate_building_quantities(building_rows)
    price_stack = data.get('price_stack') or compute_price_stack(buildings, inputs)
    return {
        'job': job,
        'inputs': inputs,
        'building_rows': building_rows,
        'totals': totals,
        'price_stack': price_stack,
    }


@app.route('/siding-estimator')
@require_login
def siding_estimator():
    return render_template(
        'siding_estimator.html',
        display_name=session.get('display_name', ''),
        pricing_defaults=_pricing_defaults(),
    )


@app.route('/siding-estimator/parse', methods=['POST'])
@require_login
def siding_estimator_parse():
    pdf_file = request.files.get('pdf')
    if not pdf_file:
        return jsonify({'error': 'No PDF uploaded'}), 400
    source = (request.form.get('source') or 'eagleview').strip()
    try:
        from estimators.siding import parse_eagleview_walls, parse_aerial_report
        pdf_bytes = pdf_file.read()
        if source == 'eagleview':
            measurements, warnings = parse_eagleview_walls(pdf_bytes)
        else:
            measurements, warnings = parse_aerial_report(pdf_bytes)
        from estimators.reliability import build_siding_reliability
        confidence = build_siding_reliability(measurements, source=source)
        return jsonify({'measurements': measurements, 'warnings': warnings, 'confidence': confidence})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/siding-estimator/generate', methods=['POST'])
@require_login
def siding_estimator_generate():
    import json as _json
    try:
        job = _json.loads(request.form.get('job', '{}'))
        buildings = _json.loads(request.form.get('buildings', '[]'))
        inputs = _json.loads(request.form.get('inputs', '{}'))
        parsed_pricing = _parse_siding_pricing_upload(request.files.get('pricing'))
        pricing = parsed_pricing.get('prices', {})
        library_rows = parsed_pricing.get('library', [])

        if not buildings:
            return jsonify({'error': 'Add at least one building'}), 400

        from estimators.siding import build_estimate_excel, compute_price_stack
        from estimators.reliability import build_siding_job_reliability
        pricing_loaded = parsed_pricing.get('loaded_count', 0)
        confidence = build_siding_job_reliability(buildings, pricing_loaded)
        price_stack = compute_price_stack(buildings, inputs)
        buf = build_estimate_excel(
            job, buildings, inputs, pricing, library_rows=library_rows, confidence=confidence
        )

        user_key = session['user_key']
        display_name = session.get('display_name', '')
        job_data = {
            'job': job,
            'buildings': buildings,
            'inputs': inputs,
            'pricing': pricing,
            'library': library_rows,
            'confidence': confidence,
            'price_stack': price_stack,
            'pricing_meta': {
                'loaded_count': pricing_loaded,
                'warnings': parsed_pricing.get('warnings', []),
            },
        }
        estimate_id = None
        conn = get_db()
        if conn:
            cur = conn.cursor()
            type_n = len(buildings)
            bldg_n = price_stack.get('total_qty') or type_n
            siding_summary = ' · '.join(x for x in [
                f"{type_n} type{'s' if type_n != 1 else ''}",
                f"{bldg_n} building{'s' if bldg_n != 1 else ''}",
                inputs.get('siding_type'),
            ] if x)
            cur.execute(
                '''INSERT INTO siding_estimate_log (
                    generated_by, display_name, property_name, property_address,
                    building_count, siding_type, summary_meta, job_data
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id''',
                (
                    user_key,
                    display_name,
                    job.get('property_name'),
                    job.get('address'),
                    len(buildings),
                    inputs.get('siding_type'),
                    siding_summary[:255],
                    _json.dumps(job_data),
                ),
            )
            estimate_id = cur.fetchone()[0]
            conn.commit()
            cur.close()
            conn.close()

        if not estimate_id:
            return jsonify({'error': 'Could not save estimate to history'}), 500

        return jsonify({
            'success': True,
            'estimate_id': estimate_id,
            'redirect_url': url_for('siding_estimator_result', estimate_id=estimate_id),
            'pricing_loaded': parsed_pricing.get('loaded_count', 0),
            'pricing_warnings': parsed_pricing.get('warnings', []),
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/siding-estimator/result/<int:estimate_id>')
@require_login
def siding_estimator_result(estimate_id):
    row = _load_siding_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _siding_preview_context(row)
    data = _siding_job_data_from_row(row)
    return render_template(
        'siding_result.html',
        estimate_id=estimate_id,
        property_name=row.get('property_name') or 'Property',
        building_count=row.get('building_count') or 1,
        filename=_siding_filename(data.get('job', {})),
        totals=ctx['totals'],
        price_stack=ctx.get('price_stack'),
        confidence=data.get('confidence'),
        user_email=session.get('user_email', ''),
    )


@app.route('/siding-estimator/preview/<int:estimate_id>')
@require_login
def siding_estimator_preview(estimate_id):
    row = _load_siding_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _siding_preview_context(row)
    return render_template(
        'siding_preview.html',
        estimate_id=estimate_id,
        **ctx,
    )


@app.route('/siding-estimator/download/<int:estimate_id>')
@require_login
def siding_estimator_download(estimate_id):
    row = _load_siding_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    try:
        data = _siding_job_data_from_row(row)
        buf = _build_siding_excel_from_row(row)
        return send_file(
            buf,
            as_attachment=True,
            download_name=_siding_filename(data.get('job', {})),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        _log_exception(e)
        return GENERIC_DOWNLOAD_ERROR, 500


@app.route('/siding-estimator/email/<int:estimate_id>', methods=['POST'])
@require_login
def siding_estimator_email(estimate_id):
    row = _load_siding_estimate_row(estimate_id, session['user_key'])
    if not row:
        return jsonify({'error': 'Estimate not found'}), 404
    payload = request.get_json(silent=True) or {}
    to_email = (payload.get('to_email') or session.get('user_email') or '').strip()
    if not to_email:
        return jsonify({'error': 'No email address'}), 400
    try:
        data = _siding_job_data_from_row(row)
        buf = _build_siding_excel_from_row(row)
        doc_b64 = base64.b64encode(buf.getvalue()).decode('ascii')
        return _send_resend_document(
            to_email=to_email,
            doc_type='siding_estimate',
            filename=_siding_filename(data.get('job', {})),
            property_name=data.get('job', {}).get('property_name', ''),
            doc_b64=doc_b64,
            sender_name=session.get('display_name', 'PPS Hub'),
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/siding-estimator/pricing-preview', methods=['POST'])
@require_login
def siding_pricing_preview():
    parsed = _parse_siding_pricing_upload(request.files.get('pricing'))
    return jsonify({
        'loaded_count': parsed.get('loaded_count', 0),
        'warnings': parsed.get('warnings', []),
        'price_count': len(parsed.get('prices', {})),
        'library_count': len(parsed.get('library', [])),
    })


@app.route('/siding-estimator/pricing-template')
@require_login
def siding_pricing_template():
    import csv
    import io
    from flask import Response

    template_rows = [
        ['item_name', 'unit_price', 'qty_per_sq'],
        ['NDX Vinyl Siding', '', ''],
        ['QA Starter', '', ''],
        ['NDX 5/8 J Channel', '', ''],
        ['NDX Inside Corner 3/4', '', ''],
        ["NDX 12' Outside Corner", '', ''],
        ['Roll of Coil Stock', '', ''],
        ['NDX Universal Trim', '', ''],
        ['House Wrap', '', ''],
        ['House Wrap Tape', '', ''],
        ['J Block Uniblock', '', ''],
        ['J Block M Block', '', ''],
        ['Exhaust Vent', '', ''],
        ['Roofing Nails', '', ''],
        ['Cap Nails', '', ''],
        ['', '', ''],
        ['NDX Vinyl Siding (per sq)', '', '1'],
        ['NDX 5/8 J Channel', '5', ''],
        ['QA Starter', '5', ''],
    ]
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerows(template_rows)
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment; filename=PPS_Siding_Pricing_Template.csv'},
    )


def _roofing_job_data_from_row(row):
    data = row.get('job_data') or {}
    if isinstance(data, str):
        data = json.loads(data)
    return data


def _load_roofing_estimate_row(estimate_id, user_key=None):
    conn = get_db()
    if not conn:
        return None
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute('SELECT * FROM roofing_estimate_log WHERE id = %s', (estimate_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    if not row:
        return None
    if user_key and row.get('generated_by') != user_key and session.get('role') != 'admin':
        return None
    return row


def _roofing_preview_context(row):
    from estimators.roofing import calculate_materials, calculate_bid_summary
    from estimators.roofing.excel_builder import REPORT_LABELS
    from estimators.roofing.material_catalog import MATERIAL_LINES

    data = _roofing_job_data_from_row(row)
    measurements = data.get('measurements', {})
    inputs = data.get('inputs', {})
    report_type = measurements.get('report_type', 'premium')
    is_quick = report_type == 'bid_perfect'
    ctx = {
        'job': data.get('job', {}),
        'inputs': inputs,
        'measurements': measurements,
        'report_label': REPORT_LABELS.get(report_type, report_type),
        'is_quick_bid': is_quick,
    }
    if is_quick:
        ctx['summary'] = calculate_bid_summary(measurements, inputs)
    else:
        qty = calculate_materials(measurements, inputs)
        ctx['summary'] = qty
        ctx['material_lines'] = [
            {'label': label, 'qty': qty.get(qk, 0)}
            for _k, label, _u, qk in MATERIAL_LINES
        ]
    return ctx


def _roofing_filename(job):
    prop = (job.get('property_name') or 'Property').replace(' ', '_')
    return f'PPS_Roof_Estimate_{prop}.xlsx'


@app.route('/roofing-estimator')
@require_login
def roofing_estimator():
    return render_template(
        'roofing_estimator.html',
        display_name=session.get('display_name', ''),
        pricing_defaults=_pricing_defaults(),
    )


@app.route('/roofing-estimator/parse', methods=['POST'])
@require_login
def roofing_estimator_parse():
    pdf_file = request.files.get('pdf')
    if not pdf_file:
        return jsonify({'error': 'No PDF uploaded'}), 400
    try:
        from estimators.roofing import parse_roof_report
        from estimators.reliability import build_roofing_reliability
        measurements, warnings = parse_roof_report(pdf_file.read())
        confidence = build_roofing_reliability(measurements, {})
        return jsonify({'measurements': measurements, 'warnings': warnings, 'confidence': confidence})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/roofing-estimator/generate', methods=['POST'])
@require_login
def roofing_estimator_generate():
    import json as _json
    try:
        job = _json.loads(request.form.get('job', '{}'))
        measurements = _json.loads(request.form.get('measurements', '{}'))
        inputs = _json.loads(request.form.get('inputs', '{}'))

        if not measurements.get('roof_area_sqft') and not measurements.get('structures'):
            return jsonify({'error': 'No roof measurements — upload a valid report PDF.'}), 400

        from estimators.roofing import build_estimate_excel
        from estimators.reliability import build_roofing_reliability
        confidence = build_roofing_reliability(measurements, inputs)
        buf = build_estimate_excel(job, measurements, inputs, {}, confidence=confidence)

        job_data = {
            'job': job,
            'measurements': measurements,
            'inputs': inputs,
            'confidence': confidence,
        }
        estimate_id = None
        conn = get_db()
        if conn:
            cur = conn.cursor()
            roofing_summary = measurements.get('report_type') or 'report'
            if measurements.get('roof_area_squares'):
                roofing_summary = f"{roofing_summary} · {measurements.get('roof_area_squares')} sq"
            cur.execute(
                '''INSERT INTO roofing_estimate_log (
                    generated_by, display_name, property_name, property_address,
                    report_type, summary_meta, job_data
                ) VALUES (%s,%s,%s,%s,%s,%s,%s) RETURNING id''',
                (
                    session['user_key'],
                    session.get('display_name', ''),
                    job.get('property_name'),
                    job.get('address'),
                    measurements.get('report_type'),
                    roofing_summary[:255],
                    _json.dumps(job_data),
                ),
            )
            estimate_id = cur.fetchone()[0]
            conn.commit()
            cur.close()
            conn.close()

        if not estimate_id:
            return jsonify({'error': 'Could not save estimate to history'}), 500

        return jsonify({
            'success': True,
            'estimate_id': estimate_id,
            'redirect_url': url_for('roofing_estimator_result', estimate_id=estimate_id),
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/roofing-estimator/result/<int:estimate_id>')
@require_login
def roofing_estimator_result(estimate_id):
    row = _load_roofing_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _roofing_preview_context(row)
    data = _roofing_job_data_from_row(row)
    return render_template(
        'roofing_result.html',
        estimate_id=estimate_id,
        property_name=row.get('property_name') or 'Property',
        filename=_roofing_filename(data.get('job', {})),
        report_label=ctx['report_label'],
        is_quick_bid=ctx['is_quick_bid'],
        summary=ctx['summary'],
        confidence=data.get('confidence'),
        user_email=session.get('user_email', ''),
    )


@app.route('/roofing-estimator/preview/<int:estimate_id>')
@require_login
def roofing_estimator_preview(estimate_id):
    row = _load_roofing_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _roofing_preview_context(row)
    return render_template('roofing_preview.html', estimate_id=estimate_id, **ctx)


@app.route('/roofing-estimator/download/<int:estimate_id>')
@require_login
def roofing_estimator_download(estimate_id):
    row = _load_roofing_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    try:
        from estimators.roofing import build_estimate_excel
        data = _roofing_job_data_from_row(row)
        buf = build_estimate_excel(
            data.get('job', {}),
            data.get('measurements', {}),
            data.get('inputs', {}),
            data.get('pricing', {}),
            confidence=data.get('confidence'),
        )
        return send_file(
            buf,
            as_attachment=True,
            download_name=_roofing_filename(data.get('job', {})),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        _log_exception(e)
        return GENERIC_DOWNLOAD_ERROR, 500


@app.route('/roofing-estimator/email/<int:estimate_id>', methods=['POST'])
@require_login
def roofing_estimator_email(estimate_id):
    row = _load_roofing_estimate_row(estimate_id, session['user_key'])
    if not row:
        return jsonify({'error': 'Estimate not found'}), 404
    payload = request.get_json(silent=True) or {}
    to_email = (payload.get('to_email') or session.get('user_email') or '').strip()
    if not to_email:
        return jsonify({'error': 'No email address'}), 400
    try:
        from estimators.roofing import build_estimate_excel
        data = _roofing_job_data_from_row(row)
        buf = build_estimate_excel(
            data.get('job', {}),
            data.get('measurements', {}),
            data.get('inputs', {}),
            data.get('pricing', {}),
            confidence=data.get('confidence'),
        )
        doc_b64 = base64.b64encode(buf.getvalue()).decode('ascii')
        return _send_resend_document(
            to_email=to_email,
            doc_type='roofing_estimate',
            filename=_roofing_filename(data.get('job', {})),
            property_name=data.get('job', {}).get('property_name', ''),
            doc_b64=doc_b64,
            sender_name=session.get('display_name', 'PPS Hub'),
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


def _gutter_job_data_from_row(row):
    data = row.get('job_data') or {}
    if isinstance(data, str):
        data = json.loads(data)
    return data


def _load_gutter_estimate_row(estimate_id, user_key=None):
    conn = get_db()
    if not conn:
        return None
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute('SELECT * FROM gutter_estimate_log WHERE id = %s', (estimate_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    if not row:
        return None
    if user_key and row.get('generated_by') != user_key and session.get('role') != 'admin':
        return None
    return row


def _gutter_preview_context(row):
    from estimators.gutter import calculate_gutter_estimate

    data = _gutter_job_data_from_row(row)
    summary = calculate_gutter_estimate(
        data.get('measurements', {}),
        data.get('inputs', {}),
    )
    return {
        'job': data.get('job', {}),
        'inputs': data.get('inputs', {}),
        'measurements': data.get('measurements', {}),
        'summary': summary,
    }


def _gutter_filename(job):
    prop = (job.get('property_name') or 'Property').replace(' ', '_')
    return f'PPS_Gutter_Estimate_{prop}.xlsx'


@app.route('/gutter-estimator')
@require_login
def gutter_estimator():
    return render_template(
        'gutter_estimator.html',
        display_name=session.get('display_name', ''),
        pricing_defaults=_pricing_defaults(),
    )


@app.route('/gutter-estimator/parse', methods=['POST'])
@require_login
def gutter_estimator_parse():
    pdf_file = request.files.get('pdf')
    if not pdf_file:
        return jsonify({'error': 'No PDF uploaded'}), 400
    try:
        from estimators.gutter import parse_gutter_measurements
        from estimators.reliability import build_gutter_reliability
        measurements, warnings = parse_gutter_measurements(pdf_file.read())
        confidence = build_gutter_reliability(measurements, {})
        return jsonify({'measurements': measurements, 'warnings': warnings, 'confidence': confidence})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/gutter-estimator/generate', methods=['POST'])
@require_login
def gutter_estimator_generate():
    import json as _json
    try:
        job = _json.loads(request.form.get('job', '{}'))
        measurements = _json.loads(request.form.get('measurements', '{}'))
        inputs = _json.loads(request.form.get('inputs', '{}'))

        gutter_lf = float(measurements.get('gutter_lf') or measurements.get('eaves_ft') or 0)
        if not gutter_lf:
            return jsonify({'error': 'Enter gutter run (LF) or upload a report with eaves length.'}), 400

        from estimators.gutter import build_estimate_excel, calculate_gutter_estimate
        from estimators.reliability import build_gutter_reliability
        user_overrides = inputs.pop('user_overrides', None) or {}
        calc = calculate_gutter_estimate(measurements, inputs)
        confidence = build_gutter_reliability(measurements, inputs, user_overrides)
        buf = build_estimate_excel(job, measurements, inputs, confidence=confidence)

        job_data = {
            'job': job,
            'measurements': measurements,
            'inputs': inputs,
            'confidence': confidence,
        }
        estimate_id = None
        conn = get_db()
        if conn:
            cur = conn.cursor()
            gutter_summary = f"{float(calc.get('gutter_lf_raw') or 0):.0f} LF"
            cur.execute(
                '''INSERT INTO gutter_estimate_log (
                    generated_by, display_name, property_name, property_address,
                    gutter_lf, summary_meta, job_data
                ) VALUES (%s,%s,%s,%s,%s,%s,%s) RETURNING id''',
                (
                    session['user_key'],
                    session.get('display_name', ''),
                    job.get('property_name'),
                    job.get('address'),
                    calc.get('gutter_lf_raw'),
                    gutter_summary[:255],
                    _json.dumps(job_data),
                ),
            )
            estimate_id = cur.fetchone()[0]
            conn.commit()
            cur.close()
            conn.close()

        if not estimate_id:
            return jsonify({'error': 'Could not save estimate to history'}), 500

        return jsonify({
            'success': True,
            'estimate_id': estimate_id,
            'redirect_url': url_for('gutter_estimator_result', estimate_id=estimate_id),
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/gutter-estimator/preview/<int:estimate_id>')
@require_login
def gutter_estimator_preview(estimate_id):
    row = _load_gutter_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _gutter_preview_context(row)
    data = _gutter_job_data_from_row(row)
    return render_template(
        'gutter_preview.html',
        estimate_id=estimate_id,
        confidence=data.get('confidence'),
        **ctx,
    )


@app.route('/gutter-estimator/result/<int:estimate_id>')
@require_login
def gutter_estimator_result(estimate_id):
    row = _load_gutter_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _gutter_preview_context(row)
    data = _gutter_job_data_from_row(row)
    return render_template(
        'gutter_result.html',
        estimate_id=estimate_id,
        property_name=row.get('property_name') or 'Property',
        summary=ctx['summary'],
        confidence=data.get('confidence'),
        user_email=session.get('user_email', ''),
    )


@app.route('/gutter-estimator/download/<int:estimate_id>')
@require_login
def gutter_estimator_download(estimate_id):
    row = _load_gutter_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    try:
        from estimators.gutter import build_estimate_excel
        data = _gutter_job_data_from_row(row)
        buf = build_estimate_excel(
            data.get('job', {}),
            data.get('measurements', {}),
            data.get('inputs', {}),
            confidence=data.get('confidence'),
        )
        return send_file(
            buf,
            as_attachment=True,
            download_name=_gutter_filename(data.get('job', {})),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        _log_exception(e)
        return GENERIC_DOWNLOAD_ERROR, 500


@app.route('/gutter-estimator/email/<int:estimate_id>', methods=['POST'])
@require_login
def gutter_estimator_email(estimate_id):
    row = _load_gutter_estimate_row(estimate_id, session['user_key'])
    if not row:
        return jsonify({'error': 'Estimate not found'}), 404
    payload = request.get_json(silent=True) or {}
    to_email = (payload.get('to_email') or session.get('user_email') or '').strip()
    if not to_email:
        return jsonify({'error': 'No email address'}), 400
    try:
        from estimators.gutter import build_estimate_excel
        data = _gutter_job_data_from_row(row)
        buf = build_estimate_excel(
            data.get('job', {}),
            data.get('measurements', {}),
            data.get('inputs', {}),
            confidence=data.get('confidence'),
        )
        doc_b64 = base64.b64encode(buf.getvalue()).decode('ascii')
        return _send_resend_document(
            to_email=to_email,
            doc_type='gutter_estimate',
            filename=_gutter_filename(data.get('job', {})),
            property_name=data.get('job', {}).get('property_name', ''),
            doc_b64=doc_b64,
            sender_name=session.get('display_name', 'PPS Hub'),
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


def _painting_job_data_from_row(row):
    data = row.get('job_data') or {}
    if isinstance(data, str):
        data = json.loads(data)
    return data


def _load_painting_estimate_row(estimate_id, user_key=None):
    conn = get_db()
    if not conn:
        return None
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute('SELECT * FROM painting_estimate_log WHERE id = %s', (estimate_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    if not row:
        return None
    if user_key and row.get('generated_by') != user_key and session.get('role') != 'admin':
        return None
    return row


def _painting_preview_context(row):
    from estimators.painting import calculate_painting_estimate

    data = _painting_job_data_from_row(row)
    summary = calculate_painting_estimate(
        data.get('line_items', []),
        data.get('inputs', {}),
    )
    return {
        'job': data.get('job', {}),
        'inputs': data.get('inputs', {}),
        'measurements': data.get('measurements', {}),
        'line_items': data.get('line_items', []),
        'summary': summary,
    }


def _painting_filename(job):
    prop = (job.get('property_name') or 'Property').replace(' ', '_')
    return f'PPS_Painting_Estimate_{prop}.xlsx'


@app.route('/painting-estimator')
@require_login
def painting_estimator():
    from estimators.painting import sections_for_ui
    return render_template(
        'painting_estimator.html',
        display_name=session.get('display_name', ''),
        pricing_defaults=_pricing_defaults(),
        sections=sections_for_ui(),
    )


@app.route('/painting-estimator/generate', methods=['POST'])
@require_login
def painting_estimator_generate():
    import json as _json
    try:
        job = _json.loads(request.form.get('job', '{}'))
        measurements = _json.loads(request.form.get('measurements', '{}'))
        line_items = _json.loads(request.form.get('line_items', '[]'))
        inputs = _json.loads(request.form.get('inputs', '{}'))

        active = [li for li in line_items if li.get('measured')]
        if not active:
            return jsonify({'error': 'Enter at least one measured quantity in the takeoff.'}), 400

        from estimators.painting import build_estimate_excel, calculate_painting_estimate
        from estimators.reliability import build_painting_reliability
        user_overrides = inputs.pop('user_overrides', None) or {}
        calc = calculate_painting_estimate(active, inputs)
        confidence = build_painting_reliability(measurements, active, user_overrides)
        buf = build_estimate_excel(job, active, inputs, confidence=confidence)

        job_data = {
            'job': job,
            'measurements': measurements,
            'line_items': active,
            'inputs': inputs,
            'confidence': confidence,
        }
        estimate_id = None
        conn = get_db()
        if conn:
            cur = conn.cursor()
            painting_parts = []
            if calc.get('line_count'):
                painting_parts.append(f"{calc.get('line_count')} lines")
            if calc.get('one_coat_bid'):
                painting_parts.append(f"${float(calc.get('one_coat_bid')):,.0f} 1-coat")
            painting_summary = ' · '.join(painting_parts)
            cur.execute(
                '''INSERT INTO painting_estimate_log (
                    generated_by, display_name, property_name, property_address,
                    line_count, one_coat_bid, two_coat_bid, summary_meta, job_data
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id''',
                (
                    session['user_key'],
                    session.get('display_name', ''),
                    job.get('property_name'),
                    job.get('address'),
                    calc.get('line_count'),
                    calc.get('one_coat_bid'),
                    calc.get('two_coat_bid'),
                    painting_summary[:255],
                    _json.dumps(job_data),
                ),
            )
            estimate_id = cur.fetchone()[0]
            conn.commit()
            cur.close()
            conn.close()

        if not estimate_id:
            return jsonify({'error': 'Could not save estimate to history'}), 500

        return jsonify({
            'success': True,
            'estimate_id': estimate_id,
            'redirect_url': url_for('painting_estimator_result', estimate_id=estimate_id),
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.route('/painting-estimator/preview/<int:estimate_id>')
@require_login
def painting_estimator_preview(estimate_id):
    row = _load_painting_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _painting_preview_context(row)
    data = _painting_job_data_from_row(row)
    return render_template(
        'painting_preview.html',
        estimate_id=estimate_id,
        confidence=data.get('confidence'),
        **ctx,
    )


@app.route('/painting-estimator/result/<int:estimate_id>')
@require_login
def painting_estimator_result(estimate_id):
    row = _load_painting_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    ctx = _painting_preview_context(row)
    data = _painting_job_data_from_row(row)
    return render_template(
        'painting_result.html',
        estimate_id=estimate_id,
        property_name=row.get('property_name') or 'Property',
        summary=ctx['summary'],
        confidence=data.get('confidence'),
        user_email=session.get('user_email', ''),
    )


@app.route('/painting-estimator/download/<int:estimate_id>')
@require_login
def painting_estimator_download(estimate_id):
    row = _load_painting_estimate_row(estimate_id, session['user_key'])
    if not row:
        return 'Estimate not found', 404
    try:
        from estimators.painting import build_estimate_excel
        data = _painting_job_data_from_row(row)
        buf = build_estimate_excel(
            data.get('job', {}),
            data.get('line_items', []),
            data.get('inputs', {}),
            confidence=data.get('confidence'),
        )
        return send_file(
            buf,
            as_attachment=True,
            download_name=_painting_filename(data.get('job', {})),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
    except Exception as e:
        _log_exception(e)
        return GENERIC_DOWNLOAD_ERROR, 500


@app.route('/painting-estimator/email/<int:estimate_id>', methods=['POST'])
@require_login
def painting_estimator_email(estimate_id):
    row = _load_painting_estimate_row(estimate_id, session['user_key'])
    if not row:
        return jsonify({'error': 'Estimate not found'}), 404
    payload = request.get_json(silent=True) or {}
    to_email = (payload.get('to_email') or session.get('user_email') or '').strip()
    if not to_email:
        return jsonify({'error': 'No email address'}), 400
    try:
        from estimators.painting import build_estimate_excel
        data = _painting_job_data_from_row(row)
        buf = build_estimate_excel(
            data.get('job', {}),
            data.get('line_items', []),
            data.get('inputs', {}),
            confidence=data.get('confidence'),
        )
        doc_b64 = base64.b64encode(buf.getvalue()).decode('ascii')
        return _send_resend_document(
            to_email=to_email,
            doc_type='painting_estimate',
            filename=_painting_filename(data.get('job', {})),
            property_name=data.get('job', {}).get('property_name', ''),
            doc_b64=doc_b64,
            sender_name=session.get('display_name', 'PPS Hub'),
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _api_error(e)


@app.errorhandler(HTTPException)
def _handle_http_exception(e):
    if _wants_json_response():
        message = e.description or GENERIC_API_ERROR
        payload = {'success': False, 'error': message}
        if request.path.startswith('/api/'):
            payload = {'success': False, 'error': message}
        return jsonify(payload), e.code
    return e


@app.errorhandler(Exception)
def _handle_uncaught_exception(e):
    if isinstance(e, HTTPException):
        return e
    _log_exception(e, request.path)
    if _wants_json_response():
        payload = {'success': False, 'error': GENERIC_API_ERROR}
        if request.path.startswith('/api/'):
            payload = {'success': False, 'error': GENERIC_API_ERROR}
        return jsonify(payload), 500
    return GENERIC_API_ERROR, 500


@app.route('/logout')
def logout():
    import urllib.parse
    session.clear()
    final = f'{HUB_PUBLIC_URL}/login'
    if PROPOSAL_URL:
        nxt = urllib.parse.quote(final, safe='')
        return redirect(f'{PROPOSAL_URL}/logout?next={nxt}')
    return redirect(url_for('login'))


def _report_password_campaign(result):
    """Tell Thomas what actually happened, unambiguously.

    The first version of this said "Skipped (excluded or no email)" for people
    who were really *already done*, and reported a failed database read
    identically to a completed run — both showed 0 emailed and everyone
    skipped. In a tool whose entire job is not emailing anyone twice, "nothing
    happened" and "nothing needed to happen" have to look different.
    """
    if not result:
        return
    try:
        if result.get('db_error'):
            _send_digest_email(
                'Hub password campaign COULD NOT RUN',
                'The campaign could not read the database, so nothing was done and\n'
                'nobody was emailed. No passwords changed. Run it again:\n\n'
                '  POST /admin/run-password-campaign\n',
                None, _hub_notify_recipients())
            return

        emailed = result.get('emailed') or []
        failed = result.get('email_failed') or []
        done = result.get('already_done') or []
        excluded = result.get('excluded') or []
        remaining = result.get('remaining') or 0

        if not emailed and not failed and not remaining:
            subject = 'Hub passwords — nothing left to do'
        elif failed:
            subject = f'Hub passwords reset — {len(failed)} need attention'
        else:
            subject = 'Hub passwords reset'

        def names(keys):
            return [f'  - {USERS.get(k, {}).get("display", k)}' for k in keys]

        lines = [f"Shared password retired ({result.get('campaign_id')}).", '']
        if emailed:
            lines += [f'Reset link sent, old password now dead: {len(emailed)}'] + names(emailed)
        if done:
            lines += ['',
                      f'Already had their reset — not emailed again: {len(done)}'] + names(done)
        if failed:
            lines += ['', 'COULD NOT EMAIL — these people KEEP their old password until they',
                      'next sign in. Reset them by hand at /admin:']
            for k in failed:
                u = USERS.get(k, {})
                lines.append(f'  - {u.get("display", k)} ({u.get("email") or "no email"})')
        if excluded:
            lines += ['', 'Excluded or no email on file: ' + ', '.join(excluded)]
        if remaining:
            lines += ['', f'{remaining} still pending — the time budget ran out.',
                      'Run POST /admin/run-password-campaign again to finish them.']
        if not emailed and not failed and not remaining:
            lines += ['', 'Everyone is done. Nothing was sent and nothing changed.']

        _send_digest_email(subject, '\n'.join(lines), None, _hub_notify_recipients())
    except Exception as e:
        print(f'password campaign report failed: {e}')


@app.route('/admin/run-password-campaign', methods=['POST'])
@require_admin
def _run_password_campaign():
    """Retire the shared password — owner-only, triggered by hand.

    This ran at import on 2026-08-21 and took the deploy down: gunicorn was
    still importing app.py while it worked through thirteen synchronous SMTP
    sends, so Render's port scan gave up first ("No open ports detected") and
    cancelled the release. Reset emails had already gone out by then, against
    the shared database, which is why run_campaign is now idempotent per person
    rather than gated on a single claim — see password_campaign.py.

    Never move this back onto the import path. Nothing that talks to SMTP or
    loops over the roster belongs in module scope; the app has to bind a port
    first. Safe to POST repeatedly: it only touches accounts still untouched,
    and stops inside a time budget, reporting how many are left.
    """
    try:
        result = password_campaign.run_campaign(
            get_db,
            USERS,
            _send_digest_email,
            lambda key, ttl: create_password_reset_token(get_db, key, ttl_hours=ttl),
            reset_url_for_token,
            lambda pw: generate_password_hash(pw, method='pbkdf2:sha256'),
            # Nobody is excluded, Thomas included (his call, 2026-08-21).
            # Leaving the owner account out would have left /admin, the vault
            # and every financial page on the one password the campaign exists
            # to retire. He is not exposed to a lockout by this: run_campaign
            # only invalidates a password after its reset email actually sends,
            # so a mail failure leaves him on his current password rather than
            # stranding him.
            exclude=(),
        )
        _report_password_campaign(result)
        return jsonify(result)
    except Exception as e:
        print(f'password campaign error: {e}')
        import traceback
        traceback.print_exc()
        return _api_error(e, ok=False)



ask_pps.register_routes(app, get_db, USERS, CLAUDE_API_KEY, CLAUDE_MODEL, require_login)
pipeline_board.register_routes(app, get_db, USERS, require_login)
office_ops.register_routes(app, get_db, USERS, require_login, send_email_fn=_send_digest_email,
                            claude_api_key=CLAUDE_API_KEY, claude_model=CLAUDE_MODEL)


if __name__ == '__main__':
    app.run(debug=True)
